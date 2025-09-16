"""
AI-Powered Test Case Generator for Healthcare/MedTech
RAG Pipeline with Gemini API - NASSCOM Hackathon
Version 6: AI-Powered Test Import with Intelligent Schema Analysis
"""

import streamlit as st
import os
import json
import uuid
import numpy as np
import faiss
import yaml
import glob
import re
import io
import tempfile
import shutil
from pathlib import Path
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from dotenv import load_dotenv
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import pandas as pd
import PyPDF2
from docx import Document
import xml.etree.ElementTree as ET
import chardet

# Load environment variables
load_dotenv()

# Configuration
DOCUMENTS_PATH = "data/documents"
UPLOADED_DOCS_PATH = "data/uploaded_documents"
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2'
GEMINI_MODEL_NAME = 'gemini-2.0-flash-exp'
OUTPUT_DIR = "data/generated_test_cases"

# NASSCOM Compliance Requirements
NASSCOM_REQUIREMENTS = {
    "document_types": {
        "prd": "Product Requirement Document",
        "user_story": "User Stories with acceptance criteria",
        "api_spec": "API Specifications (OpenAPI/Swagger)",
        "mockup": "UI/UX Mockups and wireframes",
        "test_plan": "Existing test plans or strategies"
    },
    "required_fields": [
        "Test Case ID", "Title", "Description", 
        "Test Steps", "Expected Results", "Priority"
    ],
    "compliance_standards": [
        "HIPAA", "GDPR", "FDA", "ISO 9001", 
        "ISO 13485", "ISO 27001"
    ],
    "healthcare_keywords": [
        "patient", "medical", "health", "clinical", "diagnosis",
        "treatment", "medication", "prescription", "PHI", "EHR",
        "DICOM", "HL7", "medical device", "FDA compliance"
    ],
    "document_formats": ["pdf", "docx", "doc", "xml", "yaml", "yml", "txt", "md"],
    "quality_criteria": {
        "min_content_length": 100,
        "requires_structure": True,
        "requires_healthcare_context": False  # Made optional for flexibility
    }
}

# Page configuration
st.set_page_config(
    page_title="AI Test Case Generator - NASSCOM", 
    page_icon="üß™",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 10px;
        border-radius: 5px;
        margin: 5px;
    }
    .compliance-pass {
        background-color: #d4edda;
        color: #155724;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
    .compliance-fail {
        background-color: #f8d7da;
        color: #721c24;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
    .compliance-warning {
        background-color: #fff3cd;
        color: #856404;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'test_counter' not in st.session_state:
    st.session_state.test_counter = 0
if 'generated_tests' not in st.session_state:
    st.session_state.generated_tests = []
if 'refinement_mode' not in st.session_state:
    st.session_state.refinement_mode = False
if 'test_to_refine' not in st.session_state:
    st.session_state.test_to_refine = None
if 'refinement_history' not in st.session_state:
    st.session_state.refinement_history = {}
if 'edited_test' not in st.session_state:
    st.session_state.edited_test = None
if 'test_versions' not in st.session_state:
    st.session_state.test_versions = {}
if 'uploaded_docs' not in st.session_state:
    st.session_state.uploaded_docs = []
if 'compliance_reports' not in st.session_state:
    st.session_state.compliance_reports = {}
if 'imported_tests' not in st.session_state:
    st.session_state.imported_tests = []
if 'import_history' not in st.session_state:
    st.session_state.import_history = []
if 'last_converted_tests' not in st.session_state:
    st.session_state.last_converted_tests = None
if 'last_import_report' not in st.session_state:
    st.session_state.last_import_report = None

# Configure Gemini
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from Word document"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting Word document text: {e}")
        return ""

def extract_text_from_xml(file_content: bytes) -> str:
    """Extract text content from XML file"""
    try:
        # Try to detect encoding
        detected = chardet.detect(file_content)
        encoding = detected['encoding'] or 'utf-8'
        
        text_content = file_content.decode(encoding)
        root = ET.fromstring(text_content)
        
        # Convert XML to readable format
        def extract_element_text(element, prefix=""):
            result = []
            if element.tag:
                result.append(f"{prefix}{element.tag}:")
            if element.text and element.text.strip():
                result.append(f"{prefix}  {element.text.strip()}")
            for attr_name, attr_value in element.attrib.items():
                result.append(f"{prefix}  @{attr_name}: {attr_value}")
            for child in element:
                result.extend(extract_element_text(child, prefix + "  "))
            return result
        
        text_lines = extract_element_text(root)
        return "\n".join(text_lines)
    except Exception as e:
        st.error(f"Error extracting XML text: {e}")
        return ""

def analyze_document_compliance(content: str, filename: str, file_type: str) -> Dict:
    """
    Analyze document for NASSCOM compliance using Gemini AI
    """
    prompt = f"""You are a compliance analyst for NASSCOM's AI Test Case Generation requirements.
Analyze the following document and determine its compliance with NASSCOM guidelines.

DOCUMENT NAME: {filename}
FILE TYPE: {file_type}

DOCUMENT CONTENT:
{content[:3000]}  # Limiting to first 3000 chars for analysis

NASSCOM REQUIREMENTS:
1. Document Types Accepted: PRDs, User Stories, API Specs, Mockups, Test Plans
2. Healthcare/MedTech Context: Should contain healthcare-related terminology
3. Structure: Should have clear sections and be well-organized
4. Content Quality: Sufficient detail for test case generation
5. Compliance Standards: Should reference or be compatible with HIPAA, GDPR, FDA, ISO standards

Analyze and return a JSON object with:
{{
    "is_compliant": boolean,
    "compliance_score": float (0-100),
    "document_type": string (prd/user_story/api_spec/mockup/test_plan/other),
    "detected_standards": [list of compliance standards found],
    "healthcare_relevance": float (0-100),
    "structural_quality": float (0-100),
    "content_completeness": float (0-100),
    "strengths": [list of compliance strengths],
    "improvements_needed": [list of improvements needed],
    "recommendations": [list of specific recommendations],
    "summary": string (brief compliance summary)
}}"""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        compliance_report = json.loads(response.text)
        
        # Add metadata
        compliance_report['analyzed_at'] = datetime.now().isoformat()
        compliance_report['filename'] = filename
        compliance_report['file_type'] = file_type
        compliance_report['content_length'] = len(content)
        
        return compliance_report
        
    except Exception as e:
        # Fallback basic analysis
        return {
            "is_compliant": len(content) > 100,
            "compliance_score": 50.0 if len(content) > 100 else 20.0,
            "document_type": "other",
            "detected_standards": [],
            "healthcare_relevance": 0.0,
            "structural_quality": 50.0,
            "content_completeness": 50.0,
            "strengths": ["Document uploaded successfully"],
            "improvements_needed": ["Could not perform AI analysis"],
            "recommendations": ["Please ensure document contains relevant content"],
            "summary": f"Basic analysis completed. Error in AI analysis: {str(e)}",
            "analyzed_at": datetime.now().isoformat(),
            "filename": filename,
            "file_type": file_type,
            "content_length": len(content)
        }

def process_uploaded_file(uploaded_file) -> Tuple[str, Dict]:
    """
    Process uploaded file and extract content with compliance analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    
    # Read file content
    file_content = uploaded_file.read()
    
    # Extract text based on file type
    if file_extension == 'pdf':
        text_content = extract_text_from_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        text_content = extract_text_from_docx(file_content)
    elif file_extension == 'xml':
        text_content = extract_text_from_xml(file_content)
    elif file_extension in ['yaml', 'yml']:
        try:
            text_content = file_content.decode('utf-8')
            yaml_data = yaml.safe_load(text_content)
            text_content = yaml.dump(yaml_data, default_flow_style=False)
        except:
            text_content = file_content.decode('utf-8', errors='ignore')
    elif file_extension in ['txt', 'md']:
        text_content = file_content.decode('utf-8', errors='ignore')
    else:
        # Try to decode as text
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
        except:
            text_content = ""
    
    # Analyze compliance
    compliance_report = analyze_document_compliance(
        text_content, 
        uploaded_file.name, 
        file_extension
    )
    
    return text_content, compliance_report

def save_uploaded_document(filename: str, content: str, compliance_report: Dict) -> bool:
    """
    Save uploaded document to the documents folder if compliant
    """
    try:
        # Create uploaded documents directory if it doesn't exist
        os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
        
        # Save the document
        file_path = os.path.join(UPLOADED_DOCS_PATH, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Save compliance report
        report_path = os.path.join(UPLOADED_DOCS_PATH, f"{filename}.compliance.json")
        with open(report_path, 'w') as f:
            json.dump(compliance_report, f, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Error saving document: {e}")
        return False

def display_compliance_report(report: Dict):
    """
    Display compliance analysis report with visual indicators
    """
    # Overall compliance status
    if report['is_compliant']:
        st.markdown('<div class="compliance-pass">‚úÖ Document is NASSCOM Compliant</div>', 
                   unsafe_allow_html=True)
    else:
        st.markdown('<div class="compliance-fail">‚ùå Document needs improvements for full compliance</div>', 
                   unsafe_allow_html=True)
    
    # Compliance metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        score_color = "üü¢" if report['compliance_score'] >= 70 else "üü°" if report['compliance_score'] >= 40 else "üî¥"
        st.metric("Compliance Score", f"{score_color} {report['compliance_score']:.1f}%")
    
    with col2:
        relevance_color = "üü¢" if report['healthcare_relevance'] >= 60 else "üü°" if report['healthcare_relevance'] >= 30 else "üî¥"
        st.metric("Healthcare Relevance", f"{relevance_color} {report['healthcare_relevance']:.1f}%")
    
    with col3:
        quality_color = "üü¢" if report['structural_quality'] >= 70 else "üü°" if report['structural_quality'] >= 40 else "üî¥"
        st.metric("Structural Quality", f"{quality_color} {report['structural_quality']:.1f}%")
    
    with col4:
        completeness_color = "üü¢" if report['content_completeness'] >= 70 else "üü°" if report['content_completeness'] >= 40 else "üî¥"
        st.metric("Content Completeness", f"{completeness_color} {report['content_completeness']:.1f}%")
    
    # Document type and standards
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### üìÑ Document Analysis")
        st.write(f"**Document Type:** {report.get('document_type', 'Unknown').replace('_', ' ').title()}")
        st.write(f"**File Type:** {report.get('file_type', 'Unknown').upper()}")
        st.write(f"**Content Length:** {report.get('content_length', 0):,} characters")
        
        if report.get('detected_standards'):
            st.write("**Detected Standards:**")
            for standard in report['detected_standards']:
                st.write(f"  ‚Ä¢ {standard}")
    
    with col2:
        st.markdown("### üìä Compliance Summary")
        st.info(report.get('summary', 'No summary available'))
    
    # Strengths and improvements in tabs
    tab1, tab2, tab3 = st.tabs(["‚úÖ Strengths", "‚ö†Ô∏è Improvements Needed", "üí° Recommendations"])
    
    with tab1:
        if report.get('strengths'):
            for strength in report['strengths']:
                st.success(f"‚Ä¢ {strength}")
        else:
            st.info("No specific strengths identified")
    
    with tab2:
        if report.get('improvements_needed'):
            for improvement in report['improvements_needed']:
                st.warning(f"‚Ä¢ {improvement}")
        else:
            st.success("No improvements needed")
    
    with tab3:
        if report.get('recommendations'):
            for rec in report['recommendations']:
                st.info(f"‚Ä¢ {rec}")
        else:
            st.info("No specific recommendations")

@st.cache_resource
def load_rag_system():
    """
    Load embedding model and build FAISS index from documents.
    Includes both pre-existing and uploaded documents.
    """
    # Load embedding model
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    
    # Create output directory if needed
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(DOCUMENTS_PATH, exist_ok=True)
    os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
    
    # Load and process documents from both paths
    doc_chunks = []
    doc_metadata = []
    
    # Process existing documents
    existing_files = glob.glob(f"{DOCUMENTS_PATH}/*")
    
    # Process uploaded documents
    uploaded_files = glob.glob(f"{UPLOADED_DOCS_PATH}/*")
    # Filter out compliance report files
    uploaded_files = [f for f in uploaded_files if not f.endswith('.compliance.json')]
    
    all_files = existing_files + uploaded_files
    
    if not all_files:
        st.warning(f"No documents found. Upload documents to get started.")
        return embedding_model, None, None, []
    
    # Process each document
    for file_path in all_files:
        file_name = Path(file_path).name
        file_ext = Path(file_path).suffix
        
        # Determine if it's an uploaded document
        is_uploaded = UPLOADED_DOCS_PATH in file_path
        
        try:
            if file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                    # Smart chunking for different document types
                    if 'user_story' in file_name.lower():
                        # Split by acceptance criteria
                        if 'Acceptance Criteria:' in content:
                            parts = content.split('Acceptance Criteria:')
                            if parts[0].strip():
                                doc_chunks.append(parts[0].strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'user_story_overview',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                            if len(parts) > 1:
                                doc_chunks.append(f"Acceptance Criteria:\n{parts[1].strip()}")
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'acceptance_criteria',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                        else:
                            doc_chunks.append(content)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'user_story',
                                'doc_type': 'user_story',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # Split by paragraphs for other documents
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip() and len(para.strip()) > 50:
                                doc_chunks.append(para.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'prd' if 'prd' in file_name.lower() else 'document',
                                    'doc_type': 'prd' if 'prd' in file_name.lower() else 'general',
                                    'is_uploaded': is_uploaded
                                })
                                
            elif file_ext in ['.yaml', '.yml']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    yaml_content = yaml.safe_load(f)
                    
                    # Process API specs
                    if 'paths' in yaml_content:
                        for path, methods in yaml_content.get('paths', {}).items():
                            endpoint_doc = f"API Endpoint: {path}\n"
                            endpoint_doc += yaml.dump(methods, default_flow_style=False)
                            doc_chunks.append(endpoint_doc)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'api_endpoint',
                                'endpoint': path,
                                'doc_type': 'api_specification',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # General YAML content
                        doc_chunks.append(yaml.dump(yaml_content, default_flow_style=False))
                        doc_metadata.append({
                            'source': file_name,
                            'type': 'config',
                            'doc_type': 'configuration',
                            'is_uploaded': is_uploaded
                        })
            
            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_pdf(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'pdf_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext in ['.docx', '.doc']:
                with open(file_path, 'rb') as f:
                    text = extract_text_from_docx(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'word_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext == '.xml':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_xml(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'xml_document',
                                    'doc_type': 'specification',
                                    'is_uploaded': is_uploaded
                                })
                        
        except Exception as e:
            st.error(f"Error loading {file_name}: {str(e)}")
    
    if not doc_chunks:
        st.warning("No content extracted from documents.")
        return embedding_model, None, None, []
    
    # Generate embeddings
    with st.spinner(f"Creating embeddings for {len(doc_chunks)} document chunks..."):
        contents = doc_chunks
        embeddings = embedding_model.encode(contents, show_progress_bar=True)
        embeddings = np.array(embeddings).astype('float32')
    
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    
    return embedding_model, index, doc_chunks, doc_metadata

def retrieve_context(query: str, embedding_model, index, doc_chunks, doc_metadata, k: int = 5):
    """Retrieve relevant documents for a query"""
    if index is None:
        return []
    
    query_embedding = embedding_model.encode([query]).astype('float32')
    distances, indices = index.search(query_embedding, k)
    
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(doc_chunks):
            similarity = 1 / (1 + dist)
            results.append({
                'content': doc_chunks[idx],
                'metadata': doc_metadata[idx] if idx < len(doc_metadata) else {},
                'similarity': similarity
            })
    
    return results

def generate_test_case_with_gemini(requirement: str, context_docs: List[Dict]) -> Dict:
    """
    Generate test case using Gemini with NASSCOM compliance
    """
    # Check context relevance
    avg_relevance = sum(doc['similarity'] for doc in context_docs) / len(context_docs) if context_docs else 0
    low_relevance = avg_relevance < 0.3
    
    # Build context string
    if low_relevance:
        context_str = "Note: Limited specific context found. Generating based on general best practices and requirement analysis."
    else:
        context_parts = []
        for doc in context_docs[:3]:
            context_parts.append(f"--- Source: {doc['metadata'].get('source', 'Unknown')} (Type: {doc['metadata'].get('type', 'document')}, Relevance: {doc['similarity']:.2%}) ---")
            context_parts.append(doc['content'][:800])
            context_parts.append("--- End Context ---")
        context_str = "\n\n".join(context_parts)
    
    # Enhanced prompt with NASSCOM requirements
    prompt = f"""You are a meticulous QA Engineer specializing in healthcare/MedTech testing, following NASSCOM guidelines.
Generate a comprehensive test case for the following requirement.

REQUIREMENT:
{requirement}

RELEVANT CONTEXT:
{context_str}

NASSCOM COMPLIANCE REQUIREMENTS:
- Must include all required fields: Test Case ID, Title, Description, Test Steps, Expected Results, Priority
- Should consider healthcare compliance standards: HIPAA, GDPR, FDA, ISO 13485, ISO 27001
- Should handle edge cases and boundary conditions
- Must be traceable to requirements
- Should support both manual and automated testing where feasible

Generate a test case with EXACTLY these fields:
- id: Use "TC_PLACEHOLDER" (will be replaced)
- title: Clear, descriptive title
- description: Detailed description of what is being tested
- category: One of [Functional, Security, Integration, Performance, Usability, Compliance]
- priority: One of [Critical, High, Medium, Low]
- compliance: Array of applicable standards like ["HIPAA", "GDPR", "FDA", "ISO 13485", "ISO 27001"]
- preconditions: String describing what must be true before testing
- test_steps: Array of step descriptions WITHOUT numbering (e.g., ["Navigate to login page", "Enter credentials"])
- expected_results: String describing specific expected outcomes
- test_data: JSON object with test data (e.g., {{"username": "test@example.com", "password": "Test123!"}})
- edge_cases: Array of special scenarios to consider
- negative_tests: Array of negative test scenarios
- automation_feasible: Boolean indicating if this can be automated
- estimated_duration: String like "5 minutes" or "30 minutes"
- traceability: String linking to requirement source

IMPORTANT:
1. test_steps must be an array of strings without prefixes like "Step 1:" or "1."
2. test_data must be a JSON object, not a string
3. Focus on healthcare/medical domain requirements if applicable
4. Include both positive and negative test scenarios
5. Ensure compliance with relevant healthcare standards

Return ONLY a valid JSON object."""
    
    try:
        # Configure Gemini for JSON output
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                top_k=30,
                max_output_tokens=4096
            )
        )
        
        # Safety settings for healthcare content
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
        
        response = model.generate_content(prompt, safety_settings=safety_settings)
        
        if not response.parts:
            return create_fallback_test_case(requirement, error_reason="Content blocked by safety filter")
        
        test_case = json.loads(response.text)
        
        # Generate unique ID
        unique_id = f"TC_{str(uuid.uuid4())[:8].upper()}"
        test_case['id'] = unique_id
        
        # Clean test_steps formatting
        if 'test_steps' in test_case and isinstance(test_case['test_steps'], list):
            cleaned_steps = []
            for step in test_case['test_steps']:
                cleaned = re.sub(r'^(Step\s+\d+:|^\d+\.\s*)', '', str(step)).strip()
                cleaned_steps.append(cleaned)
            test_case['test_steps'] = cleaned_steps
        
        # Ensure test_data is always a dict
        if 'test_data' in test_case:
            if isinstance(test_case['test_data'], str):
                try:
                    test_case['test_data'] = json.loads(test_case['test_data'])
                except:
                    test_case['test_data'] = {"raw_data": test_case['test_data']}
        else:
            test_case['test_data'] = {}
        
        # Add metadata and context tracking
        test_case['generated_from'] = requirement[:100]
        test_case['generation_timestamp'] = datetime.now().isoformat()
        test_case['avg_context_relevance'] = f"{avg_relevance:.2%}"
        test_case['context_sources'] = [
            doc['metadata'].get('source', 'Unknown') for doc in context_docs[:3]
        ] if not low_relevance else ["Generated using best practices"]
        test_case['low_context_confidence'] = low_relevance
        test_case['version'] = 1
        test_case['nasscom_compliant'] = True
        
        # Increment counter
        st.session_state.test_counter += 1
        
        return test_case
        
    except Exception as e:
        st.error(f"Generation error: {e}")
        return create_fallback_test_case(requirement, error_reason=str(e))

def create_fallback_test_case(requirement: str, error_reason: str = "Unknown") -> Dict:
    """
    Create a fallback test case when generation fails.
    """
    st.session_state.test_counter += 1
    return {
        'id': f'TC_FB_{st.session_state.test_counter:04d}',
        'title': f'Test: {requirement[:60]}',
        'description': f'Verify that {requirement}',
        'category': 'Functional',
        'priority': 'Medium',
        'compliance': ['General'],
        'preconditions': 'System is in stable state',
        'test_steps': [
            'Setup test environment',
            'Execute test scenario',
            'Verify results'
        ],
        'expected_results': 'System behaves as specified',
        'test_data': {'placeholder': 'Add specific test data'},
        'edge_cases': [],
        'negative_tests': [],
        'automation_feasible': False,
        'estimated_duration': '10 minutes',
        'traceability': requirement[:100],
        'generated_from': requirement[:100],
        'context_sources': ['Fallback generation'],
        'generation_timestamp': datetime.now().isoformat(),
        'nasscom_compliant': False,
        'fallback': True,
        'error': error_reason
    }

def analyze_test_suite_schema(content: str, file_type: str) -> Dict:
    """
    Use AI to analyze the schema/structure of an uploaded test suite
    """
    prompt = f"""You are an expert test case analyst. Analyze the following test suite data and understand its schema.

FILE TYPE: {file_type}
CONTENT SAMPLE (first 2000 chars):
{content[:2000]}

Analyze the test suite and return a JSON object with:
{{
    "detected_format": string (csv/json/excel/xml/custom),
    "has_headers": boolean,
    "field_mappings": {{
        "id_field": string or null,
        "title_field": string or null,
        "description_field": string or null,
        "steps_field": string or null,
        "expected_field": string or null,
        "priority_field": string or null,
        "category_field": string or null,
        "preconditions_field": string or null,
        "test_data_field": string or null
    }},
    "detected_fields": [list of all field names found],
    "row_count": integer (estimated number of test cases),
    "separator": string (for CSV - comma, tab, pipe, etc),
    "schema_confidence": float (0-100),
    "sample_data": object (one sample test case in original format),
    "recommendations": [list of recommendations for import]
}}

Be intelligent about detecting field mappings even if names don't match exactly.
For example, "Test Name" could map to "title_field", "Steps" to "steps_field", etc."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.2
            )
        )
        
        response = model.generate_content(prompt)
        schema_analysis = json.loads(response.text)
        return schema_analysis
        
    except Exception as e:
        st.error(f"Schema analysis error: {e}")
        return {
            "detected_format": "unknown",
            "has_headers": True,
            "field_mappings": {},
            "detected_fields": [],
            "row_count": 0,
            "schema_confidence": 0,
            "sample_data": {},
            "recommendations": ["Manual mapping may be required"]
        }

def convert_test_with_ai(test_data: Dict, schema: Dict, import_id: str) -> Dict:
    """
    Use AI to intelligently convert a test from any format to our standard format
    """
    prompt = f"""You are converting a test case to NASSCOM-compliant format.

ORIGINAL TEST DATA:
{json.dumps(test_data, indent=2)}

DETECTED SCHEMA MAPPINGS:
{json.dumps(schema.get('field_mappings', {}), indent=2)}

Convert this test case to our standard format with these EXACT fields:
{{
    "id": string (generate unique ID starting with {import_id}),
    "title": string (clear, descriptive title),
    "description": string (detailed description),
    "category": string (Functional/Security/Integration/Performance/Usability/Compliance),
    "priority": string (Critical/High/Medium/Low),
    "compliance": array (applicable standards like ["HIPAA", "GDPR", "FDA"]),
    "preconditions": string (what must be true before testing),
    "test_steps": array of strings (clear step descriptions),
    "expected_results": string (specific expected outcomes),
    "test_data": object (test data as JSON object),
    "edge_cases": array (special scenarios),
    "negative_tests": array (negative test scenarios),
    "automation_feasible": boolean,
    "estimated_duration": string (e.g., "10 minutes"),
    "traceability": string (original requirement reference),
    "imported": true,
    "original_format": "{schema.get('detected_format', 'unknown')}",
    "import_timestamp": "{datetime.now().isoformat()}",
    "nasscom_compliant": true
}}

Be intelligent about:
1. Parsing multi-line text into arrays for test_steps
2. Inferring missing fields from context
3. Detecting healthcare/medical context for compliance standards
4. Splitting combined fields appropriately
5. Converting string representations to proper types

Return ONLY a valid JSON object."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        converted_test = json.loads(response.text)
        
        # Ensure unique ID
        if not converted_test.get('id'):
            converted_test['id'] = f"{import_id}_{str(uuid.uuid4())[:8].upper()}"
        
        return converted_test
        
    except Exception as e:
        # Fallback conversion
        return {
            'id': f"{import_id}_{str(uuid.uuid4())[:8].upper()}",
            'title': test_data.get('title', test_data.get('name', 'Imported Test')),
            'description': test_data.get('description', str(test_data)),
            'category': 'Functional',
            'priority': 'Medium',
            'compliance': [],
            'preconditions': test_data.get('preconditions', 'N/A'),
            'test_steps': [str(test_data.get('steps', 'See original data'))],
            'expected_results': test_data.get('expected', 'As specified'),
            'test_data': test_data,
            'imported': True,
            'import_error': str(e),
            'nasscom_compliant': False
        }

def detect_duplicates(new_tests: List[Dict], existing_tests: List[Dict]) -> List[Dict]:
    """
    Detect potential duplicate test cases using AI
    """
    if not existing_tests or not new_tests:
        return []
    
    duplicates = []
    
    # Create summaries for comparison
    existing_summaries = [
        f"{t.get('title', '')} - {t.get('description', '')[:100]}"
        for t in existing_tests
    ]
    
    for new_test in new_tests:
        new_summary = f"{new_test.get('title', '')} - {new_test.get('description', '')[:100]}"
        
        # Quick check for exact matches
        if new_summary in existing_summaries:
            duplicates.append({
                'new_test': new_test,
                'match_type': 'exact',
                'confidence': 100
            })
            continue
        
        # Use AI for semantic similarity (simplified for performance)
        for i, existing_summary in enumerate(existing_summaries[:10]):  # Check first 10 for demo
            if len(new_summary) > 20 and len(existing_summary) > 20:
                # Simple similarity check
                common_words = set(new_summary.lower().split()) & set(existing_summary.lower().split())
                if len(common_words) > 5:
                    duplicates.append({
                        'new_test': new_test,
                        'existing_test': existing_tests[i],
                        'match_type': 'similar',
                        'confidence': len(common_words) * 10
                    })
    
    return duplicates

def import_test_suite_with_ai(uploaded_file) -> Tuple[List[Dict], Dict]:
    """
    Main function to import test suite using AI analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    file_content = uploaded_file.read()
    
    # Extract content based on file type
    if file_extension == 'csv':
        content = file_content.decode('utf-8', errors='ignore')
        # Parse CSV
        import csv
        from io import StringIO
        reader = csv.DictReader(StringIO(content))
        test_data = list(reader)
    elif file_extension == 'json':
        content = file_content.decode('utf-8', errors='ignore')
        test_data = json.loads(content)
        if not isinstance(test_data, list):
            test_data = [test_data]
    elif file_extension in ['xlsx', 'xls']:
        # Use pandas for Excel
        df = pd.read_excel(io.BytesIO(file_content))
        test_data = df.to_dict('records')
        content = df.to_string()
    else:
        content = file_content.decode('utf-8', errors='ignore')
        # Try to parse as structured data
        test_data = [{'raw_content': content}]
    
    # Analyze schema with AI
    with st.spinner("ü§ñ Analyzing test suite structure with AI..."):
        schema = analyze_test_suite_schema(content, file_extension)
    
    # Generate import ID
    import_id = f"IMP_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Convert tests using AI
    converted_tests = []
    with st.spinner(f"üîÑ Converting {len(test_data)} test cases..."):
        progress_bar = st.progress(0)
        for i, test in enumerate(test_data):
            converted = convert_test_with_ai(test, schema, import_id)
            converted_tests.append(converted)
            progress_bar.progress((i + 1) / len(test_data))
    
    # Check for duplicates
    with st.spinner("üîç Checking for duplicates..."):
        duplicates = detect_duplicates(converted_tests, st.session_state.generated_tests)
    
    # Create import report
    import_report = {
        'import_id': import_id,
        'filename': uploaded_file.name,
        'timestamp': datetime.now().isoformat(),
        'schema_analysis': schema,
        'total_tests': len(test_data),
        'converted_tests': len(converted_tests),
        'duplicates_found': len(duplicates),
        'success_rate': sum(1 for t in converted_tests if t.get('nasscom_compliant', False)) / len(converted_tests) * 100 if converted_tests else 0
    }
    
    return converted_tests, import_report

def display_import_interface():
    """
    Display the AI-powered test import interface
    """
    st.header("üì• AI-Powered Test Suite Import")
    st.markdown("""
    Upload your existing test suites in any format (CSV, Excel, JSON, etc.).
    Our AI will automatically:
    - üîç Analyze the structure and schema
    - üîÑ Convert tests to NASSCOM-compliant format
    - üîé Detect and handle duplicates
    - ‚úÖ Validate and enhance test quality
    """)
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose test suite file",
        type=['csv', 'xlsx', 'xls', 'json', 'xml', 'txt'],
        help="Upload CSV, Excel, JSON or other test suite formats"
    )
    
    if uploaded_file:
        # Import options
        col1, col2 = st.columns(2)
        with col1:
            merge_duplicates = st.checkbox("Merge duplicate tests", value=True)
        with col2:
            enhance_with_ai = st.checkbox("Enhance tests with AI", value=True)
        
        if st.button("üöÄ Start AI Import", type="primary"):
            # Perform import
            converted_tests, import_report = import_test_suite_with_ai(uploaded_file)
            
            # Store in session state
            st.session_state.imported_tests.extend(converted_tests)
            st.session_state.import_history.append(import_report)
            st.session_state['last_converted_tests'] = converted_tests
            st.session_state['last_import_report'] = import_report
            
            # Display results
            st.success(f"‚úÖ Successfully imported {len(converted_tests)} test cases!")
        
        # Show results and add button OUTSIDE the import button condition
        if 'last_converted_tests' in st.session_state and st.session_state['last_converted_tests']:
            # Show import report
            if 'last_import_report' in st.session_state:
                with st.expander("üìä Import Report", expanded=True):
                    import_report = st.session_state['last_import_report']
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Tests", import_report['total_tests'])
                    with col2:
                        st.metric("Converted", import_report['converted_tests'])
                    with col3:
                        st.metric("Duplicates", import_report['duplicates_found'])
                    with col4:
                        st.metric("Success Rate", f"{import_report['success_rate']:.1f}%")
                    
                    # Schema analysis
                    st.markdown("### üìã Detected Schema")
                    schema = import_report['schema_analysis']
                    st.json(schema.get('field_mappings', {}))
                    
                    if schema.get('recommendations'):
                        st.markdown("### üí° Recommendations")
                        for rec in schema['recommendations']:
                            st.info(f"‚Ä¢ {rec}")
            
            # Show converted tests preview
            converted_tests = st.session_state['last_converted_tests']
            with st.expander("üëÅÔ∏è Preview Imported Tests", expanded=False):
                for test in converted_tests[:5]:  # Show first 5
                    st.markdown(f"**{test['id']}**: {test['title']}")
                    st.write(f"Priority: {test['priority']} | Category: {test['category']}")
                if len(converted_tests) > 5:
                    st.info(f"... and {len(converted_tests) - 5} more tests")
            
            # Add to main test suite button - NOW OUTSIDE THE IMPORT CONDITION
            col1, col2 = st.columns(2)
            with col1:
                if st.button("‚úÖ Add to Test Suite", key="add_imported_to_suite", type="primary"):
                    tests_to_add = st.session_state['last_converted_tests']
                    st.session_state.generated_tests.extend(tests_to_add)
                    st.success(f"‚úÖ Added {len(tests_to_add)} tests to main suite!")
                    # Clear the temporary storage
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.balloons()
                    st.info("Go to 'üóÇÔ∏è Test Suite' tab to view your tests")
            with col2:
                if st.button("‚ùå Cancel Import", key="cancel_import"):
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.info("Import cancelled")
    
    # Show import history
    if st.session_state.import_history:
        st.markdown("### üìú Import History")
        for report in reversed(st.session_state.import_history[-5:]):  # Show last 5
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                st.write(f"üìÑ {report['filename']}")
            with col2:
                st.write(f"Tests: {report['converted_tests']}")
            with col3:
                st.write(f"‚úÖ {report['success_rate']:.0f}%")

def display_upload_interface():
    """
    Display document upload interface with NASSCOM compliance validation
    """
    st.header("üì§ Upload Documents for Test Generation")
    st.markdown("""
    Upload your documents to expand the knowledge base for test case generation.
    Documents will be analyzed for NASSCOM compliance before being added to the system.
    """)
    
    # Show accepted formats
    with st.expander("üìã NASSCOM Document Requirements", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### ‚úÖ Accepted Document Types")
            for doc_type, description in NASSCOM_REQUIREMENTS["document_types"].items():
                st.write(f"‚Ä¢ **{description}**")
            
            st.markdown("### üìÅ Supported Formats")
            formats = ", ".join([f"`.{fmt}`" for fmt in NASSCOM_REQUIREMENTS["document_formats"]])
            st.write(formats)
        
        with col2:
            st.markdown("### üè• Healthcare Standards")
            for standard in NASSCOM_REQUIREMENTS["compliance_standards"]:
                st.write(f"‚Ä¢ {standard}")
            
            st.markdown("### üìä Quality Criteria")
            st.write("‚Ä¢ Minimum 100 characters")
            st.write("‚Ä¢ Clear structure and sections")
            st.write("‚Ä¢ Healthcare context preferred")
            st.write("‚Ä¢ Traceable requirements")
    
    # File upload widget
    uploaded_files = st.file_uploader(
        "Choose documents to upload",
        type=NASSCOM_REQUIREMENTS["document_formats"],
        accept_multiple_files=True,
        help="Upload PRDs, User Stories, API Specs, Test Plans, etc."
    )
    
    if uploaded_files:
        st.markdown("### üìù Document Analysis Results")
        
        for uploaded_file in uploaded_files:
            with st.expander(f"üìÑ {uploaded_file.name}", expanded=True):
                with st.spinner(f"Analyzing {uploaded_file.name} for NASSCOM compliance..."):
                    # Process file
                    content, compliance_report = process_uploaded_file(uploaded_file)
                    
                    # Store in session state
                    st.session_state.compliance_reports[uploaded_file.name] = compliance_report
                    
                    # Display compliance report
                    display_compliance_report(compliance_report)
                    
                    # Show content preview
                    with st.expander("üìñ Content Preview", expanded=False):
                        st.text_area("", value=content[:1000] + "..." if len(content) > 1000 else content, 
                                   height=200, disabled=True, label_visibility="collapsed")
                    
                    # Action buttons
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if compliance_report['is_compliant']:
                            if st.button(f"‚úÖ Add to Knowledge Base", key=f"add_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.success(f"Added {uploaded_file.name} to knowledge base!")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    # Clear cache to reload with new documents
                                    st.cache_resource.clear()
                                    st.rerun()
                        else:
                            st.warning("Document needs improvements before adding to knowledge base")
                    
                    with col2:
                        # Download compliance report
                        report_json = json.dumps(compliance_report, indent=2)
                        st.download_button(
                            label="üì• Download Compliance Report",
                            data=report_json,
                            file_name=f"{uploaded_file.name}_compliance.json",
                            mime="application/json",
                            key=f"download_{uploaded_file.name}"
                        )
                    
                    with col3:
                        if not compliance_report['is_compliant']:
                            if st.button(f"üîÑ Override & Add Anyway", key=f"override_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.warning(f"Added {uploaded_file.name} despite compliance issues")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    st.cache_resource.clear()
                                    st.rerun()

def refine_test_case_with_feedback(test_case: Dict, feedback: str, specific_field: str = None) -> Dict:
    """
    Refine an existing test case based on human feedback.
    """
    test_case_for_prompt = {k: v for k, v in test_case.items() if k != 'retrieved_context'}
    
    if specific_field:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK FOR '{specific_field}':
{feedback}

NASSCOM REQUIREMENTS:
- Maintain all required fields
- Ensure healthcare compliance where applicable
- Include both positive and negative scenarios
- Maintain traceability

Generate an improved version of this test case, specifically improving the '{specific_field}' field based on the feedback.
Maintain all other fields unless they need adjustment to be consistent with the change.

Return ONLY a valid JSON object with the same structure."""
    else:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK:
{feedback}

Generate an improved version of this test case incorporating the feedback while maintaining NASSCOM compliance.

Return ONLY a valid JSON object."""
    
    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                max_output_tokens=4096
            )
        )
        
        response = model.generate_content(prompt)
        refined_test = json.loads(response.text)
        
        # Preserve metadata
        refined_test['id'] = test_case['id']
        refined_test['refinement_timestamp'] = datetime.now().isoformat()
        refined_test['refinement_feedback'] = feedback
        refined_test['version'] = test_case.get('version', 1) + 1
        refined_test['nasscom_compliant'] = True
        
        if 'retrieved_context' in test_case:
            refined_test['retrieved_context'] = test_case['retrieved_context']
        
        return refined_test
        
    except Exception as e:
        st.error(f"Error refining test case: {e}")
        return test_case

def display_test_case(test_case: Dict, context_docs: List[Dict]):
    """Display generated test case in a formatted way"""
    
    # Header with metrics
    st.markdown("---")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Priority", test_case.get('priority', 'N/A'))
    with col2:
        st.metric("Category", test_case.get('category', 'N/A'))
    with col3:
        compliance = test_case.get('compliance', [])
        st.metric("Compliance", len(compliance) if compliance else 'None')
    with col4:
        st.metric("Duration", test_case.get('estimated_duration', 'N/A'))
    
    # NASSCOM compliance indicator
    if test_case.get('nasscom_compliant', False):
        st.success("‚úÖ NASSCOM Compliant Test Case")
    else:
        st.warning("‚ö†Ô∏è Non-compliant test case (fallback generated)")
    
    # Main content
    st.markdown("### üìù Description")
    st.write(test_case.get('description', 'N/A'))
    
    # Traceability
    if test_case.get('traceability'):
        st.markdown("### üîó Traceability")
        st.info(test_case.get('traceability', 'N/A'))
    
    # Preconditions
    st.markdown("### ‚öôÔ∏è Preconditions")
    st.info(test_case.get('preconditions', 'None specified'))
    
    # Test steps
    st.markdown("### üìã Test Steps")
    steps = test_case.get('test_steps', [])
    for i, step in enumerate(steps, 1):
        st.write(f"{i}. {step}")
    
    # Expected results
    st.markdown("### ‚úÖ Expected Results")
    st.success(test_case.get('expected_results', 'N/A'))
    
    # Test data
    if test_case.get('test_data'):
        st.markdown("### üìä Test Data")
        st.json(test_case.get('test_data'))
    
    # Edge cases and negative tests
    col1, col2 = st.columns(2)
    with col1:
        if test_case.get('edge_cases'):
            st.markdown("### ‚ö†Ô∏è Edge Cases")
            for edge in test_case['edge_cases']:
                st.write(f"‚Ä¢ {edge}")
    
    with col2:
        if test_case.get('negative_tests'):
            st.markdown("### üö´ Negative Tests")
            for neg_test in test_case['negative_tests']:
                st.write(f"‚Ä¢ {neg_test}")
    
    # Metadata
    with st.expander("üîç Generation Details"):
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Automation Feasible:** {'Yes' if test_case.get('automation_feasible') else 'No'}")
            st.write(f"**Context Relevance:** {test_case.get('avg_context_relevance', 'N/A')}")
            if test_case.get('compliance'):
                st.write(f"**Standards:** {', '.join(test_case['compliance'])}")
        with col2:
            st.write(f"**Generated From:** {test_case.get('generated_from', 'N/A')[:50]}...")
            st.write(f"**Timestamp:** {test_case.get('generation_timestamp', 'N/A')[:19]}")
            st.write(f"**Version:** {test_case.get('version', 1)}")
    
    # Save options
    st.markdown("#### üíæ Save Options")
    col1, col2 = st.columns(2)
    with col1:
        if st.button(f"Save as JSON", key=f"save_json_{test_case['id']}"):
            filename = save_test_case(test_case, 'json')
            st.success(f"Saved to {filename}")
    with col2:
        if st.button(f"Save as CSV", key=f"save_csv_{test_case['id']}"):
            filename = save_test_case(test_case, 'csv')
            st.success(f"Saved to {filename}")

def save_test_case(test_case: Dict, format: str = 'json'):
    """Save test case to file"""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if format == 'json':
        filename = f"{OUTPUT_DIR}/test_case_{test_case['id']}_{timestamp}.json"
        with open(filename, 'w') as f:
            json.dump(test_case, f, indent=2, default=str)
        return filename
    elif format == 'csv':
        filename = f"{OUTPUT_DIR}/test_case_{test_case['id']}_{timestamp}.csv"
        df = pd.DataFrame([test_case])
        df.to_csv(filename, index=False)
        return filename
    return None

# MAIN UI
def main():
    # Header
    st.title("üß™ AI-Powered Test Case Generator")
    st.markdown("### NASSCOM Healthcare/MedTech Testing with RAG Pipeline")
    
    # Sidebar for system info
    with st.sidebar:
        st.header("üìä System Status")
        
        # Load RAG system
        embedding_model, index, doc_chunks, doc_metadata = load_rag_system()
        
        if index:
            st.success(f"‚úÖ RAG System Ready")
            st.info(f"üìö {len(doc_chunks)} document chunks indexed")
            
            # Count documents
            existing_docs = len(set(m.get('source') for m in doc_metadata if not m.get('is_uploaded', False)))
            uploaded_docs = len(set(m.get('source') for m in doc_metadata if m.get('is_uploaded', False)))
            
            st.info(f"üìÅ {existing_docs} existing documents")
            st.info(f"üì§ {uploaded_docs} uploaded documents")
            
            # Show compliance summary
            if st.session_state.compliance_reports:
                st.markdown("### üìã Compliance Summary")
                compliant_count = sum(1 for r in st.session_state.compliance_reports.values() if r['is_compliant'])
                total_count = len(st.session_state.compliance_reports)
                st.metric("Compliant Documents", f"{compliant_count}/{total_count}")
        else:
            st.warning("‚ö†Ô∏è No documents loaded")
            st.info("Upload documents to get started")
        
        st.markdown("---")
        st.metric("Test Cases Generated", st.session_state.test_counter)
        
        # Export options
        if st.session_state.generated_tests:
            st.markdown("### üíæ Export All Tests")
            if st.button("Export to JSON"):
                all_tests = {
                    "generated_on": datetime.now().isoformat(),
                    "total_tests": len(st.session_state.generated_tests),
                    "nasscom_compliant": True,
                    "test_cases": st.session_state.generated_tests
                }
                filename = f"{OUTPUT_DIR}/all_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                with open(filename, 'w') as f:
                    json.dump(all_tests, f, indent=2, default=str)
                st.success(f"Saved to {filename}")

    # Main content area with tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "üß™ Generate Tests", 
        "üì§ Upload Documents",
        "üì• Import Tests",  # NEW TAB
        "üóÇÔ∏è Test Suite", 
        "üìö Document Library"
    ])

    with tab1:
        st.header("üéØ Generate New Test Case")
        st.markdown("Describe the feature or requirement you want to test.")

        # Predefined examples
        examples = [
            "Test user registration with HIPAA-compliant data handling",
            "Verify patient data encryption meets FDA and ISO 27001 standards",
            "Test API authentication with OAuth 2.0 for medical records access",
            "Validate HL7 message processing for lab results",
            "Test DICOM image upload with PHI anonymization",
            "Verify audit trail generation for GDPR compliance"
        ]
        
        selected_example = st.selectbox("Choose an example or write your own:", 
                                       ["Custom requirement..."] + examples)
        
        if selected_example == "Custom requirement...":
            requirement = st.text_area(
                "Enter your test requirement:",
                height=100,
                placeholder="Describe what you want to test..."
            )
        else:
            requirement = selected_example
            st.info(f"Using example: {requirement}")
        
        # Generation button
        if st.button("üöÄ Generate Test Case", type="primary"):
            if not requirement or requirement == "Custom requirement...":
                st.warning("Please enter a requirement")
            elif not index:
                st.error("No documents in knowledge base. Please upload documents first.")
            else:
                with st.spinner("üîç Retrieving context and generating NASSCOM-compliant test case..."):
                    # Retrieve context
                    context_docs = retrieve_context(
                        requirement, 
                        embedding_model, 
                        index, 
                        doc_chunks, 
                        doc_metadata,
                        k=5
                    )
                    
                    # Generate test case
                    test_case = generate_test_case_with_gemini(requirement, context_docs)
                    
                    # Add context to test case object
                    test_case['retrieved_context'] = context_docs
                    
                    # Store in session
                    st.session_state.generated_tests.insert(0, test_case)
                    
                    # Display results
                    if test_case.get('nasscom_compliant', False):
                        st.success(f"‚úÖ NASSCOM-Compliant Test Case Generated: {test_case['id']}")
                    else:
                        st.warning(f"‚ö†Ô∏è Test case generated with limited compliance: {test_case['id']}")
                    
                    st.balloons()
                    
                    # Show the generated test case
                    display_test_case(test_case, context_docs)

    with tab2:
        display_upload_interface()

    with tab3:
        display_import_interface()  # NEW IMPORT INTERFACE

    with tab4:
        st.header("üóÇÔ∏è Generated Test Suite")
        
        if not st.session_state.generated_tests:
            st.info("No test cases generated yet. Go to the 'Generate Tests' tab to start.")
        else:
            st.markdown(f"You have generated **{len(st.session_state.generated_tests)}** test case(s).")
            
            # Filter options
            col1, col2, col3 = st.columns(3)
            with col1:
                filter_category = st.selectbox("Filter by Category", 
                                              ["All"] + list(set(tc.get('category', 'Unknown') for tc in st.session_state.generated_tests)))
            with col2:
                filter_priority = st.selectbox("Filter by Priority",
                                              ["All"] + list(set(tc.get('priority', 'Unknown') for tc in st.session_state.generated_tests)))
            with col3:
                filter_compliance = st.selectbox("Filter by Compliance",
                                                ["All", "NASSCOM Compliant", "Non-Compliant"])
            
            # Apply filters
            filtered_tests = st.session_state.generated_tests
            if filter_category != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('category') == filter_category]
            if filter_priority != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('priority') == filter_priority]
            if filter_compliance == "NASSCOM Compliant":
                filtered_tests = [tc for tc in filtered_tests if tc.get('nasscom_compliant', False)]
            elif filter_compliance == "Non-Compliant":
                filtered_tests = [tc for tc in filtered_tests if not tc.get('nasscom_compliant', False)]
            
            st.markdown(f"Showing **{len(filtered_tests)}** test case(s)")
            
            for i, tc in enumerate(filtered_tests):
                with st.expander(f"**{tc['id']}**: {tc.get('title', 'Untitled')} (v{tc.get('version', 1)})", expanded=(i==0)):
                    display_test_case(tc, tc.get('retrieved_context', []))

    with tab5:
        st.header("üìö Document Library")
        
        # Show both existing and uploaded documents
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("üìÅ Pre-loaded Documents")
            existing_files = glob.glob(f"{DOCUMENTS_PATH}/*")
            if existing_files:
                for file_path in existing_files:
                    file_name = Path(file_path).name
                    file_size = os.path.getsize(file_path)
                    st.write(f"‚Ä¢ {file_name} ({file_size:,} bytes)")
            else:
                st.info("No pre-loaded documents")
        
        with col2:
            st.subheader("üì§ Uploaded Documents")
            uploaded_files = glob.glob(f"{UPLOADED_DOCS_PATH}/*")
            uploaded_files = [f for f in uploaded_files if not f.endswith('.compliance.json')]
            
            if uploaded_files:
                for file_path in uploaded_files:
                    file_name = Path(file_path).name
                    file_size = os.path.getsize(file_path)
                    
                    # Check if we have compliance report
                    compliance_file = f"{file_path}.compliance.json"
                    if os.path.exists(compliance_file):
                        with open(compliance_file, 'r') as f:
                            report = json.load(f)
                        compliance_icon = "‚úÖ" if report['is_compliant'] else "‚ö†Ô∏è"
                        compliance_score = report['compliance_score']
                        st.write(f"{compliance_icon} {file_name} ({compliance_score:.0f}% compliant)")
                    else:
                        st.write(f"‚Ä¢ {file_name} ({file_size:,} bytes)")
            else:
                st.info("No uploaded documents yet")
        
        # Statistics
        st.markdown("---")
        st.subheader("üìä Knowledge Base Statistics")
        
        total_files = len(existing_files) + len(uploaded_files)
        total_size = sum(os.path.getsize(f) for f in existing_files + uploaded_files if not f.endswith('.compliance.json'))
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Documents", total_files)
        with col2:
            st.metric("Total Size", f"{total_size:,} bytes")
        with col3:
            if st.session_state.compliance_reports:
                avg_compliance = sum(r['compliance_score'] for r in st.session_state.compliance_reports.values()) / len(st.session_state.compliance_reports)
                st.metric("Avg Compliance Score", f"{avg_compliance:.1f}%")


if __name__ == "__main__":
    main()
