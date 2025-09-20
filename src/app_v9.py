"""
🧪 AI-Powered Test Case Generator - Version 9.0
Modern UI with Simplified Folder Structure
NASSCOM Healthcare/MedTech Testing Solution
"""

import streamlit as st
import os
import json
import uuid
import numpy as np
import faiss
import yaml
import glob
import re
import io
import tempfile
import shutil
import time
import random
from pathlib import Path
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from dotenv import load_dotenv
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import pandas as pd
import PyPDF2
from docx import Document
import xml.etree.ElementTree as ET
import chardet

# Load environment variables
load_dotenv()

# ==================================
# 🎨 PAGE CONFIGURATION (MUST BE FIRST)
# ==================================

st.set_page_config(
    page_title="AI Test Generator Pro",
    page_icon="🧪",
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items={
        'Get Help': 'https://nasscom.in',
        'About': "AI Test Generator v9.0 - Modern UI with Clean Architecture"
    }
)

# ==================================
# 🎨 MODERN UI STYLING AND ANIMATIONS
# ==================================

def inject_modern_css():
    """Inject clean, professional CSS with smooth transitions"""
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
    /* Modern Elegant Color Palette - Aurora Theme */
    :root {
        --primary: #8B5CF6;        /* Vibrant Purple */
        --primary-light: #A78BFA;  /* Light Purple */
        --primary-dark: #7C3AED;   /* Deep Purple */
        --secondary: #14B8A6;      /* Teal */
        --accent: #F97316;         /* Sunset Orange */
        --success: #10B981;        /* Emerald Green */
        --danger: #F43F5E;         /* Rose Red */
        --warning: #EAB308;        /* Golden Yellow */
        --info: #3B82F6;           /* Sky Blue */
        
        /* Sophisticated Dark Backgrounds */
        --bg-dark: #0A0E1A;        /* Rich Dark Navy */
        --bg-medium: #141824;      /* Deep Midnight */
        --bg-light: #1F2937;       /* Charcoal Blue */
        --bg-surface: #262B3B;     /* Surface Blue */
        
        /* Text Colors with Better Contrast */
        --text-primary: #F9FAFB;   /* Pure White */
        --text-secondary: #D1D5DB; /* Light Gray */
        --text-muted: #9CA3AF;     /* Muted Gray */
        
        /* Border & Overlay Colors */
        --border: #374151;         /* Subtle Gray Border */
        --border-light: #4B5563;   /* Light Border */
        --overlay: rgba(139, 92, 246, 0.1); /* Purple Overlay */
        
        /* Gradient Definitions */
        --gradient-primary: linear-gradient(135deg, #8B5CF6 0%, #EC4899 100%);
        --gradient-secondary: linear-gradient(135deg, #14B8A6 0%, #3B82F6 100%);
        --gradient-dark: linear-gradient(135deg, #0A0E1A 0%, #1F2937 100%);
        --gradient-surface: linear-gradient(135deg, #141824 0%, #262B3B 100%);
        --gradient-glow: linear-gradient(135deg, #8B5CF6 0%, #14B8A6 50%, #F97316 100%);
    }
        
        /* Premium gradient background with subtle animation */
        .stApp {
            background: var(--gradient-dark);
            color: var(--text-primary);
            min-height: 100vh;
            position: relative;
        }
        
        /* Animated gradient overlay */
        .stApp::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: radial-gradient(circle at 20% 80%, rgba(139, 92, 246, 0.15) 0%, transparent 50%),
                        radial-gradient(circle at 80% 20%, rgba(20, 184, 166, 0.15) 0%, transparent 50%),
                        radial-gradient(circle at 40% 40%, rgba(249, 115, 22, 0.1) 0%, transparent 50%);
            animation: gradientShift 20s ease infinite;
            pointer-events: none;
            z-index: 0;
        }
        
        @keyframes gradientShift {
            0%, 100% { transform: rotate(0deg) scale(1); }
            25% { transform: rotate(90deg) scale(1.1); }
            50% { transform: rotate(180deg) scale(1); }
            75% { transform: rotate(270deg) scale(1.1); }
        }
        
        /* Main container with glassmorphism */
        .main .block-container {
            background: linear-gradient(135deg, 
                rgba(20, 24, 36, 0.8) 0%, 
                rgba(31, 41, 55, 0.6) 50%,
                rgba(38, 43, 59, 0.8) 100%);
            border: 1px solid rgba(139, 92, 246, 0.2);
            border-radius: 16px;
            padding: 2rem;
            backdrop-filter: blur(20px) saturate(150%);
            box-shadow: 0 8px 32px rgba(139, 92, 246, 0.1),
                        inset 0 1px 0 rgba(255, 255, 255, 0.05);
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            z-index: 1;
        }
        
        .main .block-container:hover {
            box-shadow: 0 12px 48px rgba(139, 92, 246, 0.15),
                        inset 0 1px 0 rgba(255, 255, 255, 0.1);
        }
        
        /* Typography - Clean and readable */
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Inter', -apple-system, sans-serif;
            color: var(--text-primary);
            font-weight: 600;
            letter-spacing: -0.02em;
            transition: color 0.3s ease;
        }
        
        /* Premium Aurora-themed buttons */
        .stButton > button {
            background: var(--gradient-primary);
            color: white;
            border: none;
            border-radius: 12px;
            padding: 0.75rem 1.5rem;
            font-weight: 600;
            font-size: 0.9rem;
            letter-spacing: 0.025em;
            text-transform: uppercase;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
            box-shadow: 0 4px 15px rgba(139, 92, 246, 0.3),
                        inset 0 1px 0 rgba(255, 255, 255, 0.2);
            z-index: 1;
        }
        
        /* Gradient overlay effect */
        .stButton > button::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(135deg, var(--primary-light) 0%, var(--primary) 100%);
            opacity: 0;
            transition: opacity 0.3s ease;
            z-index: -1;
        }
        
        /* Ripple effect */
        .stButton > button::after {
            content: '';
            position: absolute;
            top: 50%;
            left: 50%;
            width: 0;
            height: 0;
            border-radius: 50%;
            background: rgba(255, 255, 255, 0.5);
            transform: translate(-50%, -50%);
            transition: width 0.6s, height 0.6s;
            z-index: -1;
        }
        
        /* Hover state with glow effect */
        .stButton > button:hover {
            transform: translateY(-2px) scale(1.02);
            box-shadow: 0 8px 25px rgba(139, 92, 246, 0.4),
                        0 0 40px rgba(139, 92, 246, 0.2),
                        inset 0 1px 0 rgba(255, 255, 255, 0.3);
            background: linear-gradient(135deg, #A78BFA 0%, #EC4899 100%);
        }
        
        .stButton > button:hover::before {
            opacity: 1;
        }
        
        .stButton > button:hover::after {
            width: 300px;
            height: 300px;
            opacity: 0;
        }
        
        /* Active/Click state */
        .stButton > button:active {
            transform: translateY(0) scale(0.98);
            box-shadow: 0 2px 10px rgba(139, 92, 246, 0.3);
        }
        
        /* Success button variant */
        .stButton > button[kind="primary"] {
            background: var(--gradient-secondary);
            box-shadow: 0 4px 15px rgba(20, 184, 166, 0.3);
        }
        
        .stButton > button[kind="primary"]:hover {
            box-shadow: 0 8px 25px rgba(20, 184, 166, 0.4);
            background: linear-gradient(135deg, #5EEAD4 0%, #60A5FA 100%);
        }
        
        /* Secondary button variant */
        .stButton > button[kind="secondary"] {
            background: transparent;
            border: 2px solid var(--primary);
            color: var(--primary-light);
        }
        
        .stButton > button[kind="secondary"]:hover {
            background: var(--primary);
            color: white;
            border-color: var(--primary-light);
        }
        
        /* Button loading state */
        .stButton > button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none !important;
        }
        
        /* Button icon animation */
        .stButton > button i,
        .stButton > button svg {
            transition: transform 0.3s ease;
            margin-right: 0.5rem;
        }
        
        .stButton > button:hover i,
        .stButton > button:hover svg {
            transform: rotate(5deg) scale(1.1);
        }
        
        /* Pulse animation for important buttons */
        @keyframes buttonPulse {
            0% {
                box-shadow: 0 0 0 0 rgba(79, 70, 229, 0.7);
            }
            70% {
                box-shadow: 0 0 0 10px rgba(79, 70, 229, 0);
            }
            100% {
                box-shadow: 0 0 0 0 rgba(79, 70, 229, 0);
            }
        }
        
        .stButton > button.pulse {
            animation: buttonPulse 2s infinite;
        }
        
        /* Glow effect for hover */
        @keyframes buttonGlow {
            0%, 100% {
                box-shadow: 0 2px 8px rgba(79, 70, 229, 0.25);
            }
            50% {
                box-shadow: 0 4px 16px rgba(79, 70, 229, 0.5);
            }
        }
        
        .stButton > button:focus {
            outline: none;
            box-shadow: 0 0 0 3px rgba(79, 70, 229, 0.2);
        }
        
        /* Download button special style */
        .stDownloadButton > button {
            background: linear-gradient(135deg, #06B6D4 0%, #0891B2 100%);
            border-radius: 10px;
            font-weight: 600;
            box-shadow: 0 2px 8px rgba(6, 182, 212, 0.25);
        }
        
        .stDownloadButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(6, 182, 212, 0.4);
        }
        
        /* Form submit button special style */
        button[type="submit"] {
            width: 100%;
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%) !important;
            font-size: 1rem !important;
            padding: 0.875rem !important;
            margin-top: 0.5rem;
        }
        
        /* Aurora-themed input fields */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea,
        .stSelectbox > div > div > select,
        .stNumberInput > div > div > input {
            background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-medium) 100%);
            border: 2px solid rgba(139, 92, 246, 0.2);
            border-radius: 12px;
            color: var(--text-primary);
            padding: 0.75rem 1rem;
            font-size: 0.95rem;
            font-weight: 500;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2),
                        0 1px 2px rgba(139, 92, 246, 0.1);
            position: relative;
            background-clip: padding-box;
        }
        
        .stTextInput > div > div > input::before,
        .stTextArea > div > div > textarea::before {
            content: '';
            position: absolute;
            top: 0;
            right: 0;
            bottom: 0;
            left: 0;
            z-index: -1;
            margin: -2px;
            border-radius: 10px;
            background: linear-gradient(135deg, var(--border) 0%, var(--bg-light) 100%);
        }
        
        .stTextInput > div > div > input:hover,
        .stTextArea > div > div > textarea:hover,
        .stSelectbox > div > div > select:hover {
            border-color: rgba(139, 92, 246, 0.4);
            background: linear-gradient(135deg, var(--bg-light) 0%, var(--bg-surface) 100%);
            transform: translateY(-1px);
            box-shadow: 0 4px 15px rgba(139, 92, 246, 0.15), 
                        inset 0 2px 4px rgba(139, 92, 246, 0.05);
        }
        
        .stTextInput > div > div > input:focus,
        .stTextArea > div > div > textarea:focus,
        .stSelectbox > div > div > select:focus {
            outline: none;
            border-color: var(--primary);
            box-shadow: 0 0 0 4px rgba(139, 92, 246, 0.2), 
                        0 6px 20px rgba(139, 92, 246, 0.2),
                        inset 0 2px 4px rgba(0, 0, 0, 0.1);
            background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-light) 100%);
            transform: translateY(-2px);
        }
        
        /* Animated placeholder */
        .stTextInput > div > div > input::placeholder,
        .stTextArea > div > div > textarea::placeholder {
            color: var(--text-secondary);
            transition: all 0.3s ease;
            opacity: 0.7;
        }
        
        .stTextInput > div > div > input:focus::placeholder,
        .stTextArea > div > div > textarea:focus::placeholder {
            opacity: 0.4;
            transform: translateX(5px);
        }
        
        /* Aurora-themed metrics cards */
        [data-testid="metric-container"] {
            background: var(--gradient-surface);
            border: 1px solid rgba(139, 92, 246, 0.2);
            border-radius: 12px;
            padding: 1.25rem;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2),
                        inset 0 1px 0 rgba(255, 255, 255, 0.02);
        }
        
        [data-testid="metric-container"]::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 3px;
            background: var(--gradient-primary);
            transition: left 0.3s ease;
            box-shadow: 0 0 20px rgba(139, 92, 246, 0.5);
        }
        
        [data-testid="metric-container"]:hover {
            transform: translateY(-3px) scale(1.02);
            box-shadow: 0 8px 25px rgba(139, 92, 246, 0.2);
            border-color: var(--primary-light);
            background: linear-gradient(135deg, 
                rgba(38, 43, 59, 0.9) 0%, 
                rgba(31, 41, 55, 0.9) 100%);
        }
        
        [data-testid="metric-container"]:hover::before {
            left: 0;
        }
        
        [data-testid="metric-container"] [data-testid="metric-value"] {
            background: var(--gradient-primary);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            font-weight: 700;
            font-size: 1.5rem;
            transition: all 0.3s ease;
        }
        
        /* Aurora tabs with glowing effect */
        .stTabs [data-baseweb="tab-list"] {
            background: var(--gradient-surface);
            border: 1px solid rgba(139, 92, 246, 0.2);
            border-radius: 14px;
            padding: 6px;
            box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2),
                        0 2px 8px rgba(139, 92, 246, 0.1);
            position: relative;
        }
        
        .stTabs [data-baseweb="tab"] {
            color: var(--text-secondary);
            background: transparent;
            border-radius: 10px;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            font-weight: 500;
            padding: 0.6rem 1.2rem !important;
            overflow: hidden;
        }
        
        .stTabs [data-baseweb="tab"]::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(135deg, transparent 0%, var(--overlay) 100%);
            opacity: 0;
            transition: opacity 0.3s ease;
            border-radius: 10px;
        }
        
        .stTabs [data-baseweb="tab"]:hover {
            color: var(--text-primary);
            transform: translateY(-1px);
            background: rgba(139, 92, 246, 0.1);
        }
        
        .stTabs [data-baseweb="tab"]:hover::before {
            opacity: 1;
        }
        
        .stTabs [aria-selected="true"] {
            background: var(--gradient-primary) !important;
            color: white !important;
            font-weight: 600;
            box-shadow: 0 4px 15px rgba(139, 92, 246, 0.4),
                        inset 0 1px 0 rgba(255, 255, 255, 0.2);
            transform: translateY(-1px) scale(1.02);
        }
        
        /* Tab underline animation */
        .stTabs [data-baseweb="tab-panel"] {
            animation: tabFadeIn 0.5s ease;
        }
        
        @keyframes tabFadeIn {
            from {
                opacity: 0;
                transform: translateX(10px);
            }
            to {
                opacity: 1;
                transform: translateX(0);
            }
        }
        
        /* Progress bar with animation */
        .stProgress > div > div > div > div {
            background: var(--primary);
            transition: width 0.3s ease;
            position: relative;
            overflow: hidden;
        }
        
        .stProgress > div > div > div > div::after {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 2s infinite;
        }
        
        @keyframes shimmer {
            0% { transform: translateX(-100%); }
            100% { transform: translateX(100%); }
        }
        
        /* Premium expanders with smooth animations */
        .streamlit-expanderHeader {
            background: linear-gradient(135deg, var(--bg-medium) 0%, rgba(30, 41, 59, 0.6) 100%);
            border: 1px solid var(--border);
            border-radius: 10px;
            color: var(--text-primary) !important;
            padding: 0.75rem 1rem !important;
            font-weight: 500;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }
        
        .streamlit-expanderHeader::before {
            content: '';
            position: absolute;
            left: 0;
            top: 0;
            height: 100%;
            width: 3px;
            background: var(--primary);
            transform: scaleY(0);
            transition: transform 0.3s ease;
        }
        
        .streamlit-expanderHeader:hover {
            background: linear-gradient(135deg, var(--bg-light) 0%, var(--bg-medium) 100%);
            border-color: var(--primary);
            transform: translateX(4px);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.1);
        }
        
        .streamlit-expanderHeader:hover::before {
            transform: scaleY(1);
        }
        
        /* Expander content animation */
        .streamlit-expanderContent {
            animation: expanderSlide 0.3s ease;
            background: rgba(30, 41, 59, 0.3);
            border-radius: 0 0 10px 10px;
            padding: 1rem !important;
        }
        
        @keyframes expanderSlide {
            from {
                opacity: 0;
                transform: translateY(-10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        /* Checkboxes and Radio buttons */
        .stCheckbox > label,
        .stRadio > label {
            transition: all 0.2s ease;
            padding: 0.5rem;
            border-radius: 6px;
        }
        
        .stCheckbox > label:hover,
        .stRadio > label:hover {
            background: rgba(79, 70, 229, 0.05);
            transform: translateX(2px);
        }
        
        /* Custom checkbox styling */
        .stCheckbox > label > span:first-child,
        .stRadio > div > label > span:first-child {
            border: 2px solid var(--border) !important;
            background: var(--bg-medium) !important;
            transition: all 0.3s ease !important;
        }
        
        .stCheckbox > label > span:first-child:checked,
        .stRadio > div > label > span:first-child:checked {
            background: var(--primary) !important;
            border-color: var(--primary) !important;
            animation: checkBounce 0.3s ease;
        }
        
        @keyframes checkBounce {
            0% { transform: scale(1); }
            50% { transform: scale(1.2); }
            100% { transform: scale(1); }
        }
        
        /* Slider styling */
        .stSlider > div > div > div {
            background: var(--bg-medium) !important;
        }
        
        .stSlider > div > div > div > div {
            background: var(--primary) !important;
            transition: all 0.2s ease;
        }
        
        .stSlider > div > div > div > div > div {
            background: var(--primary-light) !important;
            border: 3px solid white !important;
            box-shadow: 0 2px 8px rgba(79, 70, 229, 0.3) !important;
            transition: all 0.2s ease;
        }
        
        .stSlider > div > div > div > div > div:hover {
            transform: scale(1.2);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.5) !important;
        }
        
        /* Premium alerts with enhanced animations */
        .stAlert {
            animation: alertSlideIn 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            border-radius: 10px;
            border-left: 4px solid;
            background: linear-gradient(135deg, rgba(30, 41, 59, 0.9) 0%, rgba(30, 41, 59, 0.6) 100%);
            backdrop-filter: blur(10px);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            position: relative;
            overflow: hidden;
        }
        
        .stAlert::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 3s infinite;
        }
        
        /* Success alert */
        .stAlert[data-baseweb="notification"][kind="success"] {
            border-left-color: var(--success) !important;
            background: linear-gradient(135deg, rgba(16, 185, 129, 0.15) 0%, var(--bg-surface) 100%);
            box-shadow: 0 0 20px rgba(16, 185, 129, 0.2);
        }
        
        /* Error alert */
        .stAlert[data-baseweb="notification"][kind="error"] {
            border-left-color: var(--danger) !important;
            background: linear-gradient(135deg, rgba(244, 63, 94, 0.15) 0%, var(--bg-surface) 100%);
            box-shadow: 0 0 20px rgba(244, 63, 94, 0.2);
        }
        
        /* Warning alert */
        .stAlert[data-baseweb="notification"][kind="warning"] {
            border-left-color: var(--warning) !important;
            background: linear-gradient(135deg, rgba(234, 179, 8, 0.15) 0%, var(--bg-surface) 100%);
            box-shadow: 0 0 20px rgba(234, 179, 8, 0.2);
        }
        
        /* Info alert */
        .stAlert[data-baseweb="notification"][kind="info"] {
            border-left-color: var(--info) !important;
            background: linear-gradient(135deg, rgba(59, 130, 246, 0.15) 0%, var(--bg-surface) 100%);
            box-shadow: 0 0 20px rgba(59, 130, 246, 0.2);
        }
        
        @keyframes alertSlideIn {
            from {
                opacity: 0;
                transform: translateX(-30px) scale(0.95);
            }
            to {
                opacity: 1;
                transform: translateX(0) scale(1);
            }
        }
        
        /* Success/Error/Warning/Info messages */
        .stSuccess, .stError, .stWarning, .stInfo {
            border-radius: 10px !important;
            padding: 1rem 1.25rem !important;
            font-weight: 500;
            animation: messagePopIn 0.4s cubic-bezier(0.68, -0.55, 0.265, 1.55);
        }
        
        @keyframes messagePopIn {
            0% {
                opacity: 0;
                transform: scale(0.8) translateY(-10px);
            }
            50% {
                transform: scale(1.05) translateY(2px);
            }
            100% {
                opacity: 1;
                transform: scale(1) translateY(0);
            }
        }
        
        /* DataFrames and Tables */
        .stDataFrame {
            border-radius: 10px;
            overflow: hidden;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
        }
        
        .stDataFrame > div {
            background: var(--bg-medium) !important;
        }
        
        .stDataFrame table {
            background: transparent !important;
        }
        
        .stDataFrame thead tr th {
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%) !important;
            color: white !important;
            font-weight: 600 !important;
            text-transform: uppercase !important;
            letter-spacing: 0.05em !important;
            padding: 0.75rem !important;
            border: none !important;
        }
        
        .stDataFrame tbody tr {
            transition: all 0.2s ease;
        }
        
        .stDataFrame tbody tr:hover {
            background: rgba(79, 70, 229, 0.05) !important;
            transform: scale(1.01);
        }
        
        .stDataFrame tbody tr td {
            padding: 0.75rem !important;
            border-bottom: 1px solid var(--border) !important;
            font-weight: 500;
        }
        
        /* Code blocks */
        .stCodeBlock {
            border-radius: 10px !important;
            background: linear-gradient(135deg, #1a1f2e 0%, #0f172a 100%) !important;
            border: 1px solid var(--border) !important;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
        }
        
        .stCodeBlock pre {
            color: #e2e8f0 !important;
            font-family: 'Fira Code', 'Consolas', monospace !important;
            padding: 1rem !important;
        }
        
        /* Columns and containers */
        [data-testid="column"] {
            transition: all 0.3s ease;
            padding: 0.5rem;
        }
        
        [data-testid="column"]:hover {
            transform: translateY(-2px);
        }
        
        /* Aurora sidebar with gradient glow */
        section[data-testid="stSidebar"] {
            background: linear-gradient(135deg, 
                rgba(20, 24, 36, 0.98) 0%, 
                rgba(10, 14, 26, 0.98) 100%);
            border-right: 1px solid rgba(139, 92, 246, 0.2);
            backdrop-filter: blur(20px) saturate(150%);
            transition: all 0.3s ease;
            box-shadow: 2px 0 30px rgba(139, 92, 246, 0.1),
                        inset -1px 0 0 rgba(139, 92, 246, 0.2);
        }
        
        section[data-testid="stSidebar"] > div {
            padding: 1.5rem 1rem;
        }
        
        section[data-testid="stSidebar"] .element-container {
            margin-bottom: 1rem;
        }
        
        /* Sidebar headers */
        section[data-testid="stSidebar"] h1,
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3 {
            color: var(--primary-light) !important;
            border-bottom: 2px solid var(--primary);
            padding-bottom: 0.5rem;
            margin-bottom: 1rem;
        }
        
        /* Multiselect */
        .stMultiSelect > div {
            border-radius: 10px;
            transition: all 0.3s ease;
        }
        
        .stMultiSelect > div:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.1);
        }
        
        /* Date input */
        .stDateInput > div {
            border-radius: 10px;
            transition: all 0.3s ease;
        }
        
        .stDateInput > div:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.1);
        }
        
        /* Spinner animation */
        .stSpinner > div {
            border-color: var(--primary) !important;
        }
        
        /* Enhanced spinner */
        .stSpinner {
            position: relative;
        }
        
        .stSpinner::before {
            content: '⚙️';
            position: absolute;
            font-size: 2rem;
            animation: spin 2s linear infinite;
            opacity: 0.3;
        }
        
        @keyframes spin {
            from { transform: rotate(0deg); }
            to { transform: rotate(360deg); }
        }
        
        /* Empty state */
        .stEmpty {
            padding: 2rem;
            text-align: center;
            color: var(--text-secondary);
            font-style: italic;
        }
        
        /* Modal/Dialog styling */
        [data-baseweb="modal"] {
            background: rgba(15, 23, 42, 0.8) !important;
            backdrop-filter: blur(5px);
        }
        
        [data-baseweb="modal"] > div {
            background: linear-gradient(135deg, var(--bg-medium) 0%, var(--bg-dark) 100%) !important;
            border: 1px solid var(--border);
            border-radius: 12px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.5);
            animation: modalSlideIn 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }
        
        @keyframes modalSlideIn {
            from {
                opacity: 0;
                transform: translateY(-50px) scale(0.9);
            }
            to {
                opacity: 1;
                transform: translateY(0) scale(1);
            }
        }
        
        /* Color picker */
        .stColorPicker > div > div {
            border-radius: 10px;
            border: 2px solid var(--border);
            transition: all 0.3s ease;
        }
        
        .stColorPicker > div > div:hover {
            border-color: var(--primary);
            transform: scale(1.05);
        }
        
        /* Select slider */
        .stSelectSlider > div {
            background: var(--bg-medium);
            border-radius: 10px;
            padding: 0.5rem;
        }
        
        /* JSON viewer */
        .stJson {
            background: linear-gradient(135deg, #1a1f2e 0%, #0f172a 100%) !important;
            border: 1px solid var(--border);
            border-radius: 10px;
            padding: 1rem;
            font-family: 'Fira Code', monospace;
        }
        
        /* Special hover effects for interactive elements */
        .clickable-element {
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .clickable-element:hover {
            transform: translateY(-2px) scale(1.02);
            filter: brightness(1.1);
        }
        
        /* Tooltips */
        [data-baseweb="tooltip"] {
            background: var(--primary) !important;
            border-radius: 8px;
            padding: 0.5rem 0.75rem;
            font-weight: 500;
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
            animation: tooltipFadeIn 0.2s ease;
        }
        
        @keyframes tooltipFadeIn {
            from {
                opacity: 0;
                transform: scale(0.8);
            }
            to {
                opacity: 1;
                transform: scale(1);
            }
        }
        
        /* Premium file uploader with animations */
        [data-testid="stFileUploadDropzone"] {
            background: linear-gradient(135deg, var(--bg-medium) 0%, rgba(30, 41, 59, 0.4) 100%);
            border: 2px dashed var(--border);
            border-radius: 12px;
            padding: 2rem;
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }
        
        [data-testid="stFileUploadDropzone"]::before {
            content: '📁';
            position: absolute;
            font-size: 4rem;
            opacity: 0.1;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            transition: all 0.4s ease;
        }
        
        [data-testid="stFileUploadDropzone"]:hover {
            border-color: var(--primary);
            background: linear-gradient(135deg, rgba(79, 70, 229, 0.1) 0%, rgba(79, 70, 229, 0.05) 100%);
            transform: scale(1.02);
            box-shadow: 0 8px 24px rgba(79, 70, 229, 0.15);
        }
        
        [data-testid="stFileUploadDropzone"]:hover::before {
            opacity: 0.2;
            transform: translate(-50%, -50%) scale(1.2) rotate(5deg);
        }
        
        /* Drag over state */
        [data-testid="stFileUploadDropzone"][data-dragging="true"] {
            border-color: var(--accent);
            background: linear-gradient(135deg, rgba(16, 185, 129, 0.1) 0%, rgba(16, 185, 129, 0.05) 100%);
            animation: pulse 1s infinite;
        }
        
        /* File upload button inside */
        [data-testid="stFileUploadDropzone"] button {
            background: var(--primary) !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            transition: all 0.3s ease !important;
        }
        
        [data-testid="stFileUploadDropzone"] button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
        }
        
        /* Custom scrollbar */
        ::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }
        
        ::-webkit-scrollbar-track {
            background: var(--bg-medium);
        }
        
        ::-webkit-scrollbar-thumb {
            background: var(--border);
            border-radius: 4px;
            transition: background 0.2s ease;
        }
        
        ::-webkit-scrollbar-thumb:hover {
            background: var(--primary);
        }
        
        /* Smooth page transitions */
        * {
            transition: background-color 0.2s ease, 
                        border-color 0.2s ease,
                        box-shadow 0.2s ease;
        }
        
        /* Loading animation */
        @keyframes pulse {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.5; }
        }
        
        /* Fade in animation for content */
        @keyframes fadeIn {
            from { 
                opacity: 0; 
                transform: translateY(10px); 
            }
            to { 
                opacity: 1; 
                transform: translateY(0); 
            }
        }
        
        .element-container {
            animation: fadeIn 0.5s ease;
        }
    </style>
    """, unsafe_allow_html=True)

# AI Thinking Messages
AI_THINKING_MESSAGES = [
    "🧠 Analyzing requirements with advanced AI...",
    "🔍 Searching healthcare compliance standards...",
    "💡 Identifying edge cases and test scenarios...",
    "🏥 Applying medical device testing best practices...",
    "📊 Calculating optimal test coverage...",
    "🔬 Synthesizing comprehensive test cases...",
    "⚡ Optimizing test execution flow...",
    "🎯 Focusing on patient safety requirements...",
    "🛡️ Ensuring FDA and ISO compliance...",
    "🤖 Leveraging machine learning patterns..."
]

def show_ai_thinking(duration=3):
    """Display animated AI thinking messages with smooth transitions"""
    placeholder = st.empty()
    progress = st.progress(0)
    
    for i in range(duration * 2):
        msg = random.choice(AI_THINKING_MESSAGES)
        # Clean, animated message display
        placeholder.markdown(f"""
        <div style='text-align: center; 
                    padding: 1.5rem; 
                    background: #1E293B; 
                    border: 1px solid #475569; 
                    border-radius: 8px; 
                    margin: 1rem 0;
                    animation: slideIn 0.3s ease;'>
            <div style='font-size: 2rem; margin-bottom: 0.5rem; animation: pulse 1s ease infinite;'>
                {'🧠🔍💡🏥📊🔬⚡🎯🛡️🤖'[i % 10]}
            </div>
            <p style='color: #F1F5F9; font-size: 1rem; font-weight: 500; margin: 0;'>
                {msg}
            </p>
            <div style='margin-top: 0.5rem; color: #6366F1; font-size: 0.875rem;'>
                {int((i + 1) / (duration * 2) * 100)}% Complete
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        progress.progress((i + 1) / (duration * 2))
        time.sleep(0.5)
    
    placeholder.empty()
    progress.empty()

# Configuration - SIMPLIFIED FOLDER STRUCTURE (v9)
# Knowledge Base & Documents
DOCUMENTS_PATH = "data/documents"  # Sample docs for RAG (PRDs, user stories, API specs)
UPLOADED_DOCS_PATH = "data/uploaded_documents"  # User uploads + compliance reports

# Test Cases - SINGLE UNIFIED LOCATION
TEST_CASES_DIR = "data/test_cases"  # ALL test cases (generated, imported, manual saves)

# Vector Database
FAISS_INDEX_PATH = "data/faiss_indices"  # FAISS vector database for RAG

# Models
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2'
GEMINI_MODEL_NAME = 'gemini-2.0-flash-exp'  # Latest model

# NASSCOM Compliance Requirements
NASSCOM_REQUIREMENTS = {
    "document_types": {
        "prd": "Product Requirement Document",
        "user_story": "User Stories with acceptance criteria",
        "api_spec": "API Specifications (OpenAPI/Swagger)",
        "mockup": "UI/UX Mockups and wireframes",
        "test_plan": "Existing test plans or strategies"
    },
    "required_fields": [
        "Test Case ID", "Title", "Description", 
        "Test Steps", "Expected Results", "Priority"
    ],
    "compliance_standards": [
        "HIPAA", "GDPR", "FDA", "ISO 9001", 
        "ISO 13485", "ISO 27001"
    ],
    "healthcare_keywords": [
        "patient", "medical", "health", "clinical", "diagnosis",
        "treatment", "medication", "prescription", "PHI", "EHR",
        "DICOM", "HL7", "medical device", "FDA compliance"
    ],
    "document_formats": ["pdf", "docx", "doc", "xml", "yaml", "yml", "txt", "md"],
    "quality_criteria": {
        "min_content_length": 100,
        "requires_structure": True,
        "requires_healthcare_context": False  # Made optional for flexibility
    }
}

# Page configuration
st.set_page_config(
    page_title="AI Test Case Generator - NASSCOM", 
    page_icon="🧪",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 10px;
        border-radius: 5px;
        margin: 5px;
    }
    .compliance-pass {
        background-color: #d4edda;
        color: #155724;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
    .compliance-fail {
        background-color: #f8d7da;
        color: #721c24;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
    .compliance-warning {
        background-color: #fff3cd;
        color: #856404;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'test_counter' not in st.session_state:
    st.session_state.test_counter = 0
if 'generated_tests' not in st.session_state:
    st.session_state.generated_tests = []
if 'refinement_mode' not in st.session_state:
    st.session_state.refinement_mode = False
if 'test_to_refine' not in st.session_state:
    st.session_state.test_to_refine = None
if 'refinement_history' not in st.session_state:
    st.session_state.refinement_history = {}
if 'edited_test' not in st.session_state:
    st.session_state.edited_test = None
if 'test_versions' not in st.session_state:
    st.session_state.test_versions = {}
if 'uploaded_docs' not in st.session_state:
    st.session_state.uploaded_docs = []
if 'compliance_reports' not in st.session_state:
    st.session_state.compliance_reports = {}
if 'imported_tests' not in st.session_state:
    st.session_state.imported_tests = []
if 'import_history' not in st.session_state:
    st.session_state.import_history = []
if 'last_converted_tests' not in st.session_state:
    st.session_state.last_converted_tests = None
if 'last_import_report' not in st.session_state:
    st.session_state.last_import_report = None
if 'tests_loaded' not in st.session_state:
    st.session_state.tests_loaded = False

# Configure Gemini
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# ===============================
# Auto-Save/Load Functions (NEW IN V7)
# ===============================

def auto_save_test_cases(test_cases: List[Dict], prefix: str = "tests"):
    """Automatically save test cases to unified directory for persistence"""
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    
    # Save as a single JSON file with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{TEST_CASES_DIR}/{prefix}_{timestamp}.json"
    
    # Also save as 'latest' for easy loading
    latest_filename = f"{TEST_CASES_DIR}/{prefix}_latest.json"
    
    try:
        # Save timestamped version
        with open(filename, 'w') as f:
            json.dump(test_cases, f, indent=2, default=str)
        
        # Save/overwrite latest version
        with open(latest_filename, 'w') as f:
            json.dump(test_cases, f, indent=2, default=str)
        
        return True, filename
    except Exception as e:
        st.error(f"Failed to auto-save tests: {e}")
        return False, None

def load_saved_test_cases(prefix: str = "tests") -> List[Dict]:
    """Load previously saved test cases from unified directory"""
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    
    # Try to load the latest file first
    latest_filename = f"{TEST_CASES_DIR}/{prefix}_latest.json"
    
    try:
        if os.path.exists(latest_filename):
            with open(latest_filename, 'r') as f:
                test_cases = json.load(f)
                return test_cases if isinstance(test_cases, list) else []
    except Exception as e:
        # If latest fails, try to load most recent timestamped file
        try:
            files = glob.glob(f"{TEST_CASES_DIR}/{prefix}_*.json")
            # Filter out 'latest' file and sort by modification time
            files = [f for f in files if not f.endswith('_latest.json')]
            if files:
                files.sort(key=os.path.getmtime, reverse=True)
                with open(files[0], 'r') as f:
                    test_cases = json.load(f)
                    return test_cases if isinstance(test_cases, list) else []
        except Exception as inner_e:
            st.warning(f"Could not load saved tests: {inner_e}")
    
    return []

def cleanup_old_saves(prefix: str = "*", keep_count: int = 10):
    """Keep only the most recent saves to avoid disk bloat"""
    try:
        files = glob.glob(f"{TEST_CASES_DIR}/{prefix}_*.json")
        # Exclude 'latest' files
        files = [f for f in files if not f.endswith('_latest.json')]
        
        if len(files) > keep_count:
            # Sort by modification time
            files.sort(key=os.path.getmtime)
            # Delete oldest files
            for f in files[:-keep_count]:
                os.remove(f)
    except Exception as e:
        pass  # Silently fail cleanup

# Load saved tests on startup
@st.cache_resource(show_spinner=False)
def initialize_saved_tests():
    """Load previously saved tests on app startup - unified from single directory"""
    all_saved_tests = load_saved_test_cases("all_tests")
    if not all_saved_tests:
        # For backward compatibility, try loading old format
        generated = load_saved_test_cases("generated")
        imported = load_saved_test_cases("imported")
        all_saved_tests = generated + imported
    return all_saved_tests

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from Word document"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting Word document text: {e}")
        return ""

def extract_text_from_xml(file_content: bytes) -> str:
    """Extract text content from XML file"""
    try:
        # Try to detect encoding
        detected = chardet.detect(file_content)
        encoding = detected['encoding'] or 'utf-8'
        
        text_content = file_content.decode(encoding)
        root = ET.fromstring(text_content)
        
        # Convert XML to readable format
        def extract_element_text(element, prefix=""):
            result = []
            if element.tag:
                result.append(f"{prefix}{element.tag}:")
            if element.text and element.text.strip():
                result.append(f"{prefix}  {element.text.strip()}")
            for attr_name, attr_value in element.attrib.items():
                result.append(f"{prefix}  @{attr_name}: {attr_value}")
            for child in element:
                result.extend(extract_element_text(child, prefix + "  "))
            return result
        
        text_lines = extract_element_text(root)
        return "\n".join(text_lines)
    except Exception as e:
        st.error(f"Error extracting XML text: {e}")
        return ""

def analyze_document_compliance(content: str, filename: str, file_type: str) -> Dict:
    """
    Analyze document for NASSCOM compliance using Gemini AI
    """
    prompt = f"""You are a compliance analyst for NASSCOM's AI Test Case Generation requirements.
Analyze the following document and determine its compliance with NASSCOM guidelines.

DOCUMENT NAME: {filename}
FILE TYPE: {file_type}

DOCUMENT CONTENT:
{content[:3000]}  # Limiting to first 3000 chars for analysis

NASSCOM REQUIREMENTS:
1. Document Types Accepted: PRDs, User Stories, API Specs, Mockups, Test Plans
2. Healthcare/MedTech Context: Should contain healthcare-related terminology
3. Structure: Should have clear sections and be well-organized
4. Content Quality: Sufficient detail for test case generation
5. Compliance Standards: Should reference or be compatible with HIPAA, GDPR, FDA, ISO standards

Analyze and return a JSON object with:
{{
    "is_compliant": boolean,
    "compliance_score": float (0-100),
    "document_type": string (prd/user_story/api_spec/mockup/test_plan/other),
    "detected_standards": [list of compliance standards found],
    "healthcare_relevance": float (0-100),
    "structural_quality": float (0-100),
    "content_completeness": float (0-100),
    "strengths": [list of compliance strengths],
    "improvements_needed": [list of improvements needed],
    "recommendations": [list of specific recommendations],
    "summary": string (brief compliance summary)
}}"""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        compliance_report = json.loads(response.text)
        
        # Add metadata
        compliance_report['analyzed_at'] = datetime.now().isoformat()
        compliance_report['filename'] = filename
        compliance_report['file_type'] = file_type
        compliance_report['content_length'] = len(content)
        
        return compliance_report
        
    except Exception as e:
        # Fallback basic analysis
        return {
            "is_compliant": len(content) > 100,
            "compliance_score": 50.0 if len(content) > 100 else 20.0,
            "document_type": "other",
            "detected_standards": [],
            "healthcare_relevance": 0.0,
            "structural_quality": 50.0,
            "content_completeness": 50.0,
            "strengths": ["Document uploaded successfully"],
            "improvements_needed": ["Could not perform AI analysis"],
            "recommendations": ["Please ensure document contains relevant content"],
            "summary": f"Basic analysis completed. Error in AI analysis: {str(e)}",
            "analyzed_at": datetime.now().isoformat(),
            "filename": filename,
            "file_type": file_type,
            "content_length": len(content)
        }

def process_uploaded_file(uploaded_file) -> Tuple[str, Dict]:
    """
    Process uploaded file and extract content with compliance analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    
    # Read file content
    file_content = uploaded_file.read()
    
    # Extract text based on file type
    if file_extension == 'pdf':
        text_content = extract_text_from_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        text_content = extract_text_from_docx(file_content)
    elif file_extension == 'xml':
        text_content = extract_text_from_xml(file_content)
    elif file_extension in ['yaml', 'yml']:
        try:
            text_content = file_content.decode('utf-8')
            yaml_data = yaml.safe_load(text_content)
            text_content = yaml.dump(yaml_data, default_flow_style=False)
        except:
            text_content = file_content.decode('utf-8', errors='ignore')
    elif file_extension in ['txt', 'md']:
        text_content = file_content.decode('utf-8', errors='ignore')
    else:
        # Try to decode as text
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
        except:
            text_content = ""
    
    # Analyze compliance
    compliance_report = analyze_document_compliance(
        text_content, 
        uploaded_file.name, 
        file_extension
    )
    
    return text_content, compliance_report

def save_uploaded_document(filename: str, content: str, compliance_report: Dict) -> bool:
    """
    Save uploaded document to the documents folder if compliant
    """
    try:
        # Create uploaded documents directory if it doesn't exist
        os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
        
        # Save the document
        file_path = os.path.join(UPLOADED_DOCS_PATH, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Save compliance report
        report_path = os.path.join(UPLOADED_DOCS_PATH, f"{filename}.compliance.json")
        with open(report_path, 'w') as f:
            json.dump(compliance_report, f, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Error saving document: {e}")
        return False

def display_compliance_report(report: Dict):
    """
    Display compliance analysis report with visual indicators
    """
    # Overall compliance status
    if report['is_compliant']:
        st.markdown('<div class="compliance-pass">✅ Document is NASSCOM Compliant</div>', 
                   unsafe_allow_html=True)
    else:
        st.markdown('<div class="compliance-fail">❌ Document needs improvements for full compliance</div>', 
                   unsafe_allow_html=True)
    
    # Compliance metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        score_color = "🟢" if report['compliance_score'] >= 70 else "🟡" if report['compliance_score'] >= 40 else "🔴"
        st.metric("Compliance Score", f"{score_color} {report['compliance_score']:.1f}%")
    
    with col2:
        relevance_color = "🟢" if report['healthcare_relevance'] >= 60 else "🟡" if report['healthcare_relevance'] >= 30 else "🔴"
        st.metric("Healthcare Relevance", f"{relevance_color} {report['healthcare_relevance']:.1f}%")
    
    with col3:
        quality_color = "🟢" if report['structural_quality'] >= 70 else "🟡" if report['structural_quality'] >= 40 else "🔴"
        st.metric("Structural Quality", f"{quality_color} {report['structural_quality']:.1f}%")
    
    with col4:
        completeness_color = "🟢" if report['content_completeness'] >= 70 else "🟡" if report['content_completeness'] >= 40 else "🔴"
        st.metric("Content Completeness", f"{completeness_color} {report['content_completeness']:.1f}%")
    
    # Document type and standards
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Document Analysis")
        st.write(f"**Document Type:** {report.get('document_type', 'Unknown').replace('_', ' ').title()}")
        st.write(f"**File Type:** {report.get('file_type', 'Unknown').upper()}")
        st.write(f"**Content Length:** {report.get('content_length', 0):,} characters")
        
        if report.get('detected_standards'):
            st.write("**Detected Standards:**")
            for standard in report['detected_standards']:
                st.write(f"  • {standard}")
    
    with col2:
        st.markdown("### 📊 Compliance Summary")
        st.info(report.get('summary', 'No summary available'))
    
    # Strengths and improvements in tabs
    tab1, tab2, tab3 = st.tabs(["✅ Strengths", "⚠️ Improvements Needed", "💡 Recommendations"])
    
    with tab1:
        if report.get('strengths'):
            for strength in report['strengths']:
                st.success(f"• {strength}")
        else:
            st.info("No specific strengths identified")
    
    with tab2:
        if report.get('improvements_needed'):
            for improvement in report['improvements_needed']:
                st.warning(f"• {improvement}")
        else:
            st.success("No improvements needed")
    
    with tab3:
        if report.get('recommendations'):
            for rec in report['recommendations']:
                st.info(f"• {rec}")
        else:
            st.info("No specific recommendations")

@st.cache_resource
def load_rag_system():
    """
    Load embedding model and build FAISS index from documents.
    Includes both pre-existing and uploaded documents.
    """
    # Load embedding model
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    
    # Create necessary directories
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    os.makedirs(DOCUMENTS_PATH, exist_ok=True)
    os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
    os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
    
    # Load and process documents from both paths
    doc_chunks = []
    doc_metadata = []
    
    # Process existing documents
    existing_files = glob.glob(f"{DOCUMENTS_PATH}/*")
    
    # Process uploaded documents
    uploaded_files = glob.glob(f"{UPLOADED_DOCS_PATH}/*")
    # Filter out compliance report files
    uploaded_files = [f for f in uploaded_files if not f.endswith('.compliance.json')]
    
    all_files = existing_files + uploaded_files
    
    if not all_files:
        st.warning(f"No documents found. Upload documents to get started.")
        return embedding_model, None, None, []
    
    # Process each document
    for file_path in all_files:
        file_name = Path(file_path).name
        file_ext = Path(file_path).suffix
        
        # Determine if it's an uploaded document
        is_uploaded = UPLOADED_DOCS_PATH in file_path
        
        try:
            if file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                    # Smart chunking for different document types
                    if 'user_story' in file_name.lower():
                        # Split by acceptance criteria
                        if 'Acceptance Criteria:' in content:
                            parts = content.split('Acceptance Criteria:')
                            if parts[0].strip():
                                doc_chunks.append(parts[0].strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'user_story_overview',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                            if len(parts) > 1:
                                doc_chunks.append(f"Acceptance Criteria:\n{parts[1].strip()}")
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'acceptance_criteria',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                        else:
                            doc_chunks.append(content)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'user_story',
                                'doc_type': 'user_story',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # Split by paragraphs for other documents
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip() and len(para.strip()) > 50:
                                doc_chunks.append(para.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'prd' if 'prd' in file_name.lower() else 'document',
                                    'doc_type': 'prd' if 'prd' in file_name.lower() else 'general',
                                    'is_uploaded': is_uploaded
                                })
                                
            elif file_ext in ['.yaml', '.yml']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    yaml_content = yaml.safe_load(f)
                    
                    # Process API specs
                    if 'paths' in yaml_content:
                        for path, methods in yaml_content.get('paths', {}).items():
                            endpoint_doc = f"API Endpoint: {path}\n"
                            endpoint_doc += yaml.dump(methods, default_flow_style=False)
                            doc_chunks.append(endpoint_doc)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'api_endpoint',
                                'endpoint': path,
                                'doc_type': 'api_specification',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # General YAML content
                        doc_chunks.append(yaml.dump(yaml_content, default_flow_style=False))
                        doc_metadata.append({
                            'source': file_name,
                            'type': 'config',
                            'doc_type': 'configuration',
                            'is_uploaded': is_uploaded
                        })
            
            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_pdf(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'pdf_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext in ['.docx', '.doc']:
                with open(file_path, 'rb') as f:
                    text = extract_text_from_docx(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'word_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext == '.xml':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_xml(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'xml_document',
                                    'doc_type': 'specification',
                                    'is_uploaded': is_uploaded
                                })
                        
        except Exception as e:
            st.error(f"Error loading {file_name}: {str(e)}")
    
    if not doc_chunks:
        st.warning("No content extracted from documents.")
        return embedding_model, None, None, []
    
    # Generate embeddings
    with st.spinner(f"Creating embeddings for {len(doc_chunks)} document chunks..."):
        contents = doc_chunks
        embeddings = embedding_model.encode(contents, show_progress_bar=True)
        embeddings = np.array(embeddings).astype('float32')
    
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    
    return embedding_model, index, doc_chunks, doc_metadata

def retrieve_context(query: str, embedding_model, index, doc_chunks, doc_metadata, k: int = 5):
    """Retrieve relevant documents for a query"""
    if index is None:
        return []
    
    query_embedding = embedding_model.encode([query]).astype('float32')
    distances, indices = index.search(query_embedding, k)
    
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(doc_chunks):
            similarity = 1 / (1 + dist)
            results.append({
                'content': doc_chunks[idx],
                'metadata': doc_metadata[idx] if idx < len(doc_metadata) else {},
                'similarity': similarity
            })
    
    return results

def generate_test_case_with_gemini(requirement: str, context_docs: List[Dict]) -> Dict:
    """
    Generate test case using Gemini with NASSCOM compliance
    """
    # Check context relevance
    avg_relevance = sum(doc['similarity'] for doc in context_docs) / len(context_docs) if context_docs else 0
    low_relevance = avg_relevance < 0.3
    
    # Build context string
    if low_relevance:
        context_str = "Note: Limited specific context found. Generating based on general best practices and requirement analysis."
    else:
        context_parts = []
        for doc in context_docs[:3]:
            context_parts.append(f"--- Source: {doc['metadata'].get('source', 'Unknown')} (Type: {doc['metadata'].get('type', 'document')}, Relevance: {doc['similarity']:.2%}) ---")
            context_parts.append(doc['content'][:800])
            context_parts.append("--- End Context ---")
        context_str = "\n\n".join(context_parts)
    
    # Enhanced prompt with NASSCOM requirements
    prompt = f"""You are a meticulous QA Engineer specializing in healthcare/MedTech testing, following NASSCOM guidelines.
Generate a comprehensive test case for the following requirement.

REQUIREMENT:
{requirement}

RELEVANT CONTEXT:
{context_str}

NASSCOM COMPLIANCE REQUIREMENTS:
- Must include all required fields: Test Case ID, Title, Description, Test Steps, Expected Results, Priority
- Should consider healthcare compliance standards: HIPAA, GDPR, FDA, ISO 13485, ISO 27001
- Should handle edge cases and boundary conditions
- Must be traceable to requirements
- Should support both manual and automated testing where feasible

Generate a test case with EXACTLY these fields:
- id: Use "TC_PLACEHOLDER" (will be replaced)
- title: Clear, descriptive title
- description: Detailed description of what is being tested
- category: One of [Functional, Security, Integration, Performance, Usability, Compliance]
- priority: One of [Critical, High, Medium, Low]
- compliance: Array of applicable standards like ["HIPAA", "GDPR", "FDA", "ISO 13485", "ISO 27001"]
- preconditions: String describing what must be true before testing
- test_steps: Array of step descriptions WITHOUT numbering (e.g., ["Navigate to login page", "Enter credentials"])
- expected_results: String describing specific expected outcomes
- test_data: JSON object with test data (e.g., {{"username": "test@example.com", "password": "Test123!"}})
- edge_cases: Array of special scenarios to consider
- negative_tests: Array of negative test scenarios
- automation_feasible: Boolean indicating if this can be automated
- estimated_duration: String like "5 minutes" or "30 minutes"
- traceability: String linking to requirement source

IMPORTANT:
1. test_steps must be an array of strings without prefixes like "Step 1:" or "1."
2. test_data must be a JSON object, not a string
3. Focus on healthcare/medical domain requirements if applicable
4. Include both positive and negative test scenarios
5. Ensure compliance with relevant healthcare standards

Return ONLY a valid JSON object."""
    
    try:
        # Configure Gemini for JSON output
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                top_k=30,
                max_output_tokens=4096
            )
        )
        
        # Safety settings for healthcare content
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
        
        response = model.generate_content(prompt, safety_settings=safety_settings)
        
        if not response.parts:
            return create_fallback_test_case(requirement, error_reason="Content blocked by safety filter")
        
        test_case = json.loads(response.text)
        
        # Generate unique ID
        unique_id = f"TC_{str(uuid.uuid4())[:8].upper()}"
        test_case['id'] = unique_id
        
        # Clean test_steps formatting
        if 'test_steps' in test_case and isinstance(test_case['test_steps'], list):
            cleaned_steps = []
            for step in test_case['test_steps']:
                cleaned = re.sub(r'^(Step\s+\d+:|^\d+\.\s*)', '', str(step)).strip()
                cleaned_steps.append(cleaned)
            test_case['test_steps'] = cleaned_steps
        
        # Ensure test_data is always a dict
        if 'test_data' in test_case:
            if isinstance(test_case['test_data'], str):
                try:
                    test_case['test_data'] = json.loads(test_case['test_data'])
                except:
                    test_case['test_data'] = {"raw_data": test_case['test_data']}
        else:
            test_case['test_data'] = {}
        
        # Add metadata and context tracking
        test_case['generated_from'] = requirement[:100]
        test_case['generation_timestamp'] = datetime.now().isoformat()
        test_case['avg_context_relevance'] = f"{avg_relevance:.2%}"
        test_case['context_sources'] = [
            doc['metadata'].get('source', 'Unknown') for doc in context_docs[:3]
        ] if not low_relevance else ["Generated using best practices"]
        test_case['low_context_confidence'] = low_relevance
        test_case['version'] = 1
        test_case['nasscom_compliant'] = True
        
        # Increment counter
        st.session_state.test_counter += 1
        
        return test_case
        
    except Exception as e:
        st.error(f"Generation error: {e}")
        return create_fallback_test_case(requirement, error_reason=str(e))

def create_fallback_test_case(requirement: str, error_reason: str = "Unknown") -> Dict:
    """
    Create a fallback test case when generation fails.
    """
    st.session_state.test_counter += 1
    return {
        'id': f'TC_FB_{st.session_state.test_counter:04d}',
        'title': f'Test: {requirement[:60]}',
        'description': f'Verify that {requirement}',
        'category': 'Functional',
        'priority': 'Medium',
        'compliance': ['General'],
        'preconditions': 'System is in stable state',
        'test_steps': [
            'Setup test environment',
            'Execute test scenario',
            'Verify results'
        ],
        'expected_results': 'System behaves as specified',
        'test_data': {'placeholder': 'Add specific test data'},
        'edge_cases': [],
        'negative_tests': [],
        'automation_feasible': False,
        'estimated_duration': '10 minutes',
        'traceability': requirement[:100],
        'generated_from': requirement[:100],
        'context_sources': ['Fallback generation'],
        'generation_timestamp': datetime.now().isoformat(),
        'nasscom_compliant': False,
        'fallback': True,
        'error': error_reason
    }

def analyze_test_suite_schema(content: str, file_type: str) -> Dict:
    """
    Use AI to analyze the schema/structure of an uploaded test suite
    """
    prompt = f"""You are an expert test case analyst. Analyze the following test suite data and understand its schema.

FILE TYPE: {file_type}
CONTENT SAMPLE (first 2000 chars):
{content[:2000]}

Analyze the test suite and return a JSON object with:
{{
    "detected_format": string (csv/json/excel/xml/custom),
    "has_headers": boolean,
    "field_mappings": {{
        "id_field": string or null,
        "title_field": string or null,
        "description_field": string or null,
        "steps_field": string or null,
        "expected_field": string or null,
        "priority_field": string or null,
        "category_field": string or null,
        "preconditions_field": string or null,
        "test_data_field": string or null
    }},
    "detected_fields": [list of all field names found],
    "row_count": integer (estimated number of test cases),
    "separator": string (for CSV - comma, tab, pipe, etc),
    "schema_confidence": float (0-100),
    "sample_data": object (one sample test case in original format),
    "recommendations": [list of recommendations for import]
}}

Be intelligent about detecting field mappings even if names don't match exactly.
For example, "Test Name" could map to "title_field", "Steps" to "steps_field", etc."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.2
            )
        )
        
        response = model.generate_content(prompt)
        schema_analysis = json.loads(response.text)
        return schema_analysis
        
    except Exception as e:
        st.error(f"Schema analysis error: {e}")
        return {
            "detected_format": "unknown",
            "has_headers": True,
            "field_mappings": {},
            "detected_fields": [],
            "row_count": 0,
            "schema_confidence": 0,
            "sample_data": {},
            "recommendations": ["Manual mapping may be required"]
        }

def convert_test_with_ai(test_data: Dict, schema: Dict, import_id: str) -> Dict:
    """
    Use AI to intelligently convert a test from any format to our standard format
    """
    prompt = f"""You are converting a test case to NASSCOM-compliant format.

ORIGINAL TEST DATA:
{json.dumps(test_data, indent=2)}

DETECTED SCHEMA MAPPINGS:
{json.dumps(schema.get('field_mappings', {}), indent=2)}

Convert this test case to our standard format with these EXACT fields:
{{
    "id": string (generate unique ID starting with {import_id}),
    "title": string (clear, descriptive title),
    "description": string (detailed description),
    "category": string (Functional/Security/Integration/Performance/Usability/Compliance),
    "priority": string (Critical/High/Medium/Low),
    "compliance": array (applicable standards like ["HIPAA", "GDPR", "FDA"]),
    "preconditions": string (what must be true before testing),
    "test_steps": array of strings (clear step descriptions),
    "expected_results": string (specific expected outcomes),
    "test_data": object (test data as JSON object),
    "edge_cases": array (special scenarios),
    "negative_tests": array (negative test scenarios),
    "automation_feasible": boolean,
    "estimated_duration": string (e.g., "10 minutes"),
    "traceability": string (original requirement reference),
    "imported": true,
    "original_format": "{schema.get('detected_format', 'unknown')}",
    "import_timestamp": "{datetime.now().isoformat()}",
    "nasscom_compliant": true
}}

Be intelligent about:
1. Parsing multi-line text into arrays for test_steps
2. Inferring missing fields from context
3. Detecting healthcare/medical context for compliance standards
4. Splitting combined fields appropriately
5. Converting string representations to proper types

Return ONLY a valid JSON object."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        converted_test = json.loads(response.text)
        
        # Ensure unique ID
        if not converted_test.get('id'):
            converted_test['id'] = f"{import_id}_{str(uuid.uuid4())[:8].upper()}"
        
        return converted_test
        
    except Exception as e:
        # Fallback conversion
        return {
            'id': f"{import_id}_{str(uuid.uuid4())[:8].upper()}",
            'title': test_data.get('title', test_data.get('name', 'Imported Test')),
            'description': test_data.get('description', str(test_data)),
            'category': 'Functional',
            'priority': 'Medium',
            'compliance': [],
            'preconditions': test_data.get('preconditions', 'N/A'),
            'test_steps': [str(test_data.get('steps', 'See original data'))],
            'expected_results': test_data.get('expected', 'As specified'),
            'test_data': test_data,
            'imported': True,
            'import_error': str(e),
            'nasscom_compliant': False
        }

def detect_duplicates(new_tests: List[Dict], existing_tests: List[Dict]) -> List[Dict]:
    """
    Detect potential duplicate test cases using AI
    """
    if not existing_tests or not new_tests:
        return []
    
    duplicates = []
    
    # Create summaries for comparison
    existing_summaries = [
        f"{t.get('title', '')} - {t.get('description', '')[:100]}"
        for t in existing_tests
    ]
    
    for new_test in new_tests:
        new_summary = f"{new_test.get('title', '')} - {new_test.get('description', '')[:100]}"
        
        # Quick check for exact matches
        if new_summary in existing_summaries:
            duplicates.append({
                'new_test': new_test,
                'match_type': 'exact',
                'confidence': 100
            })
            continue
        
        # Use AI for semantic similarity (simplified for performance)
        for i, existing_summary in enumerate(existing_summaries[:10]):  # Check first 10 for demo
            if len(new_summary) > 20 and len(existing_summary) > 20:
                # Simple similarity check
                common_words = set(new_summary.lower().split()) & set(existing_summary.lower().split())
                if len(common_words) > 5:
                    duplicates.append({
                        'new_test': new_test,
                        'existing_test': existing_tests[i],
                        'match_type': 'similar',
                        'confidence': len(common_words) * 10
                    })
    
    return duplicates

def import_test_suite_with_ai(uploaded_file) -> Tuple[List[Dict], Dict]:
    """
    Main function to import test suite using AI analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    file_content = uploaded_file.read()
    
    # Extract content based on file type
    if file_extension == 'csv':
        content = file_content.decode('utf-8', errors='ignore')
        # Parse CSV
        import csv
        from io import StringIO
        reader = csv.DictReader(StringIO(content))
        test_data = list(reader)
    elif file_extension == 'json':
        content = file_content.decode('utf-8', errors='ignore')
        test_data = json.loads(content)
        if not isinstance(test_data, list):
            test_data = [test_data]
    elif file_extension in ['xlsx', 'xls']:
        # Use pandas for Excel
        df = pd.read_excel(io.BytesIO(file_content))
        test_data = df.to_dict('records')
        content = df.to_string()
    else:
        content = file_content.decode('utf-8', errors='ignore')
        # Try to parse as structured data
        test_data = [{'raw_content': content}]
    
    # Analyze schema with AI
    with st.spinner("🤖 Analyzing test suite structure with AI..."):
        schema = analyze_test_suite_schema(content, file_extension)
    
    # Generate import ID
    import_id = f"IMP_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Convert tests using AI
    converted_tests = []
    with st.spinner(f"🔄 Converting {len(test_data)} test cases..."):
        progress_bar = st.progress(0)
        for i, test in enumerate(test_data):
            converted = convert_test_with_ai(test, schema, import_id)
            converted_tests.append(converted)
            progress_bar.progress((i + 1) / len(test_data))
    
    # Check for duplicates
    with st.spinner("🔍 Checking for duplicates..."):
        duplicates = detect_duplicates(converted_tests, st.session_state.generated_tests)
    
    # Create import report
    import_report = {
        'import_id': import_id,
        'filename': uploaded_file.name,
        'timestamp': datetime.now().isoformat(),
        'schema_analysis': schema,
        'total_tests': len(test_data),
        'converted_tests': len(converted_tests),
        'duplicates_found': len(duplicates),
        'success_rate': sum(1 for t in converted_tests if t.get('nasscom_compliant', False)) / len(converted_tests) * 100 if converted_tests else 0
    }
    
    return converted_tests, import_report

def display_import_interface():
    """
    Display the AI-powered test import interface
    """
    st.header("📥 AI-Powered Test Suite Import")
    st.markdown("""
    Upload your existing test suites in any format (CSV, Excel, JSON, etc.).
    Our AI will automatically:
    - 🔍 Analyze the structure and schema
    - 🔄 Convert tests to NASSCOM-compliant format
    - 🔎 Detect and handle duplicates
    - ✅ Validate and enhance test quality
    """)
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose test suite file",
        type=['csv', 'xlsx', 'xls', 'json', 'xml', 'txt'],
        help="Upload CSV, Excel, JSON or other test suite formats"
    )
    
    if uploaded_file:
        # Import options
        col1, col2 = st.columns(2)
        with col1:
            merge_duplicates = st.checkbox("Merge duplicate tests", value=True)
        with col2:
            enhance_with_ai = st.checkbox("Enhance tests with AI", value=True)
        
        if st.button("🚀 Start AI Import", type="primary"):
            # Perform import
            converted_tests, import_report = import_test_suite_with_ai(uploaded_file)
            
            # Store in session state
            st.session_state.imported_tests.extend(converted_tests)
            st.session_state.import_history.append(import_report)
            st.session_state['last_converted_tests'] = converted_tests
            st.session_state['last_import_report'] = import_report
            
            # Display results
            st.success(f"✅ Successfully imported {len(converted_tests)} test cases!")
        
        # Show results and add button OUTSIDE the import button condition
        if 'last_converted_tests' in st.session_state and st.session_state['last_converted_tests']:
            # Show import report
            if 'last_import_report' in st.session_state:
                with st.expander("📊 Import Report", expanded=True):
                    import_report = st.session_state['last_import_report']
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Tests", import_report['total_tests'])
                    with col2:
                        st.metric("Converted", import_report['converted_tests'])
                    with col3:
                        st.metric("Duplicates", import_report['duplicates_found'])
                    with col4:
                        st.metric("Success Rate", f"{import_report['success_rate']:.1f}%")
                    
                    # Schema analysis
                    st.markdown("### 📋 Detected Schema")
                    schema = import_report['schema_analysis']
                    st.json(schema.get('field_mappings', {}))
                    
                    if schema.get('recommendations'):
                        st.markdown("### 💡 Recommendations")
                        for rec in schema['recommendations']:
                            st.info(f"• {rec}")
            
            # Show converted tests preview
            converted_tests = st.session_state['last_converted_tests']
            with st.expander("👁️ Preview Imported Tests", expanded=False):
                for test in converted_tests[:5]:  # Show first 5
                    st.markdown(f"**{test['id']}**: {test['title']}")
                    st.write(f"Priority: {test['priority']} | Category: {test['category']}")
                if len(converted_tests) > 5:
                    st.info(f"... and {len(converted_tests) - 5} more tests")
            
                # Add to main test suite button - NOW OUTSIDE THE IMPORT CONDITION
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Add to Test Suite", key="add_imported_to_suite", type="primary"):
                    tests_to_add = st.session_state['last_converted_tests']
                    st.session_state.generated_tests.extend(tests_to_add)
                    
                    # AUTO-SAVE all tests to unified location
                    saved, filename = auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                    if saved:
                        st.success(f"✅ Added {len(tests_to_add)} tests to main suite and auto-saved!")
                    else:
                        st.success(f"✅ Added {len(tests_to_add)} tests to main suite!")
                    
                    # Clean up old saves to prevent disk bloat
                    cleanup_old_saves("all_tests", keep_count=10)
                    
                    # Clear the temporary storage
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.balloons()
                    st.info("Go to '🗂️ Test Suite' tab to view your tests")
            with col2:
                if st.button("❌ Cancel Import", key="cancel_import"):
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.info("Import cancelled")
    
    # Show import history
    if st.session_state.import_history:
        st.markdown("### 📜 Import History")
        for report in reversed(st.session_state.import_history[-5:]):  # Show last 5
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                st.write(f"📄 {report['filename']}")
            with col2:
                st.write(f"Tests: {report['converted_tests']}")
            with col3:
                st.write(f"✅ {report['success_rate']:.0f}%")

def display_upload_interface():
    """
    Display document upload interface with NASSCOM compliance validation
    """
    st.header("📤 Upload Documents for Test Generation")
    st.markdown("""
    Upload your documents to expand the knowledge base for test case generation.
    Documents will be analyzed for NASSCOM compliance before being added to the system.
    """)
    
    # Show accepted formats
    with st.expander("📋 NASSCOM Document Requirements", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### ✅ Accepted Document Types")
            for doc_type, description in NASSCOM_REQUIREMENTS["document_types"].items():
                st.write(f"• **{description}**")
            
            st.markdown("### 📁 Supported Formats")
            formats = ", ".join([f"`.{fmt}`" for fmt in NASSCOM_REQUIREMENTS["document_formats"]])
            st.write(formats)
        
        with col2:
            st.markdown("### 🏥 Healthcare Standards")
            for standard in NASSCOM_REQUIREMENTS["compliance_standards"]:
                st.write(f"• {standard}")
            
            st.markdown("### 📊 Quality Criteria")
            st.write("• Minimum 100 characters")
            st.write("• Clear structure and sections")
            st.write("• Healthcare context preferred")
            st.write("• Traceable requirements")
    
    # File upload widget
    uploaded_files = st.file_uploader(
        "Choose documents to upload",
        type=NASSCOM_REQUIREMENTS["document_formats"],
        accept_multiple_files=True,
        help="Upload PRDs, User Stories, API Specs, Test Plans, etc."
    )
    
    if uploaded_files:
        st.markdown("### 📝 Document Analysis Results")
        
        for uploaded_file in uploaded_files:
            with st.expander(f"📄 {uploaded_file.name}", expanded=True):
                with st.spinner(f"Analyzing {uploaded_file.name} for NASSCOM compliance..."):
                    # Process file
                    content, compliance_report = process_uploaded_file(uploaded_file)
                    
                    # Store in session state
                    st.session_state.compliance_reports[uploaded_file.name] = compliance_report
                    
                    # Display compliance report
                    display_compliance_report(compliance_report)
                    
                    # Show content preview
                    with st.expander("📖 Content Preview", expanded=False):
                        st.text_area("", value=content[:1000] + "..." if len(content) > 1000 else content, 
                                   height=200, disabled=True, label_visibility="collapsed")
                    
                    # Action buttons
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if compliance_report['is_compliant']:
                            if st.button(f"✅ Add to Knowledge Base", key=f"add_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.success(f"Added {uploaded_file.name} to knowledge base!")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    # Clear cache to reload with new documents
                                    st.cache_resource.clear()
                                    st.rerun()
                        else:
                            st.warning("Document needs improvements before adding to knowledge base")
                    
                    with col2:
                        # Download compliance report
                        report_json = json.dumps(compliance_report, indent=2)
                        st.download_button(
                            label="📥 Download Compliance Report",
                            data=report_json,
                            file_name=f"{uploaded_file.name}_compliance.json",
                            mime="application/json",
                            key=f"download_{uploaded_file.name}"
                        )
                    
                    with col3:
                        if not compliance_report['is_compliant']:
                            if st.button(f"🔄 Override & Add Anyway", key=f"override_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.warning(f"Added {uploaded_file.name} despite compliance issues")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    st.cache_resource.clear()
                                    st.rerun()

def refine_test_case_with_feedback(test_case: Dict, feedback: str, specific_field: str = None) -> Dict:
    """
    Refine an existing test case based on human feedback.
    """
    test_case_for_prompt = {k: v for k, v in test_case.items() if k != 'retrieved_context'}
    
    if specific_field:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK FOR '{specific_field}':
{feedback}

NASSCOM REQUIREMENTS:
- Maintain all required fields
- Ensure healthcare compliance where applicable
- Include both positive and negative scenarios
- Maintain traceability

Generate an improved version of this test case, specifically improving the '{specific_field}' field based on the feedback.
Maintain all other fields unless they need adjustment to be consistent with the change.

Return ONLY a valid JSON object with the same structure."""
    else:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK:
{feedback}

Generate an improved version of this test case incorporating the feedback while maintaining NASSCOM compliance.

Return ONLY a valid JSON object."""
    
    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                max_output_tokens=4096
            )
        )
        
        response = model.generate_content(prompt)
        refined_test = json.loads(response.text)
        
        # Preserve metadata
        refined_test['id'] = test_case['id']
        refined_test['refinement_timestamp'] = datetime.now().isoformat()
        refined_test['refinement_feedback'] = feedback
        refined_test['version'] = test_case.get('version', 1) + 1
        refined_test['nasscom_compliant'] = True
        
        if 'retrieved_context' in test_case:
            refined_test['retrieved_context'] = test_case['retrieved_context']
        
        return refined_test
        
    except Exception as e:
        st.error(f"Error refining test case: {e}")
        return test_case

def display_test_case(test_case: Dict, context_docs: List[Dict], key_suffix: str = ""):
    """Display generated test case in a formatted way
    
    Args:
        test_case: The test case dictionary to display
        context_docs: Context documents used for generation
        key_suffix: Suffix to add to button keys to avoid duplicates
    """
    
    # Header with metrics
    st.markdown("---")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Priority", test_case.get('priority', 'N/A'))
    with col2:
        st.metric("Category", test_case.get('category', 'N/A'))
    with col3:
        compliance = test_case.get('compliance', [])
        st.metric("Compliance", len(compliance) if compliance else 'None')
    with col4:
        st.metric("Duration", test_case.get('estimated_duration', 'N/A'))
    
    # NASSCOM compliance indicator
    if test_case.get('nasscom_compliant', False):
        st.success("✅ NASSCOM Compliant Test Case")
    else:
        st.warning("⚠️ Non-compliant test case (fallback generated)")
    
    # Main content
    st.markdown("### 📝 Description")
    st.write(test_case.get('description', 'N/A'))
    
    # Traceability
    if test_case.get('traceability'):
        st.markdown("### 🔗 Traceability")
        st.info(test_case.get('traceability', 'N/A'))
    
    # Preconditions
    st.markdown("### ⚙️ Preconditions")
    st.info(test_case.get('preconditions', 'None specified'))
    
    # Test steps
    st.markdown("### 📋 Test Steps")
    steps = test_case.get('test_steps', [])
    for i, step in enumerate(steps, 1):
        st.write(f"{i}. {step}")
    
    # Expected results
    st.markdown("### ✅ Expected Results")
    st.success(test_case.get('expected_results', 'N/A'))
    
    # Test data
    if test_case.get('test_data'):
        st.markdown("### 📊 Test Data")
        st.json(test_case.get('test_data'))
    
    # Edge cases and negative tests
    col1, col2 = st.columns(2)
    with col1:
        if test_case.get('edge_cases'):
            st.markdown("### ⚠️ Edge Cases")
            for edge in test_case['edge_cases']:
                st.write(f"• {edge}")
    
    with col2:
        if test_case.get('negative_tests'):
            st.markdown("### 🚫 Negative Tests")
            for neg_test in test_case['negative_tests']:
                st.write(f"• {neg_test}")
    
    # Metadata
    with st.expander("🔍 Generation Details"):
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Automation Feasible:** {'Yes' if test_case.get('automation_feasible') else 'No'}")
            st.write(f"**Context Relevance:** {test_case.get('avg_context_relevance', 'N/A')}")
            if test_case.get('compliance'):
                st.write(f"**Standards:** {', '.join(test_case['compliance'])}")
        with col2:
            st.write(f"**Generated From:** {test_case.get('generated_from', 'N/A')[:50]}...")
            st.write(f"**Timestamp:** {test_case.get('generation_timestamp', 'N/A')[:19]}")
            st.write(f"**Version:** {test_case.get('version', 1)}")
    
    # Save options
    st.markdown("#### 💾 Save Options")
    col1, col2 = st.columns(2)
    with col1:
        if st.button(f"Save as JSON", key=f"save_json_{test_case['id']}{key_suffix}"):
            filename = save_test_case(test_case, 'json')
            st.success(f"Saved to {filename}")
    with col2:
        if st.button(f"Save as CSV", key=f"save_csv_{test_case['id']}{key_suffix}"):
            filename = save_test_case(test_case, 'csv')
            st.success(f"Saved to {filename}")

# ===============================
# DevOps Export Functions (NEW IN V7)
# ===============================

def convert_to_jira_format(test_cases: List[Dict]) -> pd.DataFrame:
    """Convert test cases to Jira-compatible CSV format"""
    jira_data = []
    for tc in test_cases:
        jira_data.append({
            'Issue Type': 'Test',
            'Summary': tc.get('title', 'Untitled Test'),
            'Description': tc.get('description', ''),
            'Priority': tc.get('priority', 'Medium'),
            'Labels': ', '.join(tc.get('compliance', [])),
            'Test Type': 'Manual' if not tc.get('automation_feasible', False) else 'Automated',
            'Acceptance Criteria': tc.get('expected_results', ''),
            'Test Steps': '\n'.join([f"{i+1}. {step}" for i, step in enumerate(tc.get('test_steps', []))]),
            'Test Data': json.dumps(tc.get('test_data', {})),
            'Component': tc.get('category', 'Functional'),
            'Fix Version': tc.get('version', '1.0'),
            'Story Points': '3' if tc.get('priority') == 'High' else '2' if tc.get('priority') == 'Medium' else '1',
            'Epic Link': tc.get('traceability', ''),
            'Custom Field (Compliance)': ', '.join(tc.get('compliance', [])),
            'Custom Field (Duration)': tc.get('estimated_duration', ''),
        })
    return pd.DataFrame(jira_data)

def convert_to_azure_devops_format(test_cases: List[Dict]) -> pd.DataFrame:
    """Convert test cases to Azure DevOps Test Plans format"""
    azure_data = []
    for tc in test_cases:
        azure_data.append({
            'Work Item Type': 'Test Case',
            'Title': tc.get('title', 'Untitled Test'),
            'State': 'Design',
            'Priority': {'Critical': 1, 'High': 2, 'Medium': 3, 'Low': 4}.get(tc.get('priority', 'Medium'), 3),
            'Assigned To': '',
            'Area Path': f"Healthcare\\{tc.get('category', 'General')}",
            'Iteration Path': '',
            'Description': tc.get('description', ''),
            'Steps': json.dumps([{'action': step, 'expected': tc.get('expected_results', '')} 
                                for step in tc.get('test_steps', [])]),
            'Precondition': tc.get('preconditions', ''),
            'Postcondition': tc.get('expected_results', ''),
            'Tags': '; '.join(tc.get('compliance', [])),
            'Automation Status': 'Automated' if tc.get('automation_feasible', False) else 'Not Automated',
            'Test Suite': tc.get('category', 'General'),
            'Parameters': json.dumps(tc.get('test_data', {}))
        })
    return pd.DataFrame(azure_data)

def convert_to_postman_format(test_cases: List[Dict]) -> Dict:
    """Convert test cases to Postman Collection format"""
    postman_collection = {
        "info": {
            "name": "Healthcare Test Collection",
            "description": "Generated test cases for Healthcare/MedTech system",
            "schema": "https://schema.getpostman.com/json/collection/v2.1.0/collection.json"
        },
        "item": []
    }
    
    for tc in test_cases:
        # Only include API/Integration tests
        if tc.get('category') in ['Integration', 'API', 'Performance']:
            item = {
                "name": tc.get('title', 'Untitled Test'),
                "request": {
                    "method": "GET",  # Default, should be customized
                    "header": [],
                    "description": tc.get('description', ''),
                    "url": {
                        "raw": "{{base_url}}/api/endpoint",
                        "host": ["{{base_url}}"],
                        "path": ["api", "endpoint"]
                    }
                },
                "event": [
                    {
                        "listen": "test",
                        "script": {
                            "exec": [
                                f"// Test: {tc.get('title', '')}",
                                f"// Expected: {tc.get('expected_results', '')}",
                                "pm.test('Status code is 200', function () {",
                                "    pm.response.to.have.status(200);",
                                "});",
                                "",
                                "pm.test('Response time is less than 2000ms', function () {",
                                "    pm.expect(pm.response.responseTime).to.be.below(2000);",
                                "});"
                            ]
                        }
                    }
                ],
                "response": []
            }
            
            # Add test data as variables
            if tc.get('test_data'):
                item['event'].append({
                    "listen": "prerequest",
                    "script": {
                        "exec": [
                            "// Set test data",
                            *[f"pm.variables.set('{k}', '{v}');" 
                              for k, v in tc.get('test_data', {}).items()]
                        ]
                    }
                })
            
            postman_collection["item"].append(item)
    
    return postman_collection

def convert_to_junit_xml(test_cases: List[Dict]) -> str:
    """Convert test cases to JUnit XML format"""
    xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n'
    xml_content += f'<testsuites name="Healthcare Test Suite" tests="{len(test_cases)}">\n'
    xml_content += f'  <testsuite name="Automated Tests" tests="{len(test_cases)}">\n'
    
    for tc in test_cases:
        classname = tc.get('category', 'General').replace(' ', '')
        testname = tc.get('title', 'Untitled').replace(' ', '_')
        
        xml_content += f'    <testcase classname="{classname}" name="{testname}" time="{tc.get("estimated_duration", "0")}">\n'
        
        # Add test steps as system-out
        xml_content += '      <system-out><![CDATA[\n'
        xml_content += f'        Description: {tc.get("description", "")}\n'
        xml_content += f'        Priority: {tc.get("priority", "Medium")}\n'
        xml_content += '        Steps:\n'
        for i, step in enumerate(tc.get('test_steps', []), 1):
            xml_content += f'          {i}. {step}\n'
        xml_content += f'        Expected: {tc.get("expected_results", "")}\n'
        xml_content += '      ]]></system-out>\n'
        
        # Add properties for metadata
        xml_content += '      <properties>\n'
        xml_content += f'        <property name="priority" value="{tc.get("priority", "Medium")}"/>\n'
        xml_content += f'        <property name="automated" value="{tc.get("automation_feasible", False)}"/>\n'
        for compliance in tc.get('compliance', []):
            xml_content += f'        <property name="compliance" value="{compliance}"/>\n'
        xml_content += '      </properties>\n'
        
        xml_content += '    </testcase>\n'
    
    xml_content += '  </testsuite>\n'
    xml_content += '</testsuites>'
    
    return xml_content

def convert_to_testrail_format(test_cases: List[Dict]) -> List[Dict]:
    """Convert test cases to TestRail JSON format"""
    testrail_cases = []
    
    for tc in test_cases:
        testrail_case = {
            "title": tc.get('title', 'Untitled Test'),
            "type_id": 1,  # Manual test type
            "priority_id": {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1}.get(tc.get('priority', 'Medium'), 2),
            "estimate": tc.get('estimated_duration', ''),
            "refs": tc.get('traceability', ''),
            "custom_automation_type": 0 if not tc.get('automation_feasible', False) else 1,
            "custom_preconds": tc.get('preconditions', ''),
            "custom_steps_separated": [
                {
                    "content": step,
                    "expected": tc.get('expected_results', '') if i == len(tc.get('test_steps', [])) - 1 else ''
                }
                for i, step in enumerate(tc.get('test_steps', []))
            ],
            "custom_expected": tc.get('expected_results', ''),
            "custom_test_data": json.dumps(tc.get('test_data', {})),
            "custom_tags": ', '.join(tc.get('compliance', [])),
            "section_id": 1  # Default section
        }
        testrail_cases.append(testrail_case)
    
    return testrail_cases

def customize_export_with_ai(test_cases: List[Dict], target_format: str, custom_requirements: str) -> str:
    """Use AI to customize the export format based on user requirements"""
    
    # Sample of test data for AI
    sample_test = test_cases[0] if test_cases else {}
    
    prompt = f"""You are an expert in test management tools and DevOps integration.
    
I need to export test cases to {target_format} format with these custom requirements:
{custom_requirements}

Sample test case structure:
{json.dumps(sample_test, indent=2)[:1000]}

Current standard fields I'm using:
- Jira: Issue Type, Summary, Description, Priority, Labels, Test Steps, etc.
- Azure DevOps: Work Item Type, Title, State, Priority, Steps, Tags, etc.
- Postman: Collection with requests, tests, variables
- JUnit: XML with testsuites, testcases, properties

Based on the custom requirements, provide:
1. Modified field mappings
2. Additional fields to include
3. Field transformations needed
4. Any format-specific considerations

Return as JSON with structure:
{{
    "field_mappings": {{
        "our_field": "target_field",
        ...
    }},
    "additional_fields": [
        {{"name": "field_name", "value": "how to derive it"}}
    ],
    "transformations": [
        {{"field": "field_name", "transformation": "description"}}
    ],
    "format_notes": "specific considerations for this format"
}}
"""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        customization = json.loads(response.text)
        return customization
        
    except Exception as e:
        st.error(f"AI customization error: {e}")
        return {
            "field_mappings": {},
            "additional_fields": [],
            "transformations": [],
            "format_notes": "Using default format"
        }

def display_export_interface():
    """Display the DevOps export interface with preview and customization"""
    st.header("🔌 DevOps Integration Export")
    st.markdown("""
    Export your test cases to various DevOps tools and test management platforms.
    Preview the format and customize field mappings using AI.
    """)
    
    if not st.session_state.generated_tests:
        st.warning("No test cases available to export. Generate or import tests first.")
        return
    
    # Export format selection
    col1, col2 = st.columns([2, 1])
    
    with col1:
        export_format = st.selectbox(
            "Select Export Format",
            ["Jira CSV", "Azure DevOps", "Postman Collection", "JUnit XML", "TestRail JSON", "Custom Format"],
            help="Choose the target platform for export"
        )
    
    with col2:
        tests_to_export = st.selectbox(
            "Tests to Export",
            ["All Tests", "Filtered Tests", "Selected Tests"],
            help="Choose which tests to export"
        )
    
    # Filter options
    if tests_to_export == "Filtered Tests":
        col1, col2, col3 = st.columns(3)
        with col1:
            filter_category = st.multiselect(
                "Categories",
                list(set(tc.get('category', 'Unknown') for tc in st.session_state.generated_tests))
            )
        with col2:
            filter_priority = st.multiselect(
                "Priorities",
                list(set(tc.get('priority', 'Unknown') for tc in st.session_state.generated_tests))
            )
        with col3:
            filter_automated = st.checkbox("Only Automated Tests")
        
        # Apply filters
        export_tests = st.session_state.generated_tests
        if filter_category:
            export_tests = [tc for tc in export_tests if tc.get('category') in filter_category]
        if filter_priority:
            export_tests = [tc for tc in export_tests if tc.get('priority') in filter_priority]
        if filter_automated:
            export_tests = [tc for tc in export_tests if tc.get('automation_feasible', False)]
    else:
        export_tests = st.session_state.generated_tests
    
    st.info(f"📊 {len(export_tests)} test cases will be exported")
    
    # Customization section
    with st.expander("⚙️ Customize Export Format", expanded=False):
        st.markdown("### Field Mapping Customization")
        
        custom_requirements = st.text_area(
            "Describe your customization needs",
            placeholder="Example: Add custom fields for sprint number, test environment, risk level. Map our 'compliance' field to Jira's 'Regulatory' field.",
            height=100
        )
        
        if st.button("🤖 Get AI Suggestions", key="ai_customize"):
            if custom_requirements:
                with st.spinner("AI is analyzing your requirements..."):
                    customization = customize_export_with_ai(export_tests, export_format, custom_requirements)
                    st.session_state['export_customization'] = customization
                    
                    st.success("AI suggestions ready!")
                    st.json(customization)
    
    # Preview section
    with st.expander("👁️ Preview Export Format", expanded=True):
        st.markdown(f"### Preview: {export_format}")
        
        # Generate preview based on format
        preview_data = None
        
        if export_format == "Jira CSV":
            preview_data = convert_to_jira_format(export_tests[:3])  # Preview first 3
            st.dataframe(preview_data, use_container_width=True)
            
        elif export_format == "Azure DevOps":
            preview_data = convert_to_azure_devops_format(export_tests[:3])
            st.dataframe(preview_data, use_container_width=True)
            
        elif export_format == "Postman Collection":
            preview_data = convert_to_postman_format(export_tests[:3])
            st.json(preview_data)
            
        elif export_format == "JUnit XML":
            preview_data = convert_to_junit_xml(export_tests[:3])
            st.code(preview_data[:1500], language='xml')  # Show first 1500 chars
            
        elif export_format == "TestRail JSON":
            preview_data = convert_to_testrail_format(export_tests[:3])
            st.json(preview_data[0] if preview_data else {})  # Show first test
        
        if len(export_tests) > 3:
            st.info(f"... and {len(export_tests) - 3} more tests")
    
    # Export buttons
    st.markdown("### 📥 Export Options")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Download Export", type="primary", key="download_export"):
            with st.spinner(f"Generating {export_format} export..."):
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                
                if export_format == "Jira CSV":
                    df = convert_to_jira_format(export_tests)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "Download Jira CSV",
                        csv,
                        f"jira_tests_{timestamp}.csv",
                        "text/csv"
                    )
                    
                elif export_format == "Azure DevOps":
                    df = convert_to_azure_devops_format(export_tests)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "Download Azure DevOps CSV",
                        csv,
                        f"azure_tests_{timestamp}.csv",
                        "text/csv"
                    )
                    
                elif export_format == "Postman Collection":
                    collection = convert_to_postman_format(export_tests)
                    json_str = json.dumps(collection, indent=2)
                    st.download_button(
                        "Download Postman Collection",
                        json_str,
                        f"postman_collection_{timestamp}.json",
                        "application/json"
                    )
                    
                elif export_format == "JUnit XML":
                    xml_content = convert_to_junit_xml(export_tests)
                    st.download_button(
                        "Download JUnit XML",
                        xml_content,
                        f"junit_tests_{timestamp}.xml",
                        "text/xml"
                    )
                    
                elif export_format == "TestRail JSON":
                    testrail_data = convert_to_testrail_format(export_tests)
                    json_str = json.dumps(testrail_data, indent=2)
                    st.download_button(
                        "Download TestRail JSON",
                        json_str,
                        f"testrail_tests_{timestamp}.json",
                        "application/json"
                    )
    
    with col2:
        if st.button("📋 Copy to Clipboard", key="copy_export"):
            st.info("Export copied! (Feature requires JavaScript integration)")
    
    with col3:
        if st.button("🔗 Send to Integration", key="send_integration"):
            st.info(f"Ready to integrate with {export_format.split()[0]} API")
    
    # Integration guide
    with st.expander("📚 Integration Guide", expanded=False):
        st.markdown(f"""
        ### How to import into {export_format.split()[0]}
        
        **Jira:**
        1. Go to Issues > Import Issues from CSV
        2. Upload the downloaded CSV file
        3. Map fields as needed
        4. Complete the import
        
        **Azure DevOps:**
        1. Open Test Plans
        2. Click Import Test Cases
        3. Upload the CSV file
        4. Map to your test plan
        
        **Postman:**
        1. Open Postman
        2. Click Import
        3. Select the JSON file
        4. Tests will appear in Collections
        
        **JUnit:**
        1. Place XML in your test results directory
        2. Configure CI/CD to read JUnit format
        3. Results will appear in test reports
        
        **TestRail:**
        1. Use TestRail API or Import tool
        2. Upload JSON via API endpoint
        3. Map to test suite
        """)

def save_test_case(test_case: Dict, format: str = 'json'):
    """Save test case to file"""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if format == 'json':
        filename = f"{TEST_CASES_DIR}/manual_save_{test_case['id']}_{timestamp}.json"
        with open(filename, 'w') as f:
            json.dump(test_case, f, indent=2, default=str)
        return filename
    elif format == 'csv':
        filename = f"{TEST_CASES_DIR}/manual_save_{test_case['id']}_{timestamp}.csv"
        df = pd.DataFrame([test_case])
        df.to_csv(filename, index=False)
        return filename
    return None

# MAIN UI
def main():
    # Apply modern CSS
    inject_modern_css()
    
    # Clean, professional header
    st.markdown("""
    <div style='text-align: center; padding: 2.5rem 0; animation: fadeIn 0.8s ease;'>
        <h1 style='font-size: 3rem; 
                   color: #F1F5F9;
                   font-weight: 700;
                   margin-bottom: 0.5rem;
                   letter-spacing: -0.03em;'>
            🧪 AI Test Generator Pro
        </h1>
        <p style='font-size: 1.1rem; color: #CBD5E1; margin-top: 0.5rem; font-weight: 400;'>
            NASSCOM Healthcare/MedTech Testing Solution • Version 9.0
        </p>
        <div style='margin-top: 1.5rem; display: flex; justify-content: center; gap: 0.75rem; flex-wrap: wrap;'>
            <span style='background: #1E293B; 
                         border: 1px solid #475569;
                         padding: 0.375rem 0.875rem; 
                         border-radius: 6px; 
                         color: #CBD5E1; 
                         font-size: 0.875rem;
                         font-weight: 500;
                         transition: all 0.2s ease;'>
                ✨ Clean Architecture
            </span>
            <span style='background: #1E293B; 
                         border: 1px solid #475569;
                         padding: 0.375rem 0.875rem; 
                         border-radius: 6px; 
                         color: #CBD5E1; 
                         font-size: 0.875rem;
                         font-weight: 500;
                         transition: all 0.2s ease;'>
                🚀 Enhanced Performance
            </span>
            <span style='background: #1E293B; 
                         border: 1px solid #475569;
                         padding: 0.375rem 0.875rem; 
                         border-radius: 6px; 
                         color: #CBD5E1; 
                         font-size: 0.875rem;
                         font-weight: 500;
                         transition: all 0.2s ease;'>
                ⚡ AI-Powered
            </span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Load saved tests on first run (unified from single directory)
    if 'tests_loaded' not in st.session_state:
        all_saved_tests = initialize_saved_tests()
        if all_saved_tests:
            st.session_state.generated_tests = all_saved_tests
            st.sidebar.success(f"📂 Loaded {len(all_saved_tests)} test cases from unified storage")
        st.session_state.tests_loaded = True
    
    # Sidebar for system info
    with st.sidebar:
        # Animated Logo Header
        st.markdown("""
        <div style='text-align: center; 
                    padding: 1.5rem 0; 
                    margin-bottom: 1rem;
                    background: linear-gradient(135deg, #8B5CF6 0%, #EC4899 100%);
                    border-radius: 12px;
                    box-shadow: 0 8px 32px rgba(139, 92, 246, 0.3);
                    animation: pulse 3s ease infinite;'>
            <div style='font-size: 2.5rem; margin-bottom: 0.5rem; animation: float 3s ease-in-out infinite;'>
                🧪
            </div>
            <h2 style='color: white; 
                       margin: 0; 
                       font-size: 1.2rem; 
                       font-weight: 700; 
                       letter-spacing: 0.1em;
                       text-transform: uppercase;'>
                Test Generator
            </h2>
            <p style='color: rgba(255, 255, 255, 0.9); 
                      margin: 0.25rem 0 0 0; 
                      font-size: 0.8rem;
                      font-weight: 400;'>
                Healthcare Edition v9.0
            </p>
        </div>
        
        <style>
            @keyframes float {
                0%, 100% { transform: translateY(0px); }
                50% { transform: translateY(-10px); }
            }
        </style>
        """, unsafe_allow_html=True)
        
        # Quick Actions
        st.markdown("""
        <div style='background: linear-gradient(135deg, #141824 0%, #262B3B 100%);
                    border-radius: 10px;
                    padding: 0.75rem;
                    margin-bottom: 1rem;
                    border: 1px solid rgba(139, 92, 246, 0.2);'>
            <h3 style='color: #A78BFA; 
                       font-size: 0.9rem; 
                       margin: 0 0 0.5rem 0;
                       text-transform: uppercase;
                       letter-spacing: 0.1em;'>
                ⚡ Quick Actions
            </h3>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Refresh", use_container_width=True):
                st.rerun()
        with col2:
            if st.button("💾 Save All", use_container_width=True):
                auto_save_test_cases(st.session_state.get('generated_tests', []), 
                                   'generated', "Manual save triggered")
        
        # System Status Dashboard - Start container
        with st.container():
            st.markdown("""
            <div style='background: linear-gradient(135deg, #141824 0%, #262B3B 100%);
                        border-radius: 10px;
                        padding: 1rem;
                        margin: 1rem 0;
                        border: 1px solid rgba(139, 92, 246, 0.2);'>
                <h3 style='color: #14B8A6; 
                           font-size: 0.9rem; 
                           margin: 0 0 0.75rem 0;
                           text-transform: uppercase;
                           letter-spacing: 0.1em;'>
                    📊 System Status
                </h3>
            </div>
            """, unsafe_allow_html=True)
        
        # Load RAG system
        embedding_model, index, doc_chunks, doc_metadata = load_rag_system()
        
        if index:
            # Status indicator
            st.success("🟢 **RAG System Active**")
            
            # Count documents
            existing_docs = len(set(m.get('source') for m in doc_metadata if not m.get('is_uploaded', False)))
            uploaded_docs = len(set(m.get('source') for m in doc_metadata if m.get('is_uploaded', False)))
            
            # Display metrics using Streamlit columns
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📚 Documents", len(doc_chunks), "chunks indexed")
            with col2:
                st.metric("📁 Existing", existing_docs, "pre-loaded docs")
            with col3:
                st.metric("📤 Uploaded", uploaded_docs, "user documents")
            
            # Compliance summary
            if st.session_state.compliance_reports:
                compliant_count = sum(1 for r in st.session_state.compliance_reports.values() if r['is_compliant'])
                total_count = len(st.session_state.compliance_reports)
                compliance_percentage = (compliant_count / total_count) * 100
                
                st.metric("📋 Compliance", f"{compliant_count}/{total_count}", f"{compliance_percentage:.0f}% compliant")
                st.progress(compliance_percentage / 100)
        else:
            st.warning("⚠️ **System Initializing...** Please upload documents to activate the RAG system.")
        
        # Test Statistics
        st.markdown("""
        <div style='background: linear-gradient(135deg, #141824 0%, #262B3B 100%);
                    border-radius: 10px;
                    padding: 1rem;
                    margin: 1rem 0;
                    border: 1px solid rgba(139, 92, 246, 0.2);'>
            <h3 style='color: #F97316; 
                       font-size: 0.9rem; 
                       margin: 0 0 0.75rem 0;
                       text-transform: uppercase;
                       letter-spacing: 0.1em;'>
                📈 Session Statistics
            </h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Test counter
        test_counter_html = f"""
        <div style='display: flex; 
                    justify-content: space-between; 
                    align-items: center;
                    padding: 0.5rem;
                    background: rgba(249, 115, 22, 0.1);
                    border-radius: 8px;
                    margin-bottom: 0.75rem;'>
            <span style='color: #D1D5DB; font-size: 0.85rem;'>Tests Generated</span>
            <span style='color: #F97316; font-size: 1.5rem; font-weight: 700;'>{st.session_state.test_counter}</span>
        </div>
        """
        st.markdown(test_counter_html, unsafe_allow_html=True)
        
        if st.session_state.generated_tests:
            test_count = len(st.session_state.generated_tests)
            priority_counts = {}
            for test in st.session_state.generated_tests:
                priority = test.get('priority', 'Medium')
                priority_counts[priority] = priority_counts.get(priority, 0) + 1
            
            # Priority breakdown
            priority_colors = {'High': '#F43F5E', 'Medium': '#EAB308', 'Low': '#10B981'}
            for priority, count in priority_counts.items():
                color = priority_colors.get(priority, '#9CA3AF')
                percentage = (count / test_count) * 100
                priority_html = f"""
                <div style='margin-bottom: 0.5rem;'>
                    <div style='display: flex; justify-content: space-between; margin-bottom: 0.25rem;'>
                        <span style='color: {color}; font-size: 0.8rem;'>{priority} Priority</span>
                        <span style='color: {color}; font-size: 0.8rem; font-weight: 600;'>{count}</span>
                    </div>
                    <div style='width: 100%; height: 3px; background: #1F2937; border-radius: 2px; overflow: hidden;'>
                        <div style='width: {percentage}%; 
                                    height: 100%; 
                                    background: {color};
                                    transition: width 0.5s ease;'></div>
                    </div>
                </div>
                """
                st.markdown(priority_html, unsafe_allow_html=True)
        
        # Auto-save status with better visualization
        st.markdown("""
        <div style='background: linear-gradient(135deg, #141824 0%, #262B3B 100%);
                    border-radius: 10px;
                    padding: 1rem;
                    margin: 1rem 0;
                    border: 1px solid rgba(139, 92, 246, 0.2);'>
            <h3 style='color: #10B981; 
                       font-size: 0.9rem; 
                       margin: 0 0 0.75rem 0;
                       text-transform: uppercase;
                       letter-spacing: 0.1em;'>
                💾 Auto-Save
            </h3>
        </div>
        """, unsafe_allow_html=True)
        
        if os.path.exists(f"{TEST_CASES_DIR}/all_tests_latest.json"):
            try:
                with open(f"{TEST_CASES_DIR}/all_tests_latest.json", 'r') as f:
                    saved_tests = json.load(f)
                    file_time = datetime.fromtimestamp(os.path.getmtime(f"{TEST_CASES_DIR}/all_tests_latest.json"))
                    
                    save_status_html = f"""
                    <div style='padding: 0.5rem;
                                background: rgba(16, 185, 129, 0.1);
                                border-radius: 8px;'>
                        <div style='display: flex; align-items: center; margin-bottom: 0.5rem;'>
                            <div style='width: 8px; height: 8px; 
                                        background: #10B981; 
                                        border-radius: 50%; 
                                        margin-right: 0.5rem;
                                        animation: pulse 2s infinite;'></div>
                            <span style='color: #10B981; font-weight: 600; font-size: 0.85rem;'>
                                {len(saved_tests)} tests saved
                            </span>
                        </div>
                        <div style='color: #9CA3AF; font-size: 0.75rem;'>
                            Last: {file_time.strftime('%H:%M:%S')}
                        </div>
                    </div>
                    """
                    st.markdown(save_status_html, unsafe_allow_html=True)
            except:
                st.info("No saved tests found")
        else:
            st.info("No tests saved yet")
        
        # Export section
        if st.session_state.generated_tests:
            with st.expander("📤 Export Options", expanded=False):
                if st.button("🔽 Export to JSON", use_container_width=True):
                    all_tests = {
                        "generated_on": datetime.now().isoformat(),
                        "total_tests": len(st.session_state.generated_tests),
                        "nasscom_compliant": True,
                        "test_cases": st.session_state.generated_tests
                    }
                    filename = f"{TEST_CASES_DIR}/export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                    with open(filename, 'w') as f:
                        json.dump(all_tests, f, indent=2, default=str)
                    st.success(f"Saved to {filename}")
        
        # Settings & Configuration
        with st.expander("⚙️ Configuration", expanded=False):
            st.markdown("""
            <div style='padding: 0.5rem;'>
                <h4 style='color: #A78BFA; font-size: 0.9rem; margin-bottom: 0.75rem;'>Model Settings</h4>
            </div>
            """, unsafe_allow_html=True)
            
            temperature = st.slider("Temperature", 0.0, 1.0, 0.7, 0.1, key="temp_slider")
            max_tests = st.slider("Max Tests", 1, 20, 5, key="max_tests_slider")
            
            st.markdown("""
            <div style='padding: 0.5rem; margin-top: 1rem;'>
                <h4 style='color: #14B8A6; font-size: 0.9rem; margin-bottom: 0.75rem;'>Export Settings</h4>
            </div>
            """, unsafe_allow_html=True)
            
            include_timestamps = st.checkbox("Include Timestamps", value=True, key="timestamps_check")
            include_metadata = st.checkbox("Include Metadata", value=True, key="metadata_check")
        
        # Footer with version info
        st.markdown("""
        <div style='position: relative;
                    margin-top: 2rem;
                    padding: 1rem;
                    background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, transparent 100%);
                    border-radius: 10px;
                    text-align: center;
                    border: 1px solid rgba(139, 92, 246, 0.1);'>
            <p style='color: #9CA3AF; font-size: 0.75rem; margin: 0;'>
                Built for NASSCOM Hackathon
            </p>
            <p style='color: #6366F1; font-size: 0.7rem; margin: 0.25rem 0 0 0;'>
                Aurora Theme • Version 9.0
            </p>
            <div style='margin-top: 0.75rem; 
                        display: flex; 
                        justify-content: center; 
                        gap: 1rem;'>
                <span style='color: #8B5CF6; font-size: 0.7rem;'>🚀 AI-Powered</span>
                <span style='color: #14B8A6; font-size: 0.7rem;'>✨ Healthcare Compliant</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # Main content area with tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "🧪 Generate Tests", 
        "📤 Upload Documents",
        "📥 Import Tests",
        "🔌 Export Tests",  # NEW TAB IN V7
        "🗂️ Test Suite", 
        "📚 Document Library"
    ])

    with tab1:
        st.header("🎯 Generate New Test Case")
        st.markdown("Describe the feature or requirement you want to test.")

        # Predefined examples
        examples = [
            "Test user registration with HIPAA-compliant data handling",
            "Verify patient data encryption meets FDA and ISO 27001 standards",
            "Test API authentication with OAuth 2.0 for medical records access",
            "Validate HL7 message processing for lab results",
            "Test DICOM image upload with PHI anonymization",
            "Verify audit trail generation for GDPR compliance"
        ]
        
        selected_example = st.selectbox("Choose an example or write your own:", 
                                       ["Custom requirement..."] + examples)
        
        if selected_example == "Custom requirement...":
            requirement = st.text_area(
                "Enter your test requirement:",
                height=100,
                placeholder="Describe what you want to test..."
            )
        else:
            requirement = selected_example
            st.info(f"Using example: {requirement}")
        
        # Generation button
        if st.button("🚀 Generate Test Case", type="primary"):
            if not requirement or requirement == "Custom requirement...":
                st.warning("Please enter a requirement")
            elif not index:
                st.error("No documents in knowledge base. Please upload documents first.")
            else:
                # Show AI thinking animation
                show_ai_thinking(duration=2)
                
                with st.spinner("🔍 Analyzing your requirement..."):
                    # Retrieve context
                    context_docs = retrieve_context(
                        requirement, 
                        embedding_model, 
                        index, 
                        doc_chunks, 
                        doc_metadata,
                        k=5
                    )
                    
                # More AI thinking for generation
                show_ai_thinking(duration=3)
                
                with st.spinner("✨ Generating NASSCOM-compliant test case..."):
                    # Generate test case
                    test_case = generate_test_case_with_gemini(requirement, context_docs)
                    
                    # Add context to test case object
                    test_case['retrieved_context'] = context_docs
                    
                    # Store in session
                    st.session_state.generated_tests.insert(0, test_case)
                    
                    # AUTO-SAVE all test cases to unified location
                    auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                    # Clean up old saves to prevent disk bloat
                    cleanup_old_saves("all_tests", keep_count=10)
                    
                    # Display results
                    if test_case.get('nasscom_compliant', False):
                        st.success(f"✅ NASSCOM-Compliant Test Case Generated & Auto-Saved: {test_case['id']}")
                    else:
                        st.warning(f"⚠️ Test case generated with limited compliance & auto-saved: {test_case['id']}")
                    
                    st.balloons()
                    
                    # Show the generated test case
                    display_test_case(test_case, context_docs, "_gen")

    with tab2:
        display_upload_interface()

    with tab3:
        display_import_interface()  # NEW IMPORT INTERFACE

    with tab4:
        display_export_interface()  # NEW EXPORT INTERFACE IN V7

    with tab5:
        st.header("🗂️ Generated Test Suite")
        
        if not st.session_state.generated_tests:
            st.info("No test cases generated yet. Go to the 'Generate Tests' tab to start.")
        else:
            st.markdown(f"You have generated **{len(st.session_state.generated_tests)}** test case(s).")
            
            # Filter options
            col1, col2, col3 = st.columns(3)
            with col1:
                filter_category = st.selectbox("Filter by Category", 
                                              ["All"] + list(set(tc.get('category', 'Unknown') for tc in st.session_state.generated_tests)))
            with col2:
                filter_priority = st.selectbox("Filter by Priority",
                                              ["All"] + list(set(tc.get('priority', 'Unknown') for tc in st.session_state.generated_tests)))
            with col3:
                filter_compliance = st.selectbox("Filter by Compliance",
                                                ["All", "NASSCOM Compliant", "Non-Compliant"])
            
            # Apply filters
            filtered_tests = st.session_state.generated_tests
            if filter_category != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('category') == filter_category]
            if filter_priority != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('priority') == filter_priority]
            if filter_compliance == "NASSCOM Compliant":
                filtered_tests = [tc for tc in filtered_tests if tc.get('nasscom_compliant', False)]
            elif filter_compliance == "Non-Compliant":
                filtered_tests = [tc for tc in filtered_tests if not tc.get('nasscom_compliant', False)]
            
            st.markdown(f"Showing **{len(filtered_tests)}** test case(s)")
            
            for i, tc in enumerate(filtered_tests):
                with st.expander(f"**{tc['id']}**: {tc.get('title', 'Untitled')} (v{tc.get('version', 1)})", expanded=(i==0)):
                    display_test_case(tc, tc.get('retrieved_context', []), f"_suite_{i}")

    with tab6:
        st.header("📚 Document Library")
        
        # Show both existing and uploaded documents
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📁 Pre-loaded Documents")
            existing_files = glob.glob(f"{DOCUMENTS_PATH}/*")
            if existing_files:
                for file_path in existing_files:
                    file_name = Path(file_path).name
                    file_size = os.path.getsize(file_path)
                    st.write(f"• {file_name} ({file_size:,} bytes)")
            else:
                st.info("No pre-loaded documents")
        
        with col2:
            st.subheader("📤 Uploaded Documents")
            uploaded_files = glob.glob(f"{UPLOADED_DOCS_PATH}/*")
            uploaded_files = [f for f in uploaded_files if not f.endswith('.compliance.json')]
            
            if uploaded_files:
                for file_path in uploaded_files:
                    file_name = Path(file_path).name
                    file_size = os.path.getsize(file_path)
                    
                    # Check if we have compliance report
                    compliance_file = f"{file_path}.compliance.json"
                    if os.path.exists(compliance_file):
                        with open(compliance_file, 'r') as f:
                            report = json.load(f)
                        compliance_icon = "✅" if report['is_compliant'] else "⚠️"
                        compliance_score = report['compliance_score']
                        st.write(f"{compliance_icon} {file_name} ({compliance_score:.0f}% compliant)")
                    else:
                        st.write(f"• {file_name} ({file_size:,} bytes)")
            else:
                st.info("No uploaded documents yet")
        
        # Statistics
        st.markdown("---")
        st.subheader("📊 Knowledge Base Statistics")
        
        total_files = len(existing_files) + len(uploaded_files)
        total_size = sum(os.path.getsize(f) for f in existing_files + uploaded_files if not f.endswith('.compliance.json'))
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Documents", total_files)
        with col2:
            st.metric("Total Size", f"{total_size:,} bytes")
        with col3:
            if st.session_state.compliance_reports:
                avg_compliance = sum(r['compliance_score'] for r in st.session_state.compliance_reports.values()) / len(st.session_state.compliance_reports)
                st.metric("Avg Compliance Score", f"{avg_compliance:.1f}%")


if __name__ == "__main__":
    main()
