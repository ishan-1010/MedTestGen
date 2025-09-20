"""
AI-Powered Test Case Generation Platform - Version 14.0
Enterprise Edition with Full Test Case Editing & AI Refinement
Healthcare & MedTech Compliance Testing Solution
Enhanced with Manual Editing, AI-Powered Refinement, and Version Control
"""

import streamlit as st
import os
import json
import uuid
import numpy as np
import faiss
import yaml
import glob
import re
import io
import tempfile
import shutil
import time
import random
import requests
from pathlib import Path
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from dotenv import load_dotenv
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import pandas as pd
import PyPDF2
from docx import Document
import xml.etree.ElementTree as ET
import chardet
import logging
import bcrypt
import secrets
from cryptography.fernet import Fernet

# Import MongoDB manager and API test executor
from database import MongoDBManager
from api_test_executor import APITestExecutor, MockHealthcareAPI

# Configure comprehensive logging for CLI
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()  # Output to console/CLI
    ]
)
logger = logging.getLogger(__name__)

# Log startup
logger.info("="*60)
logger.info("[APP_START] Test Case Generation Platform v14.0 Starting...")
logger.info("[SECURITY] Authentication, encryption, and user data isolation enabled")
logger.info("[FEATURES] Full test editing, AI refinement, version control, bulk operations")
logger.info("="*60)

# Load environment variables
load_dotenv()

# ==================================
# 🎨 PAGE CONFIGURATION (MUST BE FIRST)
# ==================================

st.set_page_config(
    page_title="Test Case Generation Platform v14",
    page_icon="⚡",
    layout="wide",
    menu_items={
        'Get Help': 'https://nasscom.in',
        'About': "Enterprise Test Case Generation Platform v14.0 | Full Editing & AI Refinement | Healthcare Compliance"
    }
)

# ==================================
# 🔐 AUTHENTICATION FUNCTIONS
# ==================================

def hash_password(password: str) -> str:
    """Hash a password using bcrypt"""
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    """Verify a password against its hash"""
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def validate_email(email: str) -> bool:
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_password(password: str) -> Tuple[bool, str]:
    """
    Validate password strength
    Returns: (is_valid, error_message)
    """
    if len(password) < 8:
        return False, "Password must be at least 8 characters long"
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain at least one uppercase letter"
    if not re.search(r"[a-z]", password):
        return False, "Password must contain at least one lowercase letter"
    if not re.search(r"[0-9]", password):
        return False, "Password must contain at least one number"
    return True, ""

def get_or_create_encryption_key() -> bytes:
    """Get or create encryption key for sensitive data"""
    key = os.getenv('ENCRYPTION_KEY')
    if not key:
        key = Fernet.generate_key()
        logger.warning("[SECURITY] No encryption key found in environment, generating new one")
    elif isinstance(key, str):
        key = key.encode()
    return key

def encrypt_sensitive_data(data: str) -> str:
    """Encrypt sensitive data for secure storage in MongoDB"""
    encryption_key = get_or_create_encryption_key()
    f = Fernet(encryption_key)
    encrypted = f.encrypt(data.encode()).decode()
    logger.debug(f"[ENCRYPTION] Data encrypted, length: {len(encrypted)}")
    return encrypted

def decrypt_sensitive_data(encrypted_data: str) -> str:
    """Decrypt sensitive data from MongoDB"""
    encryption_key = get_or_create_encryption_key()
    f = Fernet(encryption_key)
    decrypted = f.decrypt(encrypted_data.encode()).decode()
    logger.debug(f"[DECRYPTION] Data decrypted successfully")
    return decrypted

def check_authentication() -> bool:
    """Check if user is authenticated"""
    return 'authenticated' in st.session_state and st.session_state.authenticated

def require_authentication():
    """Decorator to require authentication for certain pages"""
    if not check_authentication():
        st.warning("Please login to access this feature")
        st.stop()

def logout():
    """Logout the current user"""
    user_id = st.session_state.get('user_id', 'unknown')
    user_email = st.session_state.get('user_email', 'unknown')
    
    logger.info(f"[LOGOUT] User logging out - ID: {user_id}, Email: {user_email}")
    
    # Clear all user-related session state
    keys_to_clear = ['authenticated', 'user_id', 'user_email', 'user_name', 'user_role', 
                      'session_id', 'generated_tests', 'last_generated_test', 
                      'last_generated_context', 'last_requirement', 'show_settings',
                      'tests_loaded_for_user', 'tests_loaded_user_id', 'tests_loaded',
                      'shown_load_success']
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    logger.info(f"[LOGOUT_SUCCESS] User {user_id} session cleared")
    st.success("Logged out successfully")
    st.rerun()

# ==================================
# 🎨 MODERN UI STYLING AND ANIMATIONS
# ==================================

def inject_modern_css():
    """Inject clean, professional CSS with automatic light/dark mode support"""
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
        /* Global smooth transitions and performance */
        * {
            scroll-behavior: smooth;
        }
        
        *:not(.no-transition) {
            transition-timing-function: cubic-bezier(0.25, 0.46, 0.45, 0.94);
        }
        
        /* Hardware acceleration for smooth animations */
        .stApp * {
            -webkit-font-smoothing: antialiased;
            -moz-osx-font-smoothing: grayscale;
            backface-visibility: hidden;
            perspective: 1000px;
        }
        
    /* =====================================================
       LIGHT MODE COLOR SCHEME (Default)
    ===================================================== */
    :root {
        /* Primary Brand Colors */
        --primary: #7C3AED;           /* Deep Purple */
        --primary-light: #8B5CF6;     /* Medium Purple */
        --primary-dark: #6D28D9;      /* Darker Purple */
        --primary-subtle: #EDE9FE;    /* Very Light Purple for backgrounds */
        
        /* Secondary Colors */
        --secondary: #059669;          /* Deep Teal */
        --secondary-light: #10B981;   /* Light Teal */
        --secondary-subtle: #D1FAE5;  /* Very Light Teal */
        
        /* Accent Colors */
        --accent: #EA580C;             /* Deep Orange */
        --accent-light: #F97316;      /* Light Orange */
        --accent-subtle: #FED7AA;     /* Very Light Orange */
        
        /* Status Colors */
        --success: #059669;            /* Green */
        --success-light: #10B981;     /* Light Green */
        --danger: #DC2626;             /* Red */
        --danger-light: #EF4444;      /* Light Red */
        --warning: #D97706;            /* Amber */
        --warning-light: #F59E0B;     /* Light Amber */
        --info: #2563EB;               /* Blue */
        --info-light: #3B82F6;        /* Light Blue */
        
        /* Light Mode Backgrounds */
        --bg-primary: #FFFFFF;         /* Pure White */
        --bg-secondary: #F9FAFB;      /* Off White */
        --bg-tertiary: #F3F4F6;       /* Light Gray */
        --bg-surface: #FFFFFF;        /* Card Background */
        --bg-elevated: #FFFFFF;       /* Elevated Surface */
        --bg-overlay: rgba(0, 0, 0, 0.05); /* Subtle Overlay */
        
        /* Light Mode Text Colors */
        --text-primary: #111827;       /* Almost Black */
        --text-secondary: #4B5563;     /* Dark Gray */
        --text-tertiary: #6B7280;      /* Medium Gray */
        --text-muted: #9CA3AF;         /* Light Gray */
        --text-on-primary: #FFFFFF;   /* White text on primary */
        
        /* Light Mode Borders */
        --border-primary: #E5E7EB;     /* Light Gray */
        --border-secondary: #D1D5DB;   /* Medium Gray */
        --border-focus: #7C3AED;       /* Purple Focus */
        --border-subtle: #F3F4F6;      /* Very Light Gray */
        
        /* Light Mode Shadows */
        --shadow-sm: 0 1px 2px rgba(0, 0, 0, 0.05);
        --shadow-md: 0 4px 6px rgba(0, 0, 0, 0.07);
        --shadow-lg: 0 10px 15px rgba(0, 0, 0, 0.1);
        --shadow-xl: 0 20px 25px rgba(0, 0, 0, 0.1);
        --shadow-glow: 0 0 20px rgba(124, 58, 237, 0.15);
        
        /* Light Mode Gradients */
        --gradient-primary: linear-gradient(135deg, #7C3AED 0%, #A78BFA 100%);
        --gradient-secondary: linear-gradient(135deg, #059669 0%, #10B981 100%);
        --gradient-accent: linear-gradient(135deg, #EA580C 0%, #F97316 100%);
        --gradient-surface: linear-gradient(135deg, #FFFFFF 0%, #F9FAFB 100%);
        --gradient-subtle: linear-gradient(135deg, #F9FAFB 0%, #F3F4F6 100%);
    }
    
    /* =====================================================
       DARK MODE COLOR SCHEME
    ===================================================== */
    @media (prefers-color-scheme: dark) {
        :root {
            /* Primary Brand Colors - More Vibrant in Dark Mode */
            --primary: #8B5CF6;           /* Vibrant Purple */
            --primary-light: #A78BFA;     /* Light Purple */
            --primary-dark: #7C3AED;      /* Deep Purple */
            --primary-subtle: rgba(139, 92, 246, 0.1); /* Subtle Purple */
            
            /* Secondary Colors */
            --secondary: #14B8A6;          /* Vibrant Teal */
            --secondary-light: #5EEAD4;   /* Light Teal */
            --secondary-subtle: rgba(20, 184, 166, 0.1); /* Subtle Teal */
            
            /* Accent Colors */
            --accent: #F97316;             /* Vibrant Orange */
            --accent-light: #FB923C;      /* Light Orange */
            --accent-subtle: rgba(249, 115, 22, 0.1); /* Subtle Orange */
            
            /* Status Colors - Adjusted for Dark Mode */
            --success: #10B981;            /* Emerald Green */
            --success-light: #34D399;     /* Light Green */
            --danger: #F43F5E;             /* Rose Red */
            --danger-light: #FB7185;      /* Light Red */
            --warning: #EAB308;            /* Golden Yellow */
            --warning-light: #FCD34D;     /* Light Yellow */
            --info: #3B82F6;               /* Sky Blue */
            --info-light: #60A5FA;        /* Light Blue */
            
            /* Dark Mode Backgrounds */
            --bg-primary: #0A0E1A;         /* Rich Dark Navy */
            --bg-secondary: #111827;       /* Dark Gray */
            --bg-tertiary: #1F2937;        /* Medium Dark */
            --bg-surface: #141824;         /* Card Background */
            --bg-elevated: #1F2937;        /* Elevated Surface */
            --bg-overlay: rgba(139, 92, 246, 0.1); /* Purple Overlay */
            
            /* Dark Mode Text Colors */
            --text-primary: #F9FAFB;       /* Pure White */
            --text-secondary: #E5E7EB;     /* Light Gray */
            --text-tertiary: #D1D5DB;      /* Medium Gray */
            --text-muted: #9CA3AF;         /* Muted Gray */
            --text-on-primary: #FFFFFF;   /* White text */
            
            /* Dark Mode Borders */
            --border-primary: #374151;     /* Gray Border */
            --border-secondary: #4B5563;   /* Light Border */
            --border-focus: #8B5CF6;       /* Purple Focus */
            --border-subtle: #1F2937;      /* Subtle Border */
            
            /* Dark Mode Shadows */
            --shadow-sm: 0 1px 2px rgba(0, 0, 0, 0.2);
            --shadow-md: 0 4px 6px rgba(0, 0, 0, 0.3);
            --shadow-lg: 0 10px 15px rgba(0, 0, 0, 0.4);
            --shadow-xl: 0 20px 25px rgba(0, 0, 0, 0.5);
            --shadow-glow: 0 0 20px rgba(139, 92, 246, 0.3);
            
            /* Dark Mode Gradients */
        --gradient-primary: linear-gradient(135deg, #8B5CF6 0%, #EC4899 100%);
        --gradient-secondary: linear-gradient(135deg, #14B8A6 0%, #3B82F6 100%);
            --gradient-accent: linear-gradient(135deg, #F97316 0%, #FCD34D 100%);
            --gradient-surface: linear-gradient(135deg, #141824 0%, #1F2937 100%);
            --gradient-subtle: linear-gradient(135deg, #0A0E1A 0%, #111827 100%);
        }
    }
        
        /* Premium adaptive background */
        .stApp {
            background: var(--gradient-subtle);
            color: var(--text-primary);
            min-height: 100vh;
            position: relative;
            transition: background 0.3s ease, color 0.3s ease;
        }
        
        /* Animated gradient overlay - adapts to theme */
        .stApp::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: radial-gradient(circle at 20% 80%, var(--primary-subtle) 0%, transparent 50%),
                        radial-gradient(circle at 80% 20%, var(--secondary-subtle) 0%, transparent 50%),
                        radial-gradient(circle at 40% 40%, var(--accent-subtle) 0%, transparent 50%);
            animation: gradientShift 20s ease infinite;
            pointer-events: none;
            z-index: 0;
            opacity: 0.5;
        }
        
        @keyframes gradientShift {
            0%, 100% { transform: rotate(0deg) scale(1); }
            25% { transform: rotate(90deg) scale(1.1); }
            50% { transform: rotate(180deg) scale(1); }
            75% { transform: rotate(270deg) scale(1.1); }
        }
        
        /* Main container with adaptive glassmorphism */
        .main .block-container {
            background: var(--bg-surface);
            border: 1px solid var(--border-primary);
            border-radius: 24px;
            padding: 2rem;
            backdrop-filter: blur(10px) saturate(150%);
            box-shadow: var(--shadow-lg);
            transition: all 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            z-index: 1;
            transform: translateZ(0);
            will-change: transform, box-shadow;
        }
        
        /* Light mode specific glass effect */
        @media (prefers-color-scheme: light) {
            .main .block-container {
                background: linear-gradient(135deg, 
                    rgba(255, 255, 255, 0.9) 0%, 
                    rgba(249, 250, 251, 0.9) 100%);
                box-shadow: 0 8px 32px rgba(0, 0, 0, 0.08),
                            inset 0 1px 0 rgba(255, 255, 255, 1);
            }
        }
        
        /* Dark mode specific glass effect */
        @media (prefers-color-scheme: dark) {
        .main .block-container {
            background: linear-gradient(135deg, 
                rgba(20, 24, 36, 0.8) 0%, 
                rgba(31, 41, 55, 0.6) 50%,
                rgba(38, 43, 59, 0.8) 100%);
            box-shadow: 0 8px 32px rgba(139, 92, 246, 0.1),
                        inset 0 1px 0 rgba(255, 255, 255, 0.05);
            }
        }
        
        .main .block-container:hover {
            box-shadow: var(--shadow-xl);
            border-color: var(--border-focus);
            transform: translateY(-1px);
        }
        
        /* Typography - Clean and readable */
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Inter', -apple-system, sans-serif;
            color: var(--text-primary);
            font-weight: 600;
            letter-spacing: -0.02em;
            transition: color 0.3s ease;
        }
        
        /* Enhanced paragraph and text visibility */
        p, .stMarkdown, .element-container {
            color: var(--text-primary) !important;
        }
        
        /* Light mode text enhancements */
        @media (prefers-color-scheme: light) {
            .stMarkdown p, .stMarkdown div {
                color: #1F2937 !important;
                font-weight: 500;
            }
            
            /* Bold text should be even darker in light mode */
            .stMarkdown strong, .stMarkdown b {
                color: #111827 !important;
                font-weight: 700;
            }
            
            /* Improve markdown text inside containers */
            .element-container .stMarkdown {
                color: #1F2937 !important;
                font-size: 1rem !important;
                line-height: 1.6 !important;
            }
            
            /* Tab content text visibility */
            [data-testid="stTab"] .stMarkdown {
                color: #374151 !important;
                font-weight: 500 !important;
            }
        }
        
        /* Premium adaptive buttons */
        .stButton > button {
            background: var(--gradient-primary);
            color: var(--text-on-primary);
            border: none;
            border-radius: 20px;
            padding: 0.75rem 1.5rem;
            font-weight: 600;
            font-size: 0.9rem;
            letter-spacing: 0.025em;
            text-transform: uppercase;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            overflow: hidden;
            box-shadow: var(--shadow-md);
            z-index: 1;
            transform: translateZ(0);
            will-change: transform, box-shadow, background;
        }
        
        /* Gradient overlay effect */
        .stButton > button::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(135deg, var(--primary-light) 0%, var(--primary) 100%);
            opacity: 0;
            transition: opacity 0.3s ease;
            z-index: -1;
        }
        
        /* Ripple effect */
        .stButton > button::after {
            content: '';
            position: absolute;
            top: 50%;
            left: 50%;
            width: 0;
            height: 0;
            border-radius: 50%;
            background: rgba(255, 255, 255, 0.3);
            transform: translate(-50%, -50%);
            transition: width 0.6s, height 0.6s;
            z-index: -1;
        }
        
        /* Hover state with adaptive glow */
        .stButton > button:hover {
            transform: translateY(-3px) scale(1.03);
            box-shadow: var(--shadow-glow);
            filter: brightness(1.1);
            transition: all 0.2s ease-out;
        }
        
        .stButton > button:hover::before {
            opacity: 1;
            transition: opacity 0.3s ease;
        }
        
        .stButton > button:hover::after {
            width: 300px;
            height: 300px;
            opacity: 0;
            transition: all 0.6s ease-out;
        }
        
        /* Active/Click state */
        .stButton > button:active {
            transform: translateY(-1px) scale(0.99);
            box-shadow: var(--shadow-sm);
            transition: all 0.1s ease;
        }
        
        /* Success button variant */
        .stButton > button[kind="primary"] {
            background: var(--gradient-secondary);
            box-shadow: 0 4px 15px rgba(20, 184, 166, 0.3);
        }
        
        .stButton > button[kind="primary"]:hover {
            box-shadow: 0 8px 25px rgba(20, 184, 166, 0.4);
            background: linear-gradient(135deg, #5EEAD4 0%, #60A5FA 100%);
        }
        
        /* Secondary button variant */
        .stButton > button[kind="secondary"] {
            background: transparent;
            border: 2px solid var(--primary);
            color: var(--primary-light);
        }
        
        .stButton > button[kind="secondary"]:hover {
            background: var(--primary);
            color: white;
            border-color: var(--primary-light);
        }
        
        /* Button loading state */
        .stButton > button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none !important;
        }
        
        /* Button icon animation */
        .stButton > button i,
        .stButton > button svg {
            transition: transform 0.3s ease;
            margin-right: 0.5rem;
        }
        
        .stButton > button:hover i,
        .stButton > button:hover svg {
            transform: rotate(5deg) scale(1.1);
        }
        
        /* Pulse animation for important buttons */
        @keyframes buttonPulse {
            0% {
                box-shadow: 0 0 0 0 rgba(79, 70, 229, 0.7);
            }
            70% {
                box-shadow: 0 0 0 10px rgba(79, 70, 229, 0);
            }
            100% {
                box-shadow: 0 0 0 0 rgba(79, 70, 229, 0);
            }
        }
        
        .stButton > button.emphasis {
            animation: buttonEmphasis 3s infinite ease-in-out;
        }
        
        @keyframes buttonEmphasis {
            0%, 100% {
                transform: scale(1);
                box-shadow: var(--shadow-md);
            }
            50% {
                transform: scale(1.02);
                box-shadow: var(--shadow-lg);
            }
        }
        
        /* Glow effect for hover */
        @keyframes buttonGlow {
            0%, 100% {
                box-shadow: 0 2px 8px rgba(79, 70, 229, 0.25);
            }
            50% {
                box-shadow: 0 4px 16px rgba(79, 70, 229, 0.5);
            }
        }
        
        .stButton > button:focus {
            outline: none;
            box-shadow: 0 0 0 3px rgba(79, 70, 229, 0.2);
        }
        
        /* Download button special style */
        .stDownloadButton > button {
            background: linear-gradient(135deg, #06B6D4 0%, #0891B2 100%);
            border-radius: 20px;
            font-weight: 600;
            box-shadow: 0 2px 8px rgba(6, 182, 212, 0.25);
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            transform: translateZ(0);
            will-change: transform, box-shadow;
        }
        
        .stDownloadButton > button:hover {
            transform: translateY(-4px) scale(1.03);
            box-shadow: 0 8px 25px rgba(6, 182, 212, 0.45);
            transition: all 0.2s ease-out;
        }
        
        .stDownloadButton > button:active {
            transform: translateY(-1px) scale(0.99);
            transition: all 0.1s ease;
        }
        
        /* Form submit button special style */
        button[type="submit"] {
            width: 100%;
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%) !important;
            font-size: 1rem !important;
            padding: 0.875rem !important;
            margin-top: 0.5rem;
        }
        
        /* Adaptive input fields */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea,
        .stSelectbox > div > div > select,
        .stNumberInput > div > div > input {
            background: var(--bg-elevated);
            border: 2px solid var(--border-primary);
            border-radius: 16px;
            color: var(--text-primary);
            padding: 0.75rem 1rem;
            font-size: 0.95rem;
            font-weight: 500;
            transition: all 0.35s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            box-shadow: var(--shadow-sm);
            position: relative;
            background-clip: padding-box;
            transform: translateZ(0);
            will-change: transform, border-color, box-shadow;
        }
        
        /* Light mode input styling */
        @media (prefers-color-scheme: light) {
            .stTextInput > div > div > input,
            .stTextArea > div > div > textarea,
            .stSelectbox > div > div > select,
            .stNumberInput > div > div > input {
                background: var(--bg-primary);
                border-color: var(--border-secondary);
                box-shadow: inset 0 1px 2px rgba(0, 0, 0, 0.05);
            }
        }
        
        /* Dark mode input styling */
        @media (prefers-color-scheme: dark) {
            .stTextInput > div > div > input,
            .stTextArea > div > div > textarea,
            .stSelectbox > div > div > select,
            .stNumberInput > div > div > input {
                background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-tertiary) 100%);
                box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2);
            }
        }
        
        .stTextInput > div > div > input:hover,
        .stTextArea > div > div > textarea:hover,
        .stSelectbox > div > div > select:hover {
            border-color: var(--border-focus);
            transform: translateY(-2px) scale(1.01);
            box-shadow: var(--shadow-md);
            transition: all 0.2s ease-out;
        }
        
        .stTextInput > div > div > input:focus,
        .stTextArea > div > div > textarea:focus,
        .stSelectbox > div > div > select:focus {
            outline: none;
            border-color: var(--primary);
            box-shadow: 0 0 0 4px var(--primary-subtle), var(--shadow-lg);
            background: var(--bg-elevated);
            transform: translateY(-3px) scale(1.02);
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        
        /* Animated placeholder */
        .stTextInput > div > div > input::placeholder,
        .stTextArea > div > div > textarea::placeholder {
            color: var(--text-secondary);
            transition: all 0.3s ease;
            opacity: 0.7;
        }
        
        .stTextInput > div > div > input:focus::placeholder,
        .stTextArea > div > div > textarea:focus::placeholder {
            opacity: 0.4;
            transform: translateX(5px);
        }
        
        /* Adaptive metrics cards */
        [data-testid="metric-container"] {
            background: var(--bg-surface);
            border: 1px solid var(--border-primary);
            border-radius: 20px;
            padding: 1.25rem;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            overflow: hidden;
            box-shadow: var(--shadow-md);
            transform: translateZ(0);
            will-change: transform, box-shadow;
        }
        
        /* Light mode metrics */
        @media (prefers-color-scheme: light) {
            [data-testid="metric-container"] {
                background: linear-gradient(135deg, var(--bg-primary) 0%, var(--bg-secondary) 100%);
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05),
                            inset 0 1px 0 rgba(255, 255, 255, 1);
            }
        }
        
        /* Dark mode metrics */
        @media (prefers-color-scheme: dark) {
            [data-testid="metric-container"] {
                background: var(--gradient-surface);
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2),
                        inset 0 1px 0 rgba(255, 255, 255, 0.02);
            }
        }
        
        [data-testid="metric-container"]::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 3px;
            background: var(--gradient-primary);
            transition: left 0.3s ease;
            box-shadow: var(--shadow-glow);
        }
        
        [data-testid="metric-container"]:hover {
            transform: translateY(-5px) scale(1.03);
            box-shadow: var(--shadow-lg);
            border-color: var(--primary);
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        
        [data-testid="metric-container"]:hover::before {
            left: 0;
            transition: left 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
        }
        
        [data-testid="metric-container"] [data-testid="metric-value"] {
            background: var(--gradient-primary);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            font-weight: 700;
            font-size: 1.5rem;
            transition: all 0.3s ease;
        }
        
        /* Adaptive tabs */
        .stTabs [data-baseweb="tab-list"] {
            background: var(--bg-surface);
            border: 1px solid var(--border-primary);
            border-radius: 20px;
            padding: 6px;
            box-shadow: var(--shadow-sm);
            position: relative;
        }
        
        /* Light mode tabs */
        @media (prefers-color-scheme: light) {
            .stTabs [data-baseweb="tab-list"] {
                background: var(--bg-secondary);
                box-shadow: inset 0 1px 2px rgba(0, 0, 0, 0.05);
            }
        }
        
        /* Dark mode tabs */
        @media (prefers-color-scheme: dark) {
            .stTabs [data-baseweb="tab-list"] {
                background: var(--gradient-surface);
            box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2),
                        0 2px 8px rgba(139, 92, 246, 0.1);
            }
        }
        
        .stTabs [data-baseweb="tab"] {
            color: var(--text-secondary);
            background: transparent;
            border-radius: 14px;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            font-weight: 500;
            padding: 0.6rem 1.2rem !important;
            overflow: hidden;
            transform: translateZ(0);
            will-change: transform, background, color;
        }
        
        .stTabs [data-baseweb="tab"]::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: var(--bg-overlay);
            opacity: 0;
            transition: opacity 0.3s ease;
            border-radius: 10px;
        }
        
        .stTabs [data-baseweb="tab"]:hover {
            color: var(--text-primary);
            transform: translateY(-2px) scale(1.02);
            background: var(--primary-subtle);
            transition: all 0.2s ease-out;
        }
        
        .stTabs [data-baseweb="tab"]:hover::before {
            opacity: 1;
            transition: opacity 0.3s ease;
        }
        
        .stTabs [aria-selected="true"] {
            background: var(--gradient-primary) !important;
            color: var(--text-on-primary) !important;
            font-weight: 600;
            box-shadow: var(--shadow-md);
            transform: translateY(-2px) scale(1.05);
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        
        /* Tab underline animation */
        .stTabs [data-baseweb="tab-panel"] {
            animation: tabFadeIn 0.6s cubic-bezier(0.25, 0.46, 0.45, 0.94);
        }
        
        @keyframes tabFadeIn {
            from {
                opacity: 0;
                transform: translateX(20px);
            }
            to {
                opacity: 1;
                transform: translateX(0);
            }
        }
        
        /* Progress bar with animation */
        .stProgress > div > div > div > div {
            background: var(--primary);
            transition: width 0.3s ease;
            position: relative;
            overflow: hidden;
        }
        
        .stProgress > div > div > div > div::after {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 2s infinite;
        }
        
        @keyframes shimmer {
            0% { transform: translateX(-100%); }
            100% { transform: translateX(100%); }
        }
        
        /* Adaptive expanders */
        .streamlit-expanderHeader {
            background: var(--bg-elevated);
            border: 1px solid var(--border-primary);
            border-radius: 16px;
            color: var(--text-primary) !important;
            padding: 0.75rem 1rem !important;
            font-weight: 500;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            overflow: hidden;
            transform: translateZ(0);
            will-change: transform, background, border-color;
            cursor: pointer;
        }
        
        /* Light mode expander */
        @media (prefers-color-scheme: light) {
            .streamlit-expanderHeader {
                background: linear-gradient(135deg, var(--bg-secondary) 0%, var(--bg-tertiary) 100%);
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
            }
        }
        
        /* Dark mode expander */
        @media (prefers-color-scheme: dark) {
            .streamlit-expanderHeader {
                background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-tertiary) 100%);
            }
        }
        
        .streamlit-expanderHeader::before {
            content: '';
            position: absolute;
            left: 0;
            top: 0;
            height: 100%;
            width: 3px;
            background: var(--primary);
            transform: scaleY(0);
            transition: transform 0.3s ease;
        }
        
        .streamlit-expanderHeader:hover {
            border-color: var(--primary);
            transform: translateX(6px) scale(1.01);
            box-shadow: var(--shadow-md);
            transition: all 0.2s ease-out;
        }
        
        .streamlit-expanderHeader:hover::before {
            transform: scaleY(1);
            transition: transform 0.3s cubic-bezier(0.25, 0.46, 0.45, 0.94);
        }
        
        /* Expander content animation */
        .streamlit-expanderContent {
            animation: expanderSlide 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            background: var(--bg-secondary);
            border-radius: 0 0 16px 16px;
            padding: 1rem !important;
            transform-origin: top;
        }
        
        @keyframes expanderSlide {
            from {
                opacity: 0;
                transform: translateY(-20px) scaleY(0.95);
            }
            to {
                opacity: 1;
                transform: translateY(0) scaleY(1);
            }
        }
        
        /* Checkboxes and Radio buttons */
        .stCheckbox > label,
        .stRadio > label {
            transition: all 0.3s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            padding: 0.5rem;
            border-radius: 12px;
            cursor: pointer;
        }
        
        .stCheckbox > label:hover,
        .stRadio > label:hover {
            background: rgba(79, 70, 229, 0.08);
            transform: translateX(4px) scale(1.02);
            border-radius: 12px;
            transition: all 0.2s ease-out;
        }
        
        /* Custom checkbox styling */
        .stCheckbox > label > span:first-child,
        .stRadio > div > label > span:first-child {
            border: 2px solid var(--border) !important;
            background: var(--bg-medium) !important;
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            border-radius: 6px !important;
            transform: scale(1);
        }
        
        .stCheckbox > label > span:first-child:checked,
        .stRadio > div > label > span:first-child:checked {
            background: var(--primary) !important;
            border-color: var(--primary) !important;
            animation: checkConfirm 0.3s ease-out;
        }
        
        @keyframes checkConfirm {
            0% { transform: scale(0.8); opacity: 0.5; }
            50% { transform: scale(1.05); }
            100% { transform: scale(1); opacity: 1; }
        }
        
        /* Enhanced Slider styling - Force purple theme */
        .stSlider > div > div > div {
            background: rgba(107, 114, 128, 0.3) !important;
            border-radius: 8px !important;
            height: 6px !important;
        }
        
        .stSlider > div > div > div > div {
            background: linear-gradient(90deg, #8B5CF6, #7C3AED) !important;
            transition: all 0.2s ease;
            border-radius: 8px !important;
            height: 6px !important;
        }
        
        .stSlider > div > div > div > div > div {
            background: #8B5CF6 !important;
            border: 3px solid #FFFFFF !important;
            border-radius: 50% !important;
            box-shadow: 0 2px 8px rgba(139, 92, 246, 0.4) !important;
            transition: all 0.2s ease;
            width: 20px !important;
            height: 20px !important;
        }
        
        /* Override Streamlit's default colors */
        div[data-baseweb="slider"] {
            --slider-color-0: #8B5CF6 !important;
            --slider-color-1: #7C3AED !important;
        }
        
        .stSlider > div > div > div > div > div:hover {
            transform: scale(1.2);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.5) !important;
        }
        
        /* Text status displays for checkboxes and sliders */
        .status-indicator {
            font-family: 'SF Mono', 'Monaco', 'Courier New', monospace;
            font-weight: 600;
            padding: 4px 8px;
            border-radius: 6px;
            display: inline-block;
        }
        
        /* Slider value text display */
        h3 code {
            color: var(--primary) !important;
            background: var(--bg-medium) !important;
            padding: 8px 12px !important;
            border-radius: 8px !important;
            font-size: 1.2rem !important;
        }
        
        /* Premium alerts with enhanced animations */
        .stAlert {
            animation: alertSlideIn 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            border-radius: 10px;
            border-left: 4px solid;
            background: linear-gradient(135deg, rgba(30, 41, 59, 0.9) 0%, rgba(30, 41, 59, 0.6) 100%);
            backdrop-filter: blur(10px);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            position: relative;
            overflow: hidden;
        }
        
        .stAlert::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 3s infinite;
        }
        
        /* Adaptive alerts */
        .stAlert {
            border-radius: 16px;
            border: 1px solid var(--border-primary);
            animation: alertSlideIn 0.5s ease;
        }
        
        /* Success alert */
        .stAlert[data-baseweb="notification"][kind="success"] {
            border-left: 4px solid var(--success) !important;
            color: var(--text-primary);
        }
        
        /* Error alert */
        .stAlert[data-baseweb="notification"][kind="error"] {
            border-left: 4px solid var(--danger) !important;
            color: var(--text-primary);
        }
        
        /* Warning alert */
        .stAlert[data-baseweb="notification"][kind="warning"] {
            border-left: 4px solid var(--warning) !important;
            color: var(--text-primary);
        }
        
        /* Info alert */
        .stAlert[data-baseweb="notification"][kind="info"] {
            border-left: 4px solid var(--info) !important;
            color: var(--text-primary);
        }
        
        /* Light mode alert backgrounds */
        @media (prefers-color-scheme: light) {
            .stAlert[data-baseweb="notification"][kind="success"] {
                background: linear-gradient(135deg, #D1FAE5 0%, var(--bg-primary) 100%);
                box-shadow: 0 2px 8px rgba(16, 185, 129, 0.15);
            }
            .stAlert[data-baseweb="notification"][kind="error"] {
                background: linear-gradient(135deg, #FEE2E2 0%, var(--bg-primary) 100%);
                box-shadow: 0 2px 8px rgba(244, 63, 94, 0.15);
            }
            .stAlert[data-baseweb="notification"][kind="warning"] {
                background: linear-gradient(135deg, #FEF3C7 0%, var(--bg-primary) 100%);
                box-shadow: 0 2px 8px rgba(234, 179, 8, 0.15);
            }
            .stAlert[data-baseweb="notification"][kind="info"] {
                background: linear-gradient(135deg, #DBEAFE 0%, var(--bg-primary) 100%);
                box-shadow: 0 2px 8px rgba(59, 130, 246, 0.15);
            }
        }
        
        /* Dark mode alert backgrounds */
        @media (prefers-color-scheme: dark) {
            .stAlert[data-baseweb="notification"][kind="success"] {
                background: linear-gradient(135deg, rgba(16, 185, 129, 0.15) 0%, var(--bg-surface) 100%);
                box-shadow: 0 0 20px rgba(16, 185, 129, 0.2);
            }
            .stAlert[data-baseweb="notification"][kind="error"] {
                background: linear-gradient(135deg, rgba(244, 63, 94, 0.15) 0%, var(--bg-surface) 100%);
                box-shadow: 0 0 20px rgba(244, 63, 94, 0.2);
            }
            .stAlert[data-baseweb="notification"][kind="warning"] {
                background: linear-gradient(135deg, rgba(234, 179, 8, 0.15) 0%, var(--bg-surface) 100%);
                box-shadow: 0 0 20px rgba(234, 179, 8, 0.2);
            }
            .stAlert[data-baseweb="notification"][kind="info"] {
            background: linear-gradient(135deg, rgba(59, 130, 246, 0.15) 0%, var(--bg-surface) 100%);
            box-shadow: 0 0 20px rgba(59, 130, 246, 0.2);
            }
        }
        
        @keyframes alertSlideIn {
            from {
                opacity: 0;
                transform: translateX(-15px);
            }
            to {
                opacity: 1;
                transform: translateX(0);
            }
        }
        
        /* Success/Error/Warning/Info messages */
        .stSuccess, .stError, .stWarning, .stInfo {
            border-radius: 14px !important;
            padding: 1rem 1.25rem !important;
            font-weight: 500;
            animation: messagePopIn 0.5s cubic-bezier(0.34, 1.56, 0.64, 1);
            transform: translateZ(0);
            will-change: transform, opacity;
        }
        
        /* Enhanced visibility for info messages */
        .stInfo {
            background: var(--info-light) !important;
            color: var(--text-primary) !important;
            border: 2px solid var(--info) !important;
            font-weight: 600 !important;
        }
        
        /* Enhanced visibility for warning messages */
        .stWarning {
            background: var(--warning-light) !important;
            color: var(--text-primary) !important;
            border: 2px solid var(--warning) !important;
            font-weight: 600 !important;
        }
        
        /* Light mode specific enhancements for better visibility */
        @media (prefers-color-scheme: light) {
            .stInfo {
                background: linear-gradient(135deg, #EFF6FF 0%, #DBEAFE 100%) !important;
                color: #1E40AF !important;
                border: 2px solid #3B82F6 !important;
                box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15) !important;
                font-size: 1rem !important;
                font-weight: 600 !important;
            }
            
            .stWarning {
                background: linear-gradient(135deg, #FFFBEB 0%, #FEF3C7 100%) !important;
                color: #92400E !important;
                border: 2px solid #F59E0B !important;
                box-shadow: 0 4px 12px rgba(245, 158, 11, 0.15) !important;
                font-size: 1rem !important;
                font-weight: 600 !important;
            }
            
            .stSuccess {
                background: linear-gradient(135deg, #F0FDF4 0%, #DCFCE7 100%) !important;
                color: #14532D !important;
                border: 2px solid #22C55E !important;
                box-shadow: 0 4px 12px rgba(34, 197, 94, 0.15) !important;
                font-size: 1rem !important;
                font-weight: 600 !important;
            }
            
            .stError {
                background: linear-gradient(135deg, #FEF2F2 0%, #FEE2E2 100%) !important;
                color: #7F1D1D !important;
                border: 2px solid #EF4444 !important;
                box-shadow: 0 4px 12px rgba(239, 68, 68, 0.15) !important;
                font-size: 1rem !important;
                font-weight: 600 !important;
            }
            
            /* Ensure text inside alerts is highly visible */
            .stInfo p, .stWarning p, .stSuccess p, .stError p {
                color: inherit !important;
                font-weight: 600 !important;
                margin: 0 !important;
            }
            
            /* Additional Streamlit alert container styling for light mode */
            [data-testid="stAlert"] {
                background: linear-gradient(135deg, #F9FAFB 0%, #F3F4F6 100%) !important;
                border: 2px solid var(--primary) !important;
                border-radius: 16px !important;
                padding: 1.25rem !important;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08) !important;
            }
            
            [data-testid="stAlert"] > div {
                color: #111827 !important;
                font-weight: 600 !important;
                font-size: 1rem !important;
            }
            
            /* Info specific */
            div[kind="info"] [data-testid="stAlert"],
            .element-container:has(.stInfo) {
                background: linear-gradient(135deg, #EFF6FF 0%, #DBEAFE 100%) !important;
                border-left: 4px solid #3B82F6 !important;
            }
            
            /* Warning specific */
            div[kind="warning"] [data-testid="stAlert"],
            .element-container:has(.stWarning) {
                background: linear-gradient(135deg, #FFFBEB 0%, #FEF3C7 100%) !important;
                border-left: 4px solid #F59E0B !important;
            }
        }
        
        /* Dark mode specific styles for messages */
        @media (prefers-color-scheme: dark) {
            .stInfo {
                background: linear-gradient(135deg, rgba(59, 130, 246, 0.2) 0%, rgba(59, 130, 246, 0.1) 100%) !important;
                color: #93BBFC !important;
                border: 2px solid #3B82F6 !important;
            }
            
            .stWarning {
                background: linear-gradient(135deg, rgba(245, 158, 11, 0.2) 0%, rgba(245, 158, 11, 0.1) 100%) !important;
                color: #FCD34D !important;
                border: 2px solid #F59E0B !important;
            }
            
            .stSuccess {
                background: linear-gradient(135deg, rgba(34, 197, 94, 0.2) 0%, rgba(34, 197, 94, 0.1) 100%) !important;
                color: #86EFAC !important;
                border: 2px solid #22C55E !important;
            }
            
            .stError {
                background: linear-gradient(135deg, rgba(239, 68, 68, 0.2) 0%, rgba(239, 68, 68, 0.1) 100%) !important;
                color: #FCA5A5 !important;
                border: 2px solid #EF4444 !important;
            }
        }
        
        @keyframes messagePopIn {
            0% {
                opacity: 0;
                transform: scale(0.95) translateY(-8px);
            }
            60% {
                transform: scale(1.02) translateY(2px);
            }
            100% {
                opacity: 1;
                transform: scale(1) translateY(0);
            }
        }
        
        /* Adaptive DataFrames and Tables */
        .stDataFrame {
            border-radius: 16px;
            overflow: hidden;
            box-shadow: var(--shadow-md);
            border: 1px solid var(--border-primary);
        }
        
        .stDataFrame > div {
            background: var(--bg-surface) !important;
        }
        
        .stDataFrame table {
            background: transparent !important;
            color: var(--text-primary) !important;
        }
        
        .stDataFrame thead tr th {
            background: var(--gradient-primary) !important;
            color: var(--text-on-primary) !important;
            font-weight: 600 !important;
            text-transform: uppercase !important;
            letter-spacing: 0.05em !important;
            padding: 0.75rem !important;
            border: none !important;
        }
        
        .stDataFrame tbody tr {
            transition: all 0.2s ease;
            background: var(--bg-surface);
        }
        
        .stDataFrame tbody tr:hover {
            background: var(--primary-subtle) !important;
            transform: scale(1.01);
        }
        
        .stDataFrame tbody tr td {
            padding: 0.75rem !important;
            border-bottom: 1px solid var(--border-primary) !important;
            font-weight: 500;
            color: var(--text-primary) !important;
        }
        
        /* Light mode table striping */
        @media (prefers-color-scheme: light) {
            .stDataFrame tbody tr:nth-child(even) {
                background: var(--bg-secondary);
            }
        }
        
        /* Adaptive Code blocks */
        .stCodeBlock {
            border-radius: 16px !important;
            border: 1px solid var(--border-primary) !important;
            box-shadow: var(--shadow-md);
        }
        
        /* Light mode code blocks */
        @media (prefers-color-scheme: light) {
            .stCodeBlock {
                background: linear-gradient(135deg, #F8F9FA 0%, #E9ECEF 100%) !important;
            }
            
            .stCodeBlock pre {
                color: #212529 !important;
                font-family: 'Fira Code', 'Consolas', monospace !important;
                padding: 1rem !important;
            }
        }
        
        /* Dark mode code blocks */
        @media (prefers-color-scheme: dark) {
            .stCodeBlock {
            background: linear-gradient(135deg, #1a1f2e 0%, #0f172a 100%) !important;
        }
        
        .stCodeBlock pre {
            color: #e2e8f0 !important;
            font-family: 'Fira Code', 'Consolas', monospace !important;
            padding: 1rem !important;
            }
        }
        
        /* Columns and containers */
        [data-testid="column"] {
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            padding: 0.5rem;
            transform: translateZ(0);
            will-change: transform;
        }
        
        [data-testid="column"]:hover {
            transform: translateY(-3px) scale(1.005);
            transition: all 0.2s ease-out;
        }
        
        
        /* Multiselect */
        .stMultiSelect > div {
            border-radius: 16px;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            transform: translateZ(0);
            will-change: transform, box-shadow;
        }
        
        .stMultiSelect > div:hover {
            transform: translateY(-3px) scale(1.01);
            box-shadow: 0 6px 20px rgba(79, 70, 229, 0.15);
            transition: all 0.2s ease-out;
        }
        
        /* Date input */
        .stDateInput > div {
            border-radius: 10px;
            transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            transform: translateZ(0);
            will-change: transform, box-shadow;
        }
        
        .stDateInput > div:hover {
            transform: translateY(-3px) scale(1.01);
            box-shadow: 0 6px 16px rgba(79, 70, 229, 0.15);
            transition: all 0.2s ease-out;
        }
        
        /* Spinner animation */
        .stSpinner > div {
            border-color: var(--primary) !important;
        }
        
        /* Enhanced spinner */
        .stSpinner {
            position: relative;
        }
        
        .stSpinner::before {
            content: '⚙️';
            position: absolute;
            font-size: 2rem;
            animation: spin 2s linear infinite;
            opacity: 0.3;
        }
        
        @keyframes spin {
            from { transform: rotate(0deg); }
            to { transform: rotate(360deg); }
        }
        
        /* Empty state */
        .stEmpty {
            padding: 2rem;
            text-align: center;
            color: var(--text-secondary);
            font-style: italic;
        }
        
        /* Modal/Dialog styling */
        [data-baseweb="modal"] {
            background: rgba(15, 23, 42, 0.8) !important;
            backdrop-filter: blur(5px);
        }
        
        [data-baseweb="modal"] > div {
            background: linear-gradient(135deg, var(--bg-medium) 0%, var(--bg-dark) 100%) !important;
            border: 1px solid var(--border);
            border-radius: 12px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.5);
            animation: modalSlideIn 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }
        
        @keyframes modalSlideIn {
            from {
                opacity: 0;
                transform: translateY(-50px) scale(0.9);
            }
            to {
                opacity: 1;
                transform: translateY(0) scale(1);
            }
        }
        
        /* Color picker */
        .stColorPicker > div > div {
            border-radius: 10px;
            border: 2px solid var(--border);
            transition: all 0.3s ease;
        }
        
        .stColorPicker > div > div:hover {
            border-color: var(--primary);
            transform: scale(1.05);
        }
        
        /* Select slider */
        .stSelectSlider > div {
            background: var(--bg-medium);
            border-radius: 16px;
            padding: 0.5rem;
        }
        
        /* JSON viewer */
        .stJson {
            background: linear-gradient(135deg, #1a1f2e 0%, #0f172a 100%) !important;
            border: 1px solid var(--border);
            border-radius: 16px;
            padding: 1rem;
            font-family: 'Fira Code', monospace;
        }
        
        /* Special hover effects for interactive elements */
        .clickable-element {
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .clickable-element:hover {
            transform: translateY(-2px) scale(1.02);
            filter: brightness(1.1);
        }
        
        /* Tooltips */
        [data-baseweb="tooltip"] {
            background: var(--primary) !important;
            border-radius: 8px;
            padding: 0.5rem 0.75rem;
            font-weight: 500;
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
            animation: tooltipFadeIn 0.2s ease;
        }
        
        @keyframes tooltipFadeIn {
            from {
                opacity: 0;
                transform: scale(0.8);
            }
            to {
                opacity: 1;
                transform: scale(1);
            }
        }
        
        /* Premium file uploader with animations */
        [data-testid="stFileUploadDropzone"] {
            background: linear-gradient(135deg, var(--bg-medium) 0%, rgba(30, 41, 59, 0.4) 100%);
            border: 2px dashed var(--border);
            border-radius: 20px;
            padding: 2rem;
            transition: all 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            position: relative;
            overflow: hidden;
            transform: translateZ(0);
            will-change: transform, border-color, background;
        }
        
        [data-testid="stFileUploadDropzone"]::before {
            content: '📁';
            position: absolute;
            font-size: 4rem;
            opacity: 0.1;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            transition: all 0.4s ease;
        }
        
        [data-testid="stFileUploadDropzone"]:hover {
            border-color: var(--primary);
            background: linear-gradient(135deg, rgba(79, 70, 229, 0.1) 0%, rgba(79, 70, 229, 0.05) 100%);
            transform: scale(1.03) translateY(-3px);
            box-shadow: 0 12px 32px rgba(79, 70, 229, 0.2);
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        
        [data-testid="stFileUploadDropzone"]:hover::before {
            opacity: 0.2;
            transform: translate(-50%, -50%) scale(1.3) rotate(10deg);
            transition: all 0.4s cubic-bezier(0.68, -0.55, 0.265, 1.55);
        }
        
        /* Drag over state */
        [data-testid="stFileUploadDropzone"][data-dragging="true"] {
            border-color: var(--accent);
            background: linear-gradient(135deg, rgba(16, 185, 129, 0.1) 0%, rgba(16, 185, 129, 0.05) 100%);
            animation: subtleFade 2s infinite;
        }
        
        /* File upload button inside */
        [data-testid="stFileUploadDropzone"] button {
            background: var(--primary) !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            transition: all 0.3s ease !important;
        }
        
        [data-testid="stFileUploadDropzone"] button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
        }
        
        /* Custom scrollbar */
        ::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }
        
        ::-webkit-scrollbar-track {
            background: var(--bg-medium);
        }
        
        ::-webkit-scrollbar-thumb {
            background: var(--border);
            border-radius: 4px;
            transition: background 0.2s ease;
        }
        
        ::-webkit-scrollbar-thumb:hover {
            background: var(--primary);
        }
        
        /* Smooth page transitions */
        * {
            transition: background-color 0.2s ease, 
                        border-color 0.2s ease,
                        box-shadow 0.2s ease;
        }
        
        /* Loading animation */
        /* Main header and title styles for light/dark mode */
        .main-title {
            font-size: 3rem;
            font-weight: 700;
            margin-bottom: 0.5rem;
            letter-spacing: -0.03em;
            background: var(--gradient-primary);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .main-subtitle {
            font-size: 1.1rem;
            color: var(--text-secondary);
            margin-top: 0.5rem;
            font-weight: 400;
        }
        
        .feature-badge {
            background: var(--bg-elevated);
            border: 1px solid var(--border-primary);
            padding: 0.375rem 0.875rem;
            border-radius: 12px;
            color: var(--text-secondary);
            font-size: 0.875rem;
            font-weight: 500;
            transition: all 0.2s ease;
            display: inline-block;
        }
        
        .feature-badge:hover {
            background: var(--primary-subtle);
            border-color: var(--primary);
            color: var(--text-primary);
            transform: translateY(-2px);
        }
        
        /* Light mode specific adjustments for main header */
        @media (prefers-color-scheme: light) {
            .main-title {
                /* Ensure gradient text is visible in light mode */
                background: linear-gradient(135deg, var(--primary-dark) 0%, var(--primary) 100%);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                background-clip: text;
            }
            
            .feature-badge {
                background: linear-gradient(135deg, #F3F4F6 0%, #E5E7EB 100%);
                border: 2px solid var(--primary);
                color: var(--primary-dark);
                font-weight: 600;
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
            }
            
            .feature-badge:hover {
                background: var(--primary);
                color: white;
                transform: translateY(-3px);
                box-shadow: 0 4px 12px rgba(124, 58, 237, 0.25);
            }
        }
        
        /* Dark mode specific adjustments for main header */
        @media (prefers-color-scheme: dark) {
            .main-title {
                background: var(--gradient-primary);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                background-clip: text;
            }
            
            .feature-badge {
                background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-tertiary) 100%);
                border: 1px solid var(--border-secondary);
            }
        }
        
        /* Enhanced AI thinking animation container */
        .ai-thinking-container {
            text-align: center;
            padding: 2rem;
            background: var(--bg-elevated);
            border: 2px solid var(--border-primary);
            border-radius: 20px;
            margin: 1.5rem 0;
            box-shadow: var(--shadow-lg);
            position: relative;
            overflow: hidden;
            animation: slideInScale 0.5s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        
        /* Light mode specific for AI thinking */
        @media (prefers-color-scheme: light) {
            .ai-thinking-container {
                background: linear-gradient(135deg, 
                    rgba(255, 255, 255, 0.98) 0%, 
                    rgba(249, 250, 251, 0.98) 100%);
                border: 2px solid rgba(124, 58, 237, 0.2);
                box-shadow: 0 10px 40px rgba(124, 58, 237, 0.08),
                           0 4px 12px rgba(0, 0, 0, 0.04),
                           0 0 0 1px rgba(124, 58, 237, 0.05) inset;
                backdrop-filter: blur(10px);
            }
            
            .ai-thinking-icon {
                color: var(--primary);
                text-shadow: 0 2px 8px rgba(124, 58, 237, 0.2);
            }
            
            .loading-message {
                color: #1F2937 !important;
                text-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
            }
        }
        
        /* Dark mode specific for AI thinking */
        @media (prefers-color-scheme: dark) {
            .ai-thinking-container {
                background: linear-gradient(135deg, 
                    rgba(30, 41, 59, 0.98) 0%, 
                    rgba(15, 23, 42, 0.98) 100%);
                border: 2px solid rgba(139, 92, 246, 0.3);
                box-shadow: 0 12px 48px rgba(139, 92, 246, 0.15),
                           0 0 80px rgba(139, 92, 246, 0.1),
                           0 0 0 1px rgba(139, 92, 246, 0.1) inset;
                backdrop-filter: blur(12px);
            }
            
            .ai-thinking-icon {
                color: var(--primary-light);
                filter: drop-shadow(0 0 20px rgba(139, 92, 246, 0.5));
            }
            
            .loading-message {
                color: #E5E7EB !important;
                text-shadow: 0 2px 4px rgba(0, 0, 0, 0.3);
            }
        }
        
        /* Animated background gradient */
        .ai-thinking-container::before {
            content: '';
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(
                45deg,
                transparent,
                var(--primary-subtle),
                transparent
            );
            animation: shimmerRotate 6s linear infinite;
            opacity: 0.3;
        }
        
        /* Icon animation */
        .ai-thinking-icon {
            font-size: 3rem;
            margin-bottom: 1rem;
            display: inline-block;
            animation: iconRotate 4s ease-in-out infinite;
            filter: drop-shadow(0 4px 8px rgba(124, 58, 237, 0.3));
        }
        
        /* Loading message styles */
        .loading-message {
            color: var(--text-primary);
            font-size: 1.1rem;
            font-weight: 600;
            margin: 0.5rem 0;
            position: relative;
            z-index: 1;
            letter-spacing: 0.5px;
        }
        
        /* Progress indicator */
        .loading-progress {
            margin-top: 1rem;
            padding: 0.5rem 1rem;
            background: var(--gradient-primary);
            color: var(--text-on-primary);
            font-size: 0.9rem;
            font-weight: 600;
            border-radius: 12px;
            display: inline-block;
            position: relative;
            z-index: 1;
            box-shadow: 0 2px 8px rgba(124, 58, 237, 0.2);
            animation: progressGlow 2s ease-in-out infinite;
        }
        
        /* Progress bar background */
        .progress-bar-container {
            width: 100%;
            height: 6px;
            background: var(--bg-secondary);
            border-radius: 3px;
            margin-top: 1rem;
            overflow: hidden;
            position: relative;
        }
        
        .progress-bar-fill {
            height: 100%;
            background: var(--gradient-primary);
            border-radius: 3px;
            transition: width 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
            box-shadow: 0 0 10px var(--primary);
        }
        
        /* Enhanced animations */
        @keyframes slideInScale {
            from { 
                opacity: 0; 
                transform: translateY(10px) scale(0.97); 
            }
            to { 
                opacity: 1; 
                transform: translateY(0) scale(1); 
            }
        }
        
        @keyframes iconRotate {
            0%, 100% { 
                transform: translateY(0) rotate(0deg) scale(1); 
            }
            25% { 
                transform: translateY(-5px) rotate(-3deg) scale(1.05); 
            }
            50% { 
                transform: translateY(0) rotate(0deg) scale(1); 
            }
            75% { 
                transform: translateY(-3px) rotate(3deg) scale(1.02); 
            }
        }
        
        @keyframes shimmerRotate {
            0% { 
                transform: rotate(0deg) translateX(-100%);
                opacity: 0.3; 
            }
            50% {
                opacity: 0.5;
            }
            100% { 
                transform: rotate(180deg) translateX(100%);
                opacity: 0.3; 
            }
        }
        
        @keyframes progressGlow {
            0%, 100% { 
                transform: scale(1);
                box-shadow: 0 2px 8px rgba(124, 58, 237, 0.15);
                opacity: 0.95;
            }
            50% { 
                transform: scale(1.02);
                box-shadow: 0 3px 12px rgba(124, 58, 237, 0.25);
                opacity: 1;
            }
        }
        
        @keyframes subtleFade {
            0%, 100% { 
                opacity: 1; 
                transform: scale(1);
            }
            50% { 
                opacity: 0.9; 
                transform: scale(0.98);
            }
        }
        
        /* Typing dots animation */
        .typing-dots {
            display: inline-flex;
            gap: 4px;
            margin-left: 8px;
        }
        
        .typing-dot {
            width: 6px;
            height: 6px;
            background: var(--primary);
            border-radius: 50%;
            animation: typingDot 1.6s infinite ease-in-out;
            opacity: 0.6;
        }
        
        .typing-dot:nth-child(2) {
            animation-delay: 0.15s;
        }
        
        .typing-dot:nth-child(3) {
            animation-delay: 0.3s;
        }
        
        @keyframes typingDot {
            0%, 80%, 100% {
                transform: scale(1);
                opacity: 0.8;
            }
            40% {
                transform: scale(1.2);
                opacity: 1;
            }
        }
        
        /* Fade in animation for content */
        @keyframes fadeIn {
            from { 
                opacity: 0; 
                transform: translateY(10px); 
            }
            to { 
                opacity: 1; 
                transform: translateY(0); 
            }
        }
        
        .element-container {
            animation: fadeIn 0.5s ease;
        }
        
        /* Enhanced Streamlit spinner styles */
        .stSpinner > div {
            text-align: center !important;
        }
        
        .stSpinner svg {
            width: 40px !important;
            height: 40px !important;
            animation: spinnerRotate 1s cubic-bezier(0.68, -0.55, 0.265, 1.55) infinite !important;
        }
        
        /* Light mode spinner */
        @media (prefers-color-scheme: light) {
            .stSpinner svg circle {
                stroke: var(--primary) !important;
                stroke-width: 3px !important;
                filter: drop-shadow(0 2px 4px rgba(124, 58, 237, 0.2));
            }
        }
        
        /* Dark mode spinner */
        @media (prefers-color-scheme: dark) {
            .stSpinner svg circle {
                stroke: var(--primary-light) !important;
                stroke-width: 3px !important;
                filter: drop-shadow(0 2px 8px rgba(139, 92, 246, 0.4));
            }
        }
        
        @keyframes spinnerRotate {
            0% { 
                transform: rotate(0deg); 
            }
            100% { 
                transform: rotate(360deg); 
            }
        }
        
        /* Spinner text styling */
        .stSpinner > div > div {
            color: var(--text-primary) !important;
            font-weight: 600 !important;
            font-size: 1rem !important;
            margin-top: 1rem !important;
            letter-spacing: 0.5px !important;
        }
    </style>
    """, unsafe_allow_html=True)

# AI Thinking Messages
AI_THINKING_MESSAGES = [
    "Analyzing requirement specifications...",
    "Searching compliance standards database...",
    "Identifying critical test scenarios...",
    "Applying healthcare testing protocols...",
    "Calculating test coverage metrics...",
    "Synthesizing test case components...",
    "Optimizing test execution sequence...",
    "Validating patient safety requirements...",
    "Ensuring regulatory compliance...",
    "Processing domain knowledge base..."
]

def show_ai_thinking(duration=3, style="default"):
    """Display animated AI thinking messages with smooth transitions
    
    Args:
        duration: Duration in seconds
        style: Animation style - "default", "minimal", or "detailed"
    """
    placeholder = st.empty()
    
    if style == "minimal":
        # Simpler, cleaner animation
        for i in range(duration * 2):
            msg = random.choice(AI_THINKING_MESSAGES)
            placeholder.markdown(f"""
            <div class='ai-thinking-container' style='padding: 1.5rem;'>
                <p class='loading-message' style='margin: 0;'>
                    {msg}
                    <span class='typing-dots'>
                        <span class='typing-dot'></span>
                        <span class='typing-dot'></span>
                        <span class='typing-dot'></span>
                    </span>
                </p>
            </div>
            """, unsafe_allow_html=True)
            time.sleep(0.5)
    
    elif style == "detailed":
        # More elaborate animation with multiple elements
        for i in range(duration * 2):
            msg = random.choice(AI_THINKING_MESSAGES)
            icon = '🧠🔍💡🏥📊🔬⚡🎯🛡️🤖'[i % 10]
            secondary_msg = random.choice([
                "Processing healthcare compliance rules...",
                "Analyzing test patterns...",
                "Optimizing test coverage...",
                "Validating requirements..."
            ])
            progress_val = (i + 1) / (duration * 2)
            
            placeholder.markdown(f"""
            <div class='ai-thinking-container'>
                <div class='ai-thinking-icon'>
                    {icon}
                </div>
                <p class='loading-message'>
                {msg}
            </p>
                <p class='loading-message' style='font-size: 0.9rem; opacity: 0.8; font-weight: 400;'>
                    {secondary_msg}
                </p>
                <div class='loading-progress'>
                    {int(progress_val * 100)}% Complete
                </div>
                <div class='progress-bar-container'>
                    <div class='progress-bar-fill' style='width: {int(progress_val * 100)}%;'></div>
            </div>
            </div>
            """, unsafe_allow_html=True)
            time.sleep(0.6)
    
    else:
        # Default animation style
        for i in range(duration * 2):
            msg = random.choice(AI_THINKING_MESSAGES)
            icon = '🧠🔍💡🏥📊🔬⚡🎯🛡️🤖'[i % 10]
            progress_val = (i + 1) / (duration * 2)
            
            placeholder.markdown(f"""
            <div class='ai-thinking-container'>
                <div class='ai-thinking-icon'>
                    {icon}
                </div>
                <p class='loading-message'>
                    {msg}
                    <span class='typing-dots'>
                        <span class='typing-dot'></span>
                        <span class='typing-dot'></span>
                        <span class='typing-dot'></span>
                    </span>
                </p>
                <div class='loading-progress'>
                    {int(progress_val * 100)}% Complete
                </div>
                <div class='progress-bar-container'>
                    <div class='progress-bar-fill' style='width: {int(progress_val * 100)}%;'></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        time.sleep(0.5)
    
    # Final success animation
    placeholder.markdown(f"""
    <div class='ai-thinking-container' style='border-color: var(--success);'>
        <div class='ai-thinking-icon' style='animation: none; color: var(--success); font-size: 1.5rem;'>
            ✓
        </div>
        <p class='loading-message'>
            <strong>Analysis Complete</strong>
        </p>
        <div class='loading-progress' style='background: var(--gradient-secondary);'>
            Processing Complete
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    time.sleep(0.5)
    placeholder.empty()

def show_quick_loader(message="Processing...", icon="⚡"):
    """Show a quick, simple loading indicator"""
    return st.markdown(f"""
    <div class='ai-thinking-container' style='padding: 1rem; margin: 0.5rem 0;'>
        <div style='display: flex; align-items: center; justify-content: center; gap: 1rem;'>
            <div class='ai-thinking-icon' style='font-size: 1.5rem; margin: 0;'>
                {icon}
            </div>
            <p class='loading-message' style='margin: 0;'>
                {message}
                <span class='typing-dots'>
                    <span class='typing-dot'></span>
                    <span class='typing-dot'></span>
                    <span class='typing-dot'></span>
                </span>
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Configuration - SIMPLIFIED FOLDER STRUCTURE (v9)
# Knowledge Base & Documents
DOCUMENTS_PATH = "data/documents"  # Sample docs for RAG (PRDs, user stories, API specs)
UPLOADED_DOCS_PATH = "data/uploaded_documents"  # User uploads + compliance reports

# Test Cases - SINGLE UNIFIED LOCATION
TEST_CASES_DIR = "data/test_cases"  # ALL test cases (generated, imported, manual saves)

# Vector Database
FAISS_INDEX_PATH = "data/faiss_indices"  # FAISS vector database for RAG

# Models
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2'
GEMINI_MODEL_NAME = 'gemini-2.0-flash-exp'  # Latest model

# NASSCOM Compliance Requirements
NASSCOM_REQUIREMENTS = {
    "document_types": {
        "prd": "Product Requirement Document",
        "user_story": "User Stories with acceptance criteria",
        "api_spec": "API Specifications (OpenAPI/Swagger)",
        "mockup": "UI/UX Mockups and wireframes",
        "test_plan": "Existing test plans or strategies"
    },
    "required_fields": [
        "Test Case ID", "Title", "Description", 
        "Test Steps", "Expected Results", "Priority"
    ],
    "compliance_standards": [
        "HIPAA", "GDPR", "FDA", "ISO 9001", 
        "ISO 13485", "ISO 27001"
    ],
    "healthcare_keywords": [
        "patient", "medical", "health", "clinical", "diagnosis",
        "treatment", "medication", "prescription", "PHI", "EHR",
        "DICOM", "HL7", "medical device", "FDA compliance"
    ],
    "document_formats": ["pdf", "docx", "doc", "xml", "yaml", "yml", "txt", "md"],
    "quality_criteria": {
        "min_content_length": 100,
        "requires_structure": True,
        "requires_healthcare_context": False  # Made optional for flexibility
    }
}

# Page configuration already set at line 48
# Removed duplicate st.set_page_config() call

# Custom CSS for better styling (removed - now handled in inject_modern_css)
st.markdown("""
<style>
    /* Adaptive metric card styles */
    .metric-card {
        background: var(--bg-surface);
        padding: 10px;
        border-radius: 12px;
        margin: 5px;
        border: 1px solid var(--border-primary);
    }
    
    /* Adaptive compliance status styles */
    .compliance-pass {
        background: var(--success-subtle, rgba(16, 185, 129, 0.1));
        color: var(--success);
        padding: 10px;
        border-radius: 12px;
        margin: 5px 0;
        border: 1px solid var(--success);
    }
    
    .compliance-fail {
        background: var(--danger-subtle, rgba(244, 63, 94, 0.1));
        color: var(--danger);
        padding: 10px;
        border-radius: 12px;
        margin: 5px 0;
        border: 1px solid var(--danger);
    }
    
    .compliance-warning {
        background: var(--warning-subtle, rgba(234, 179, 8, 0.1));
        color: var(--warning);
        padding: 10px;
        border-radius: 12px;
        margin: 5px 0;
        border: 1px solid var(--warning);
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'test_counter' not in st.session_state:
    st.session_state.test_counter = 0
# Don't initialize generated_tests here - it will be loaded after authentication
# if 'generated_tests' not in st.session_state:
#     st.session_state.generated_tests = []
if 'refinement_mode' not in st.session_state:
    st.session_state.refinement_mode = False
if 'test_to_refine' not in st.session_state:
    st.session_state.test_to_refine = None
if 'refinement_history' not in st.session_state:
    st.session_state.refinement_history = {}
if 'edited_test' not in st.session_state:
    st.session_state.edited_test = None
if 'test_versions' not in st.session_state:
    st.session_state.test_versions = {}
if 'uploaded_docs' not in st.session_state:
    st.session_state.uploaded_docs = []
if 'compliance_reports' not in st.session_state:
    st.session_state.compliance_reports = {}
if 'imported_tests' not in st.session_state:
    st.session_state.imported_tests = []
if 'import_history' not in st.session_state:
    st.session_state.import_history = []
if 'last_converted_tests' not in st.session_state:
    st.session_state.last_converted_tests = None
if 'last_import_report' not in st.session_state:
    st.session_state.last_import_report = None
if 'tests_loaded' not in st.session_state:
    st.session_state.tests_loaded = False

# Configure Gemini
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# ===============================
# 🗄️ MONGODB INTEGRATION (NEW IN V11)
# ===============================

@st.cache_resource
def init_mongodb():
    """Initialize MongoDB connection (cached to persist across reruns)"""
    try:
        # Get MongoDB URI from environment
        mongodb_uri = os.getenv('MONGODB_URI')
        
        if not mongodb_uri:
            logger.warning("MongoDB URI not found in environment variables")
            st.warning("MongoDB not configured. Using local file storage only.")
            return None
        
        # Ensure database name is in URI (append if missing)
        if 'mongodb.net/' in mongodb_uri and '/genai_hack_app' not in mongodb_uri:
            # Add database name if missing
            if mongodb_uri.endswith('?retryWrites=true&w=majority'):
                mongodb_uri = mongodb_uri.replace('?retryWrites=true&w=majority', '/genai_hack_app?retryWrites=true&w=majority')
            elif '?' in mongodb_uri:
                mongodb_uri = mongodb_uri.replace('?', '/genai_hack_app?')
            else:
                mongodb_uri += '/genai_hack_app'
            logger.info("Added database name to MongoDB URI")
        
        db = MongoDBManager(mongodb_uri)
        logger.info("✅ MongoDB initialized successfully")
        return db
    except Exception as e:
        logger.error(f"❌ MongoDB initialization failed: {e}")
        st.error(f"Failed to connect to MongoDB: {e}")
        return None

def convert_numpy_to_python(obj):
    """Recursively convert numpy types to Python types for MongoDB compatibility"""
    import numpy as np
    
    if isinstance(obj, dict):
        return {key: convert_numpy_to_python(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_to_python(item) for item in obj]
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, np.bool_):
        return bool(obj)
    elif isinstance(obj, (np.int8, np.int16, np.int32, np.int64)):
        return int(obj)
    elif isinstance(obj, (np.float16, np.float32, np.float64)):
        return float(obj)
    elif isinstance(obj, np.str_):
        return str(obj)
    else:
        return obj

def get_or_create_session():
    """Get existing session for user or create new one with user ownership"""
    if 'session_id' not in st.session_state:
        if 'db' in st.session_state and st.session_state.db:
            # Get user_id from session state (set during login)
            user_id = st.session_state.get('user_id')
            if user_id:
                # Get or create session linked to user
                st.session_state.session_id = st.session_state.db.get_or_create_session_for_user(user_id)
                logger.info(f"[SESSION] User {user_id} session: {st.session_state.session_id}")
            else:
                # Fallback for non-authenticated users (shouldn't happen in v13)
                st.session_state.session_id = st.session_state.db.create_session()
        else:
            st.session_state.session_id = f"LOCAL_{uuid.uuid4().hex[:8]}"
    return st.session_state.session_id

def save_test_to_mongodb(test_case: Dict) -> bool:
    """Save test case to MongoDB with user ownership"""
    if 'db' not in st.session_state or not st.session_state.db:
        return False
    
    try:
        # Convert numpy types to Python types for MongoDB compatibility
        clean_test = convert_numpy_to_python(test_case)
        
        session_id = get_or_create_session()
        user_id = st.session_state.get('user_id')
        success, test_id = st.session_state.db.save_test_case(clean_test, session_id, user_id)
        if success:
            logger.info(f"✅ Test case {test_id} saved to MongoDB for user {user_id}")
        return success
    except Exception as e:
        logger.error(f"Failed to save to MongoDB: {e}")
        return False

def load_tests_from_mongodb(limit: int = 100) -> List[Dict]:
    """Load test cases from MongoDB for current user"""
    if 'db' not in st.session_state or not st.session_state.db:
        logger.warning("[LOAD_TESTS] No database connection available")
        return []
    
    try:
        session_id = get_or_create_session()
        user_id = st.session_state.get('user_id')
        logger.info(f"[LOAD_TESTS] Loading tests for user_id={user_id}, session_id={session_id}")
        
        # Load only user's test cases for data isolation
        tests = st.session_state.db.get_all_test_cases(session_id=session_id, user_id=user_id, limit=limit)
        logger.info(f"[LOAD_TESTS] Raw query returned {len(tests)} tests")
        
        # Clean numpy types from loaded tests
        tests = [convert_numpy_to_python(test) for test in tests]
        logger.info(f"[LOAD_TESTS] After cleaning: {len(tests)} tests for user {user_id}")
        return tests
    except Exception as e:
        logger.error(f"[LOAD_TESTS] Failed to load from MongoDB: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return []

def save_document_to_mongodb(filename: str, content: str, doc_type: str, metadata: Dict = None) -> bool:
    """Save document to MongoDB with user ownership"""
    if 'db' not in st.session_state or not st.session_state.db:
        return False
    
    try:
        session_id = get_or_create_session()
        user_id = st.session_state.get('user_id')
        doc_id = st.session_state.db.save_document(filename, content, doc_type, metadata, session_id, user_id)
        logger.info(f"Document {filename} saved for user {user_id}")
        return doc_id is not None
    except Exception as e:
        logger.error(f"Failed to save document to MongoDB: {e}")
        return False

def save_compliance_to_mongodb(report: Dict) -> bool:
    """Save compliance report to MongoDB with user ownership"""
    if 'db' not in st.session_state or not st.session_state.db:
        return False
    
    try:
        session_id = get_or_create_session()
        user_id = st.session_state.get('user_id')
        report_id = st.session_state.db.save_compliance_report(report, session_id, user_id)
        logger.info(f"Compliance report saved for user {user_id}")
        return report_id is not None
    except Exception as e:
        logger.error(f"Failed to save compliance report to MongoDB: {e}")
        return False

# ===============================
# Auto-Save/Load Functions (NEW IN V7)
# ===============================

def auto_save_test_cases(test_cases: List[Dict], prefix: str = "tests"):
    """Automatically save test cases to unified directory for persistence"""
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    
    # Save as a single JSON file with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{TEST_CASES_DIR}/{prefix}_{timestamp}.json"
    
    # Also save as 'latest' for easy loading
    latest_filename = f"{TEST_CASES_DIR}/{prefix}_latest.json"
    
    try:
        # Save timestamped version
        with open(filename, 'w') as f:
            json.dump(test_cases, f, indent=2, default=str)
        
        # Save/overwrite latest version
        with open(latest_filename, 'w') as f:
            json.dump(test_cases, f, indent=2, default=str)
        
        return True, filename
    except Exception as e:
        st.error(f"Failed to auto-save tests: {e}")
        return False, None

def load_saved_test_cases(prefix: str = "tests") -> List[Dict]:
    """Load previously saved test cases from unified directory"""
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    
    # Try to load the latest file first
    latest_filename = f"{TEST_CASES_DIR}/{prefix}_latest.json"
    
    try:
        if os.path.exists(latest_filename):
            with open(latest_filename, 'r') as f:
                test_cases = json.load(f)
                return test_cases if isinstance(test_cases, list) else []
    except Exception as e:
        # If latest fails, try to load most recent timestamped file
        try:
            files = glob.glob(f"{TEST_CASES_DIR}/{prefix}_*.json")
            # Filter out 'latest' file and sort by modification time
            files = [f for f in files if not f.endswith('_latest.json')]
            if files:
                files.sort(key=os.path.getmtime, reverse=True)
                with open(files[0], 'r') as f:
                    test_cases = json.load(f)
                    return test_cases if isinstance(test_cases, list) else []
        except Exception as inner_e:
            st.warning(f"Could not load saved tests: {inner_e}")
    
    return []

def cleanup_old_saves(prefix: str = "*", keep_count: int = 10):
    """Keep only the most recent saves to avoid disk bloat"""
    try:
        files = glob.glob(f"{TEST_CASES_DIR}/{prefix}_*.json")
        # Exclude 'latest' files
        files = [f for f in files if not f.endswith('_latest.json')]
        
        if len(files) > keep_count:
            # Sort by modification time
            files.sort(key=os.path.getmtime)
            # Delete oldest files
            for f in files[:-keep_count]:
                os.remove(f)
    except Exception as e:
        pass  # Silently fail cleanup

# Load saved tests on startup
@st.cache_resource(show_spinner=False)
def initialize_saved_tests():
    """Load previously saved tests on app startup - unified from single directory"""
    all_saved_tests = load_saved_test_cases("all_tests")
    if not all_saved_tests:
        # For backward compatibility, try loading old format
        generated = load_saved_test_cases("generated")
        imported = load_saved_test_cases("imported")
        all_saved_tests = generated + imported
    return all_saved_tests

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from Word document"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting Word document text: {e}")
        return ""

def extract_text_from_xml(file_content: bytes) -> str:
    """Extract text content from XML file"""
    try:
        # Try to detect encoding
        detected = chardet.detect(file_content)
        encoding = detected['encoding'] or 'utf-8'
        
        text_content = file_content.decode(encoding)
        root = ET.fromstring(text_content)
        
        # Convert XML to readable format
        def extract_element_text(element, prefix=""):
            result = []
            if element.tag:
                result.append(f"{prefix}{element.tag}:")
            if element.text and element.text.strip():
                result.append(f"{prefix}  {element.text.strip()}")
            for attr_name, attr_value in element.attrib.items():
                result.append(f"{prefix}  @{attr_name}: {attr_value}")
            for child in element:
                result.extend(extract_element_text(child, prefix + "  "))
            return result
        
        text_lines = extract_element_text(root)
        return "\n".join(text_lines)
    except Exception as e:
        st.error(f"Error extracting XML text: {e}")
        return ""

def analyze_document_compliance(content: str, filename: str, file_type: str) -> Dict:
    """
    Analyze document for NASSCOM compliance using Gemini AI
    """
    prompt = f"""You are a compliance analyst for NASSCOM's AI Test Case Generation requirements.
Analyze the following document and determine its compliance with NASSCOM guidelines.

DOCUMENT NAME: {filename}
FILE TYPE: {file_type}

DOCUMENT CONTENT:
{content[:3000]}  # Limiting to first 3000 chars for analysis

NASSCOM REQUIREMENTS:
1. Document Types Accepted: PRDs, User Stories, API Specs, Mockups, Test Plans
2. Healthcare/MedTech Context: Should contain healthcare-related terminology
3. Structure: Should have clear sections and be well-organized
4. Content Quality: Sufficient detail for test case generation
5. Compliance Standards: Should reference or be compatible with HIPAA, GDPR, FDA, ISO standards

Analyze and return a JSON object with:
{{
    "is_compliant": boolean,
    "compliance_score": float (0-100),
    "document_type": string (prd/user_story/api_spec/mockup/test_plan/other),
    "detected_standards": [list of compliance standards found],
    "healthcare_relevance": float (0-100),
    "structural_quality": float (0-100),
    "content_completeness": float (0-100),
    "strengths": [list of compliance strengths],
    "improvements_needed": [list of improvements needed],
    "recommendations": [list of specific recommendations],
    "summary": string (brief compliance summary)
}}"""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        compliance_report = json.loads(response.text)
        
        # Add metadata
        compliance_report['analyzed_at'] = datetime.now().isoformat()
        compliance_report['filename'] = filename
        compliance_report['file_type'] = file_type
        compliance_report['content_length'] = len(content)
        
        return compliance_report
        
    except Exception as e:
        # Fallback basic analysis
        return {
            "is_compliant": len(content) > 100,
            "compliance_score": 50.0 if len(content) > 100 else 20.0,
            "document_type": "other",
            "detected_standards": [],
            "healthcare_relevance": 0.0,
            "structural_quality": 50.0,
            "content_completeness": 50.0,
            "strengths": ["Document uploaded successfully"],
            "improvements_needed": ["Could not perform AI analysis"],
            "recommendations": ["Please ensure document contains relevant content"],
            "summary": f"Basic analysis completed. Error in AI analysis: {str(e)}",
            "analyzed_at": datetime.now().isoformat(),
            "filename": filename,
            "file_type": file_type,
            "content_length": len(content)
        }

def process_uploaded_file(uploaded_file) -> Tuple[str, Dict]:
    """
    Process uploaded file and extract content with compliance analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    
    # Read file content
    file_content = uploaded_file.read()
    
    # Extract text based on file type
    if file_extension == 'pdf':
        text_content = extract_text_from_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        text_content = extract_text_from_docx(file_content)
    elif file_extension == 'xml':
        text_content = extract_text_from_xml(file_content)
    elif file_extension in ['yaml', 'yml']:
        try:
            text_content = file_content.decode('utf-8')
            yaml_data = yaml.safe_load(text_content)
            text_content = yaml.dump(yaml_data, default_flow_style=False)
        except:
            text_content = file_content.decode('utf-8', errors='ignore')
    elif file_extension in ['txt', 'md']:
        text_content = file_content.decode('utf-8', errors='ignore')
    else:
        # Try to decode as text
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
        except:
            text_content = ""
    
    # Analyze compliance
    compliance_report = analyze_document_compliance(
        text_content, 
        uploaded_file.name, 
        file_extension
    )
    
    return text_content, compliance_report

def save_uploaded_document(filename: str, content: str, compliance_report: Dict) -> bool:
    """
    Save uploaded document to MongoDB if available, otherwise to file system
    """
    try:
        # Try MongoDB first
        if st.session_state.db:
            # Determine document type
            file_ext = filename.split('.')[-1].lower()
            
            # Save document to MongoDB
            doc_saved = save_document_to_mongodb(
                filename=filename,
                content=content,
                doc_type=file_ext,
                metadata={'compliance_score': compliance_report.get('compliance_score', 0)}
            )
            
            # Save compliance report to MongoDB
            compliance_saved = save_compliance_to_mongodb(compliance_report)
            
            if doc_saved and compliance_saved:
                logger.info(f"✅ Document {filename} saved to MongoDB")
                return True
        
        # Fallback to file system
        # Create uploaded documents directory if it doesn't exist
        os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
        
        # Save the document
        file_path = os.path.join(UPLOADED_DOCS_PATH, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Save compliance report
        report_path = os.path.join(UPLOADED_DOCS_PATH, f"{filename}.compliance.json")
        with open(report_path, 'w') as f:
            json.dump(compliance_report, f, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Error saving document: {e}")
        logger.error(f"Error saving document {filename}: {e}")
        return False

def display_compliance_report(report: Dict):
    """
    Display compliance analysis report with visual indicators
    """
    # Overall compliance status
    if report['is_compliant']:
        st.markdown('<div class="compliance-pass">✅ Document is NASSCOM Compliant</div>', 
                   unsafe_allow_html=True)
    else:
        st.markdown('<div class="compliance-fail">❌ Document needs improvements for full compliance</div>', 
                   unsafe_allow_html=True)
    
    # Compliance metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        score_color = "🟢" if report['compliance_score'] >= 70 else "🟡" if report['compliance_score'] >= 40 else "🔴"
        st.metric("Compliance Score", f"{score_color} {report['compliance_score']:.1f}%")
    
    with col2:
        relevance_color = "🟢" if report['healthcare_relevance'] >= 60 else "🟡" if report['healthcare_relevance'] >= 30 else "🔴"
        st.metric("Healthcare Relevance", f"{relevance_color} {report['healthcare_relevance']:.1f}%")
    
    with col3:
        quality_color = "🟢" if report['structural_quality'] >= 70 else "🟡" if report['structural_quality'] >= 40 else "🔴"
        st.metric("Structural Quality", f"{quality_color} {report['structural_quality']:.1f}%")
    
    with col4:
        completeness_color = "🟢" if report['content_completeness'] >= 70 else "🟡" if report['content_completeness'] >= 40 else "🔴"
        st.metric("Content Completeness", f"{completeness_color} {report['content_completeness']:.1f}%")
    
    # Document type and standards
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Document Analysis")
        st.write(f"**Document Type:** {report.get('document_type', 'Unknown').replace('_', ' ').title()}")
        st.write(f"**File Type:** {report.get('file_type', 'Unknown').upper()}")
        st.write(f"**Content Length:** {report.get('content_length', 0):,} characters")
        
        if report.get('detected_standards'):
            st.write("**Detected Standards:**")
            for standard in report['detected_standards']:
                st.write(f"  • {standard}")
    
    with col2:
        st.markdown("### 📊 Compliance Summary")
        st.info(report.get('summary', 'No summary available'))
    
    # Strengths and improvements in tabs
    tab1, tab2, tab3 = st.tabs(["✅ Strengths", "⚠️ Improvements Needed", "💡 Recommendations"])
    
    with tab1:
        if report.get('strengths'):
            for strength in report['strengths']:
                st.success(f"• {strength}")
        else:
            st.info("No specific strengths identified")
    
    with tab2:
        if report.get('improvements_needed'):
            for improvement in report['improvements_needed']:
                st.warning(f"• {improvement}")
        else:
            st.success("No improvements needed")
    
    with tab3:
        if report.get('recommendations'):
            for rec in report['recommendations']:
                st.info(f"• {rec}")
        else:
            st.info("No specific recommendations")

@st.cache_resource
def load_rag_system():
    """
    Load embedding model and build FAISS index from documents.
    Includes both pre-existing and uploaded documents.
    """
    # Load embedding model
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    
    # Create necessary directories
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    os.makedirs(DOCUMENTS_PATH, exist_ok=True)
    os.makedirs(UPLOADED_DOCS_PATH, exist_ok=True)
    os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
    
    # Load and process documents from multiple sources
    doc_chunks = []
    doc_metadata = []
    all_files = []
    
    # 1. Load shared documents from MongoDB (available to all users)
    if st.session_state.db:
        try:
            shared_docs = st.session_state.db.get_all_shared_documents(limit=100)
            logger.info(f"[RAG] Loading {len(shared_docs)} shared documents from MongoDB")
            
            # Process each shared document
            for doc_info in shared_docs:
                try:
                    # Get full document content
                    full_doc = st.session_state.db.get_shared_document(doc_info['_id'])
                    if full_doc and 'content' in full_doc:
                        content = full_doc['content']
                        file_name = full_doc.get('filename', 'shared_doc')
                        
                        # Process content based on type
                        if 'user_stor' in file_name.lower():
                            # Split by user stories
                            stories = content.split('\n#### User Story:')
                            for story in stories:
                                if story.strip():
                                    doc_chunks.append(story.strip())
                                    doc_metadata.append({
                                        'source': f"[Shared] {file_name}",
                                        'type': 'user_story',
                                        'doc_type': 'shared_document',
                                        'is_shared': True
                                    })
                        else:
                            # Standard chunking
                            chunks = [content[i:i+1000] for i in range(0, len(content), 800)]
                            for chunk in chunks:
                                if chunk.strip():
                                    doc_chunks.append(chunk.strip())
                                    doc_metadata.append({
                                        'source': f"[Shared] {file_name}",
                                        'type': full_doc.get('doc_type', 'document'),
                                        'doc_type': 'shared_document',
                                        'is_shared': True
                                    })
                        logger.info(f"[RAG] Processed shared document: {file_name}")
                except Exception as e:
                    logger.error(f"[RAG] Error processing shared document {doc_info.get('filename', 'unknown')}: {e}")
        except Exception as e:
            logger.error(f"[RAG] Failed to load shared documents from MongoDB: {e}")
    
    # 2. Process file-based documents (fallback and legacy support)
    existing_files = glob.glob(f"{DOCUMENTS_PATH}/*")
    uploaded_files = glob.glob(f"{UPLOADED_DOCS_PATH}/*")
    uploaded_files = [f for f in uploaded_files if not f.endswith('.compliance.json')]
    
    all_files = existing_files + uploaded_files
    
    # If no MongoDB shared docs and no file docs, show warning
    if not doc_chunks and not all_files:
        st.warning(f"No documents found. Upload documents to get started or check MongoDB connection.")
        return embedding_model, None, None, []
    
    # Process each document
    for file_path in all_files:
        file_name = Path(file_path).name
        file_ext = Path(file_path).suffix
        
        # Determine if it's an uploaded document
        is_uploaded = UPLOADED_DOCS_PATH in file_path
        
        try:
            if file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                    # Smart chunking for different document types
                    if 'user_story' in file_name.lower():
                        # Split by acceptance criteria
                        if 'Acceptance Criteria:' in content:
                            parts = content.split('Acceptance Criteria:')
                            if parts[0].strip():
                                doc_chunks.append(parts[0].strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'user_story_overview',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                            if len(parts) > 1:
                                doc_chunks.append(f"Acceptance Criteria:\n{parts[1].strip()}")
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'acceptance_criteria',
                                    'doc_type': 'user_story',
                                    'is_uploaded': is_uploaded
                                })
                        else:
                            doc_chunks.append(content)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'user_story',
                                'doc_type': 'user_story',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # Split by paragraphs for other documents
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip() and len(para.strip()) > 50:
                                doc_chunks.append(para.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'prd' if 'prd' in file_name.lower() else 'document',
                                    'doc_type': 'prd' if 'prd' in file_name.lower() else 'general',
                                    'is_uploaded': is_uploaded
                                })
                                
            elif file_ext in ['.yaml', '.yml']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    yaml_content = yaml.safe_load(f)
                    
                    # Process API specs
                    if 'paths' in yaml_content:
                        for path, methods in yaml_content.get('paths', {}).items():
                            endpoint_doc = f"API Endpoint: {path}\n"
                            endpoint_doc += yaml.dump(methods, default_flow_style=False)
                            doc_chunks.append(endpoint_doc)
                            doc_metadata.append({
                                'source': file_name,
                                'type': 'api_endpoint',
                                'endpoint': path,
                                'doc_type': 'api_specification',
                                'is_uploaded': is_uploaded
                            })
                    else:
                        # General YAML content
                        doc_chunks.append(yaml.dump(yaml_content, default_flow_style=False))
                        doc_metadata.append({
                            'source': file_name,
                            'type': 'config',
                            'doc_type': 'configuration',
                            'is_uploaded': is_uploaded
                        })
            
            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_pdf(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'pdf_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext in ['.docx', '.doc']:
                with open(file_path, 'rb') as f:
                    text = extract_text_from_docx(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'word_document',
                                    'doc_type': 'document',
                                    'is_uploaded': is_uploaded
                                })
            
            elif file_ext == '.xml':
                with open(file_path, 'rb') as f:
                    text = extract_text_from_xml(f.read())
                    if text:
                        # Split into chunks
                        chunks = [text[i:i+1000] for i in range(0, len(text), 800)]
                        for chunk in chunks:
                            if chunk.strip():
                                doc_chunks.append(chunk.strip())
                                doc_metadata.append({
                                    'source': file_name,
                                    'type': 'xml_document',
                                    'doc_type': 'specification',
                                    'is_uploaded': is_uploaded
                                })
                        
        except Exception as e:
            st.error(f"Error loading {file_name}: {str(e)}")
    
    if not doc_chunks:
        st.warning("No content extracted from documents.")
        return embedding_model, None, None, []
    
    # Generate embeddings
    with st.spinner(f"Creating embeddings for {len(doc_chunks)} document chunks..."):
        contents = doc_chunks
        embeddings = embedding_model.encode(contents, show_progress_bar=True)
        embeddings = np.array(embeddings).astype('float32')
    
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    
    return embedding_model, index, doc_chunks, doc_metadata

def retrieve_context(query: str, embedding_model, index, doc_chunks, doc_metadata, k: int = 5):
    """Retrieve relevant documents for a query"""
    if index is None:
        return []
    
    query_embedding = embedding_model.encode([query]).astype('float32')
    distances, indices = index.search(query_embedding, k)
    
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(doc_chunks):
            similarity = 1 / (1 + dist)
            results.append({
                'content': doc_chunks[idx],
                'metadata': doc_metadata[idx] if idx < len(doc_metadata) else {},
                'similarity': similarity
            })
    
    return results

def generate_test_case_with_gemini(requirement: str, context_docs: List[Dict]) -> Dict:
    """
    Generate test case using Gemini with NASSCOM compliance
    """
    # Check if this is an API test request
    api_keywords = ['api', 'endpoint', 'rest', 'http', 'get', 'post', 'put', 'delete', 
                   'request', 'response', 'webhook', 'graphql', '/api/', 'status code']
    is_api_test = any(keyword in requirement.lower() for keyword in api_keywords)
    
    # Check context relevance
    avg_relevance = sum(doc['similarity'] for doc in context_docs) / len(context_docs) if context_docs else 0
    low_relevance = avg_relevance < 0.3
    
    # Build context string
    if low_relevance:
        context_str = "Note: Limited specific context found. Generating based on general best practices and requirement analysis."
    else:
        context_parts = []
        for doc in context_docs[:3]:
            context_parts.append(f"--- Source: {doc['metadata'].get('source', 'Unknown')} (Type: {doc['metadata'].get('type', 'document')}, Relevance: {doc['similarity']:.2%}) ---")
            context_parts.append(doc['content'][:800])
            context_parts.append("--- End Context ---")
        context_str = "\n\n".join(context_parts)
    
    # Add API-specific instructions if needed
    api_instructions = """
ADDITIONAL API TEST REQUIREMENTS:
- Set category to "API" or "Integration"  
- test_data MUST include: method (GET/POST/PUT/DELETE), endpoint (e.g., /api/patients), headers, body, expected_status_code
- Include authentication details in test_data.headers if needed
- Test steps should include: prepare request, send request, verify response
- Consider response time validation
""" if is_api_test else ""
    
    # Enhanced prompt with NASSCOM requirements
    prompt = f"""You are a meticulous QA Engineer specializing in healthcare/MedTech testing, following NASSCOM guidelines.
Generate a comprehensive test case for the following requirement.

REQUIREMENT:
{requirement}

RELEVANT CONTEXT:
{context_str}
{api_instructions}
NASSCOM COMPLIANCE REQUIREMENTS:
- Must include all required fields: Test Case ID, Title, Description, Test Steps, Expected Results, Priority
- Should consider healthcare compliance standards: HIPAA, GDPR, FDA, ISO 13485, ISO 27001
- Should handle edge cases and boundary conditions
- Must be traceable to requirements
- Should support both manual and automated testing where feasible

Generate a test case with EXACTLY these fields:
- id: Use "TC_PLACEHOLDER" (will be replaced)
- title: Clear, descriptive title
- description: Detailed description of what is being tested
- category: One of [Functional, Security, Integration, Performance, Usability, Compliance]
- priority: One of [Critical, High, Medium, Low]
- compliance: Array of applicable standards like ["HIPAA", "GDPR", "FDA", "ISO 13485", "ISO 27001"]
- preconditions: String describing what must be true before testing
- test_steps: Array of step descriptions WITHOUT numbering (e.g., ["Navigate to login page", "Enter credentials"])
- expected_results: String describing specific expected outcomes
- test_data: JSON object with test data. For API tests, include: {{"method": "GET", "endpoint": "/api/endpoint", "headers": {{}}, "body": {{}}, "expected_status_code": 200}}. For other tests: {{"username": "test@example.com", "password": "Test123!"}}
- edge_cases: Array of special scenarios to consider
- negative_tests: Array of negative test scenarios
- automation_feasible: Boolean indicating if this can be automated
- estimated_duration: String like "5 minutes" or "30 minutes"
- traceability: String linking to requirement source

IMPORTANT:
1. test_steps must be an array of strings without prefixes like "Step 1:" or "1."
2. test_data must be a JSON object, not a string
3. Focus on healthcare/medical domain requirements if applicable
4. Include both positive and negative test scenarios
5. Ensure compliance with relevant healthcare standards

Return ONLY a valid JSON object."""
    
    try:
        # Configure Gemini for JSON output
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                top_k=30,
                max_output_tokens=4096
            )
        )
        
        # Safety settings for healthcare content
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
        
        response = model.generate_content(prompt, safety_settings=safety_settings)
        
        if not response.parts:
            return create_fallback_test_case(requirement, error_reason="Content blocked by safety filter")
        
        test_case = json.loads(response.text)
        
        # Generate unique ID
        unique_id = f"TC_{str(uuid.uuid4())[:8].upper()}"
        test_case['id'] = unique_id
        
        # Clean test_steps formatting
        if 'test_steps' in test_case and isinstance(test_case['test_steps'], list):
            cleaned_steps = []
            for step in test_case['test_steps']:
                cleaned = re.sub(r'^(Step\s+\d+:|^\d+\.\s*)', '', str(step)).strip()
                cleaned_steps.append(cleaned)
            test_case['test_steps'] = cleaned_steps
        
        # Ensure test_data is always a dict
        if 'test_data' in test_case:
            if isinstance(test_case['test_data'], str):
                try:
                    test_case['test_data'] = json.loads(test_case['test_data'])
                except:
                    test_case['test_data'] = {"raw_data": test_case['test_data']}
        else:
            test_case['test_data'] = {}
        
        # Add metadata and context tracking
        test_case['generated_from'] = requirement[:100]
        test_case['generation_timestamp'] = datetime.now().isoformat()
        test_case['avg_context_relevance'] = f"{avg_relevance:.2%}"
        test_case['context_sources'] = [
            doc['metadata'].get('source', 'Unknown') for doc in context_docs[:3]
        ] if not low_relevance else ["Generated using best practices"]
        test_case['low_context_confidence'] = low_relevance
        test_case['version'] = 1
        test_case['nasscom_compliant'] = True
        
        # Increment counter
        st.session_state.test_counter += 1
        
        return test_case
        
    except Exception as e:
        st.error(f"Generation error: {e}")
        return create_fallback_test_case(requirement, error_reason=str(e))

def create_fallback_test_case(requirement: str, error_reason: str = "Unknown") -> Dict:
    """
    Create a fallback test case when generation fails.
    """
    st.session_state.test_counter += 1
    return {
        'id': f'TC_FB_{st.session_state.test_counter:04d}',
        'title': f'Test: {requirement[:60]}',
        'description': f'Verify that {requirement}',
        'category': 'Functional',
        'priority': 'Medium',
        'compliance': ['General'],
        'preconditions': 'System is in stable state',
        'test_steps': [
            'Setup test environment',
            'Execute test scenario',
            'Verify results'
        ],
        'expected_results': 'System behaves as specified',
        'test_data': {'placeholder': 'Add specific test data'},
        'edge_cases': [],
        'negative_tests': [],
        'automation_feasible': False,
        'estimated_duration': '10 minutes',
        'traceability': requirement[:100],
        'generated_from': requirement[:100],
        'context_sources': ['Fallback generation'],
        'generation_timestamp': datetime.now().isoformat(),
        'nasscom_compliant': False,
        'fallback': True,
        'error': error_reason
    }

def analyze_test_suite_schema(content: str, file_type: str) -> Dict:
    """
    Use AI to analyze the schema/structure of an uploaded test suite
    """
    prompt = f"""You are an expert test case analyst. Analyze the following test suite data and understand its schema.

FILE TYPE: {file_type}
CONTENT SAMPLE (first 2000 chars):
{content[:2000]}

Analyze the test suite and return a JSON object with:
{{
    "detected_format": string (csv/json/excel/xml/custom),
    "has_headers": boolean,
    "field_mappings": {{
        "id_field": string or null,
        "title_field": string or null,
        "description_field": string or null,
        "steps_field": string or null,
        "expected_field": string or null,
        "priority_field": string or null,
        "category_field": string or null,
        "preconditions_field": string or null,
        "test_data_field": string or null
    }},
    "detected_fields": [list of all field names found],
    "row_count": integer (estimated number of test cases),
    "separator": string (for CSV - comma, tab, pipe, etc),
    "schema_confidence": float (0-100),
    "sample_data": object (one sample test case in original format),
    "recommendations": [list of recommendations for import]
}}

Be intelligent about detecting field mappings even if names don't match exactly.
For example, "Test Name" could map to "title_field", "Steps" to "steps_field", etc."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.2
            )
        )
        
        response = model.generate_content(prompt)
        schema_analysis = json.loads(response.text)
        return schema_analysis
        
    except Exception as e:
        st.error(f"Schema analysis error: {e}")
        return {
            "detected_format": "unknown",
            "has_headers": True,
            "field_mappings": {},
            "detected_fields": [],
            "row_count": 0,
            "schema_confidence": 0,
            "sample_data": {},
            "recommendations": ["Manual mapping may be required"]
        }

def convert_test_with_ai(test_data: Dict, schema: Dict, import_id: str) -> Dict:
    """
    Use AI to intelligently convert a test from any format to our standard format
    """
    prompt = f"""You are converting a test case to NASSCOM-compliant format.

ORIGINAL TEST DATA:
{json.dumps(test_data, indent=2)}

DETECTED SCHEMA MAPPINGS:
{json.dumps(schema.get('field_mappings', {}), indent=2)}

Convert this test case to our standard format with these EXACT fields:
{{
    "id": string (generate unique ID starting with {import_id}),
    "title": string (clear, descriptive title),
    "description": string (detailed description),
    "category": string (Functional/Security/Integration/Performance/Usability/Compliance),
    "priority": string (Critical/High/Medium/Low),
    "compliance": array (applicable standards like ["HIPAA", "GDPR", "FDA"]),
    "preconditions": string (what must be true before testing),
    "test_steps": array of strings (clear step descriptions),
    "expected_results": string (specific expected outcomes),
    "test_data": object (test data as JSON object),
    "edge_cases": array (special scenarios),
    "negative_tests": array (negative test scenarios),
    "automation_feasible": boolean,
    "estimated_duration": string (e.g., "10 minutes"),
    "traceability": string (original requirement reference),
    "imported": true,
    "original_format": "{schema.get('detected_format', 'unknown')}",
    "import_timestamp": "{datetime.now().isoformat()}",
    "nasscom_compliant": true
}}

Be intelligent about:
1. Parsing multi-line text into arrays for test_steps
2. Inferring missing fields from context
3. Detecting healthcare/medical context for compliance standards
4. Splitting combined fields appropriately
5. Converting string representations to proper types

Return ONLY a valid JSON object."""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        converted_test = json.loads(response.text)
        
        # Ensure unique ID
        if not converted_test.get('id'):
            converted_test['id'] = f"{import_id}_{str(uuid.uuid4())[:8].upper()}"
        
        return converted_test
        
    except Exception as e:
        # Fallback conversion
        return {
            'id': f"{import_id}_{str(uuid.uuid4())[:8].upper()}",
            'title': test_data.get('title', test_data.get('name', 'Imported Test')),
            'description': test_data.get('description', str(test_data)),
            'category': 'Functional',
            'priority': 'Medium',
            'compliance': [],
            'preconditions': test_data.get('preconditions', 'N/A'),
            'test_steps': [str(test_data.get('steps', 'See original data'))],
            'expected_results': test_data.get('expected', 'As specified'),
            'test_data': test_data,
            'imported': True,
            'import_error': str(e),
            'nasscom_compliant': False
        }

def detect_duplicates(new_tests: List[Dict], existing_tests: List[Dict]) -> List[Dict]:
    """
    Detect potential duplicate test cases using AI
    """
    if not existing_tests or not new_tests:
        return []
    
    duplicates = []
    
    # Create summaries for comparison
    existing_summaries = [
        f"{t.get('title', '')} - {t.get('description', '')[:100]}"
        for t in existing_tests
    ]
    
    for new_test in new_tests:
        new_summary = f"{new_test.get('title', '')} - {new_test.get('description', '')[:100]}"
        
        # Quick check for exact matches
        if new_summary in existing_summaries:
            duplicates.append({
                'new_test': new_test,
                'match_type': 'exact',
                'confidence': 100
            })
            continue
        
        # Use AI for semantic similarity (simplified for performance)
        for i, existing_summary in enumerate(existing_summaries[:10]):  # Check first 10 for demo
            if len(new_summary) > 20 and len(existing_summary) > 20:
                # Simple similarity check
                common_words = set(new_summary.lower().split()) & set(existing_summary.lower().split())
                if len(common_words) > 5:
                    duplicates.append({
                        'new_test': new_test,
                        'existing_test': existing_tests[i],
                        'match_type': 'similar',
                        'confidence': len(common_words) * 10
                    })
    
    return duplicates

def import_test_suite_with_ai(uploaded_file) -> Tuple[List[Dict], Dict]:
    """
    Main function to import test suite using AI analysis
    """
    file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
    file_content = uploaded_file.read()
    
    # Extract content based on file type
    if file_extension == 'csv':
        content = file_content.decode('utf-8', errors='ignore')
        # Parse CSV
        import csv
        from io import StringIO
        reader = csv.DictReader(StringIO(content))
        test_data = list(reader)
    elif file_extension == 'json':
        content = file_content.decode('utf-8', errors='ignore')
        test_data = json.loads(content)
        if not isinstance(test_data, list):
            test_data = [test_data]
    elif file_extension in ['xlsx', 'xls']:
        # Use pandas for Excel
        df = pd.read_excel(io.BytesIO(file_content))
        test_data = df.to_dict('records')
        content = df.to_string()
    else:
        content = file_content.decode('utf-8', errors='ignore')
        # Try to parse as structured data
        test_data = [{'raw_content': content}]
    
    # Analyze schema with AI
    with st.spinner("🤖 Analyzing test suite structure with AI..."):
        schema = analyze_test_suite_schema(content, file_extension)
    
    # Generate import ID
    import_id = f"IMP_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Convert tests using AI
    converted_tests = []
    with st.spinner(f"🔄 Converting {len(test_data)} test cases..."):
        progress_bar = st.progress(0)
        for i, test in enumerate(test_data):
            converted = convert_test_with_ai(test, schema, import_id)
            converted_tests.append(converted)
            progress_bar.progress((i + 1) / len(test_data))
    
    # Check for duplicates
    with st.spinner("🔍 Checking for duplicates..."):
        duplicates = detect_duplicates(converted_tests, st.session_state.generated_tests)
    
    # Create import report
    import_report = {
        'import_id': import_id,
        'filename': uploaded_file.name,
        'timestamp': datetime.now().isoformat(),
        'schema_analysis': schema,
        'total_tests': len(test_data),
        'converted_tests': len(converted_tests),
        'duplicates_found': len(duplicates),
        'success_rate': sum(1 for t in converted_tests if t.get('nasscom_compliant', False)) / len(converted_tests) * 100 if converted_tests else 0
    }
    
    return converted_tests, import_report

def display_import_interface():
    """
    Display the AI-powered test import interface
    """
    st.markdown("""
    Import existing test suites from various formats including CSV, Excel, and JSON.
    The system provides:
    - Automatic structure and schema analysis
    - Conversion to standardized compliance format
    - Duplicate detection and management
    - Quality validation and enhancement
    """)
    
    # Sample files section
    with st.expander("📚 **Download Sample Test Files** (For Testing Import Feature)", expanded=True):
        st.info("Use these pre-formatted sample files to test the import functionality:")
        
        if st.session_state.db:
            # Load sample test files from MongoDB
            sample_files = st.session_state.db.get_all_shared_documents(limit=20)
            sample_test_files = [f for f in sample_files if f.get('metadata', {}).get('is_sample', False)]
            
            if sample_test_files:
                for sample in sample_test_files:
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.write(f"**{sample['filename']}**")
                        desc = sample.get('metadata', {}).get('description', '')
                        if desc:
                            st.caption(desc)
                        usage = sample.get('metadata', {}).get('usage', '')
                        if usage:
                            st.caption(f"💡 {usage}")
                    with col2:
                        # Get the full document with content
                        full_doc = st.session_state.db.get_shared_document(sample['_id'])
                        if full_doc and full_doc.get('content'):
                            st.download_button(
                                "📥 Download",
                                full_doc['content'],
                                file_name=sample['filename'],
                                key=f"dl_sample_{sample['_id']}",
                                help="Download this sample file"
                            )
                st.markdown("---")
                st.success("**Quick Start:** Download any sample → Review format → Upload below → See import results!")
            else:
                st.warning("Sample files not loaded yet. They will be available after running the setup script.")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose test suite file",
        type=['csv', 'xlsx', 'xls', 'json', 'xml', 'txt'],
        help="Upload CSV, Excel, JSON or other test suite formats"
    )
    
    if uploaded_file:
        # Import options
        col1, col2 = st.columns(2)
        with col1:
            merge_duplicates = st.checkbox("Merge duplicate tests", value=True)
            merge_status = "✅ Enabled" if merge_duplicates else "⭕ Disabled"
            st.caption(f"Status: {merge_status}")
        with col2:
            enhance_with_ai = st.checkbox("Enhance tests with AI", value=True)
            ai_status = "✅ Enabled" if enhance_with_ai else "⭕ Disabled"
            st.caption(f"Status: {ai_status}")
        
        if st.button("🚀 Start AI Import", type="primary"):
            # Perform import
            converted_tests, import_report = import_test_suite_with_ai(uploaded_file)
            
            # Store in session state
            st.session_state.imported_tests.extend(converted_tests)
            st.session_state.import_history.append(import_report)
            st.session_state['last_converted_tests'] = converted_tests
            st.session_state['last_import_report'] = import_report
            
            # Display results
            st.success(f"✅ Successfully imported {len(converted_tests)} test cases!")
        
        # Show results and add button OUTSIDE the import button condition
        if 'last_converted_tests' in st.session_state and st.session_state['last_converted_tests']:
            # Show import report
            if 'last_import_report' in st.session_state:
                with st.expander("📊 Import Report", expanded=True):
                    import_report = st.session_state['last_import_report']
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Tests", import_report['total_tests'])
                    with col2:
                        st.metric("Converted", import_report['converted_tests'])
                    with col3:
                        st.metric("Duplicates", import_report['duplicates_found'])
                    with col4:
                        st.metric("Success Rate", f"{import_report['success_rate']:.1f}%")
                    
                    # Schema analysis
                    st.markdown("### 📋 Detected Schema")
                    schema = import_report['schema_analysis']
                    st.json(schema.get('field_mappings', {}))
                    
                    if schema.get('recommendations'):
                        st.markdown("### 💡 Recommendations")
                        for rec in schema['recommendations']:
                            st.info(f"• {rec}")
            
            # Show converted tests preview
            converted_tests = st.session_state['last_converted_tests']
            with st.expander("👁️ Preview Imported Tests", expanded=False):
                for test in converted_tests[:5]:  # Show first 5
                    st.markdown(f"**{test['id']}**: {test['title']}")
                    st.write(f"Priority: {test['priority']} | Category: {test['category']}")
                if len(converted_tests) > 5:
                    st.info(f"... and {len(converted_tests) - 5} more tests")
            
                # Add to main test suite button - NOW OUTSIDE THE IMPORT CONDITION
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Add to Test Suite", key="add_imported_to_suite", type="primary"):
                    tests_to_add = st.session_state['last_converted_tests']
                    # Clean numpy types from imported tests
                    tests_to_add = [convert_numpy_to_python(test) for test in tests_to_add]
                    st.session_state.generated_tests.extend(tests_to_add)
                    
                    # Save to MongoDB if available
                    if st.session_state.db and st.session_state.get('user_id'):
                        with st.spinner("Saving imported tests to database..."):
                            session_id = get_or_create_session()
                            user_id = st.session_state.get('user_id')
                            success, ids = st.session_state.db.save_test_cases_batch(
                                tests_to_add, session_id, user_id
                            )
                            if success:
                                st.success(f"✅ Added {len(tests_to_add)} tests to main suite and saved to database!")
                                logger.info(f"[IMPORT] Saved {len(ids)} imported tests to MongoDB for user {user_id}")
                            else:
                                # Fallback to file save
                                saved, filename = auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                                if saved:
                                    st.success(f"✅ Added {len(tests_to_add)} tests to main suite and auto-saved!")
                    else:
                        # No MongoDB, use file save
                        saved, filename = auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                        if saved:
                            st.success(f"✅ Added {len(tests_to_add)} tests to main suite and auto-saved!")
                        else:
                            st.success(f"✅ Added {len(tests_to_add)} tests to main suite!")
                    
                    # Clean up old saves to prevent disk bloat
                    cleanup_old_saves("all_tests", keep_count=10)
                    
                    # Clear the temporary storage
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.success("✅ Tests imported successfully!")
                    st.info("Navigate to Test Suite Management to review imported test cases.")
            with col2:
                if st.button("❌ Cancel Import", key="cancel_import"):
                    del st.session_state['last_converted_tests']
                    if 'last_import_report' in st.session_state:
                        del st.session_state['last_import_report']
                    st.info("Import cancelled")
    
    # Show import history
    if st.session_state.import_history:
        st.markdown("### 📜 Import History")
        for report in reversed(st.session_state.import_history[-5:]):  # Show last 5
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                st.write(f"📄 {report['filename']}")
            with col2:
                st.write(f"Tests: {report['converted_tests']}")
            with col3:
                st.write(f"✅ {report['success_rate']:.0f}%")

def display_upload_interface():
    """
    Display document upload interface with NASSCOM compliance validation
    """
    st.markdown("""
    Select and upload project documentation to establish context for test generation.
    All documents undergo compliance validation before integration into the knowledge base.
    """)
    
    # Sample documents section
    with st.expander("📚 **Download Sample Documents** (For Testing Upload Feature)", expanded=True):
        st.info("Use these sample documents to test the document upload and compliance analysis features:")
        
        if st.session_state.db:
            # Load sample documents from MongoDB
            sample_files = st.session_state.db.get_all_shared_documents(limit=20)
            sample_docs = [f for f in sample_files if f.get('metadata', {}).get('is_document', False)]
            
            if sample_docs:
                for doc in sample_docs:
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.write(f"**{doc['filename']}**")
                        desc = doc.get('metadata', {}).get('description', '')
                        if desc:
                            st.caption(desc)
                        usage = doc.get('metadata', {}).get('usage', '')
                        if usage:
                            st.caption(f"💡 {usage}")
                    with col2:
                        # Get the full document with content
                        full_doc = st.session_state.db.get_shared_document(doc['_id'])
                        if full_doc and full_doc.get('content'):
                            st.download_button(
                                "📥 Download",
                                full_doc['content'],
                                file_name=doc['filename'],
                                key=f"dl_doc_sample_{doc['_id']}",
                                help="Download this sample document"
                            )
                st.markdown("---")
                st.success("**How to test:** Download a sample → Review its structure → Upload below → See compliance analysis → Generate test cases from it!")
            else:
                st.info("Sample documents will appear here after setup. These help you test the upload functionality.")
    
    # Show accepted formats
    with st.expander("📋 NASSCOM Document Requirements", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### ✅ Accepted Document Types")
            for doc_type, description in NASSCOM_REQUIREMENTS["document_types"].items():
                st.write(f"• **{description}**")
            
            st.markdown("### 📁 Supported Formats")
            formats = ", ".join([f"`.{fmt}`" for fmt in NASSCOM_REQUIREMENTS["document_formats"]])
            st.write(formats)
        
        with col2:
            st.markdown("### 🏥 Healthcare Standards")
            for standard in NASSCOM_REQUIREMENTS["compliance_standards"]:
                st.write(f"• {standard}")
            
            st.markdown("### 📊 Quality Criteria")
            st.write("• Minimum 100 characters")
            st.write("• Clear structure and sections")
            st.write("• Healthcare context preferred")
            st.write("• Traceable requirements")
    
    # File upload widget
    uploaded_files = st.file_uploader(
        "Choose documents to upload",
        type=NASSCOM_REQUIREMENTS["document_formats"],
        accept_multiple_files=True,
        help="Upload PRDs, User Stories, API Specs, Test Plans, etc."
    )
    
    if uploaded_files:
        st.markdown("### 📝 Document Analysis Results")
        
        for uploaded_file in uploaded_files:
            with st.expander(f"📄 {uploaded_file.name}", expanded=True):
                with st.spinner(f"Analyzing {uploaded_file.name} for NASSCOM compliance..."):
                    # Process file
                    content, compliance_report = process_uploaded_file(uploaded_file)
                    
                    # Store in session state
                    st.session_state.compliance_reports[uploaded_file.name] = compliance_report
                    
                    # Display compliance report
                    display_compliance_report(compliance_report)
                    
                    # Show content preview
                    with st.expander("📖 Content Preview", expanded=False):
                        st.text_area("", value=content[:1000] + "..." if len(content) > 1000 else content, 
                                   height=200, disabled=True, label_visibility="collapsed")
                    
                    # Action buttons
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if compliance_report['is_compliant']:
                            if st.button(f"✅ Add to Knowledge Base", key=f"add_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.success(f"Added {uploaded_file.name} to knowledge base!")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    # Clear cache to reload with new documents
                                    st.cache_resource.clear()
                                    st.rerun()
                        else:
                            st.warning("Document needs improvements before adding to knowledge base")
                    
                    with col2:
                        # Download compliance report
                        report_json = json.dumps(compliance_report, indent=2)
                        st.download_button(
                            label="📥 Download Compliance Report",
                            data=report_json,
                            file_name=f"{uploaded_file.name}_compliance.json",
                            mime="application/json",
                            key=f"download_{uploaded_file.name}"
                        )
                    
                    with col3:
                        if not compliance_report['is_compliant']:
                            if st.button(f"🔄 Override & Add Anyway", key=f"override_{uploaded_file.name}"):
                                if save_uploaded_document(uploaded_file.name, content, compliance_report):
                                    st.warning(f"Added {uploaded_file.name} despite compliance issues")
                                    st.session_state.uploaded_docs.append(uploaded_file.name)
                                    st.cache_resource.clear()
                                    st.rerun()

def refine_test_case_with_feedback(test_case: Dict, feedback: str, specific_field: str = None) -> Dict:
    """
    Refine an existing test case based on human feedback.
    """
    test_case_for_prompt = {k: v for k, v in test_case.items() if k != 'retrieved_context'}
    
    if specific_field:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK FOR '{specific_field}':
{feedback}

NASSCOM REQUIREMENTS:
- Maintain all required fields
- Ensure healthcare compliance where applicable
- Include both positive and negative scenarios
- Maintain traceability

Generate an improved version of this test case, specifically improving the '{specific_field}' field based on the feedback.
Maintain all other fields unless they need adjustment to be consistent with the change.

Return ONLY a valid JSON object with the same structure."""
    else:
        prompt = f"""You are a QA Engineer following NASSCOM guidelines. Refine the following test case based on user feedback.

CURRENT TEST CASE:
{json.dumps(test_case_for_prompt, indent=2, default=str)}

USER FEEDBACK:
{feedback}

Generate an improved version of this test case incorporating the feedback while maintaining NASSCOM compliance.

Return ONLY a valid JSON object."""
    
    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3,
                top_p=0.9,
                max_output_tokens=4096
            )
        )
        
        response = model.generate_content(prompt)
        refined_test = json.loads(response.text)
        
        # Preserve metadata
        refined_test['id'] = test_case['id']
        refined_test['refinement_timestamp'] = datetime.now().isoformat()
        refined_test['refinement_feedback'] = feedback
        refined_test['version'] = test_case.get('version', 1) + 1
        refined_test['nasscom_compliant'] = True
        
        if 'retrieved_context' in test_case:
            refined_test['retrieved_context'] = test_case['retrieved_context']
        
        # Clean numpy types before returning
        refined_test = convert_numpy_to_python(refined_test)
        
        return refined_test
        
    except Exception as e:
        st.error(f"Error refining test case: {e}")
        return test_case

def display_test_case(test_case: Dict, context_docs: List[Dict], key_suffix: str = ""):
    """Display generated test case in a formatted way
    
    Args:
        test_case: The test case dictionary to display
        context_docs: Context documents used for generation
        key_suffix: Suffix to add to button keys to avoid duplicates
    """
    
    # Header with metrics
    st.markdown("---")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Priority", test_case.get('priority', 'N/A'))
    with col2:
        st.metric("Category", test_case.get('category', 'N/A'))
    with col3:
        compliance = test_case.get('compliance', [])
        st.metric("Compliance", len(compliance) if compliance else 'None')
    with col4:
        st.metric("Duration", test_case.get('estimated_duration', 'N/A'))
    
    # NASSCOM compliance indicator
    if test_case.get('nasscom_compliant', False):
        st.success("✅ NASSCOM Compliant Test Case")
    else:
        st.warning("⚠️ Non-compliant test case (fallback generated)")
    
    # Main content
    st.markdown("### 📝 Description")
    st.write(test_case.get('description', 'N/A'))
    
    # Traceability
    if test_case.get('traceability'):
        st.markdown("### 🔗 Traceability")
        st.info(test_case.get('traceability', 'N/A'))
    
    # Preconditions
    st.markdown("### ⚙️ Preconditions")
    st.info(test_case.get('preconditions', 'None specified'))
    
    # Test steps
    st.markdown("### 📋 Test Steps")
    steps = test_case.get('test_steps', [])
    for i, step in enumerate(steps, 1):
        st.write(f"{i}. {step}")
    
    # Expected results
    st.markdown("### ✅ Expected Results")
    st.success(test_case.get('expected_results', 'N/A'))
    
    # Test data
    if test_case.get('test_data'):
        st.markdown("### 📊 Test Data")
        st.json(test_case.get('test_data'))
    
    # Edge cases and negative tests
    col1, col2 = st.columns(2)
    with col1:
        if test_case.get('edge_cases'):
            st.markdown("### ⚠️ Edge Cases")
            for edge in test_case['edge_cases']:
                st.write(f"• {edge}")
    
    with col2:
        if test_case.get('negative_tests'):
            st.markdown("### 🚫 Negative Tests")
            for neg_test in test_case['negative_tests']:
                st.write(f"• {neg_test}")
    
    # Metadata
    with st.expander("🔍 Generation Details"):
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Automation Feasible:** {'Yes' if test_case.get('automation_feasible') else 'No'}")
            st.write(f"**Context Relevance:** {test_case.get('avg_context_relevance', 'N/A')}")
            if test_case.get('compliance'):
                st.write(f"**Standards:** {', '.join(test_case['compliance'])}")
        with col2:
            st.write(f"**Generated From:** {test_case.get('generated_from', 'N/A')[:50]}...")
            st.write(f"**Timestamp:** {test_case.get('generation_timestamp', 'N/A')[:19]}")
            st.write(f"**Version:** {test_case.get('version', 1)}")
    
    # Edit and Refinement Options
    st.markdown("#### 🛠️ Test Case Actions")
    action_col1, action_col2, action_col3 = st.columns(3)
    
    with action_col1:
        if st.button(f"✏️ Edit Manually", key=f"edit_{test_case['id']}{key_suffix}", 
                     help="Manually edit this test case"):
            st.session_state[f'editing_{test_case["id"]}'] = True
            st.rerun()
    
    with action_col2:
        if st.button(f"🤖 Refine with AI", key=f"refine_{test_case['id']}{key_suffix}",
                     help="Use AI to improve this test case"):
            st.session_state[f'refining_{test_case["id"]}'] = True
            st.rerun()
    
    with action_col3:
        if st.button(f"📋 Clone Test", key=f"clone_{test_case['id']}{key_suffix}",
                     help="Create a copy of this test case"):
            # Clone the test case
            cloned_test = test_case.copy()
            cloned_test['id'] = f"TC_{uuid.uuid4().hex[:8].upper()}"
            cloned_test['title'] = f"{cloned_test.get('title', 'Untitled')} (Copy)"
            cloned_test['version'] = 1
            cloned_test['generation_timestamp'] = datetime.now().isoformat()
            
            # Clean numpy types from cloned test
            cloned_test = convert_numpy_to_python(cloned_test)
            st.session_state.generated_tests.append(cloned_test)
            
            # Save cloned test to MongoDB
            if st.session_state.db and st.session_state.get('user_id'):
                if save_test_to_mongodb(cloned_test):
                    logger.info(f"[CLONE] Test case {cloned_test['id']} cloned and saved to MongoDB")
            
            st.success(f"✅ Test case cloned: {cloned_test['id']}")
            time.sleep(1)
            st.rerun()
    
    # Manual Edit Form
    if st.session_state.get(f'editing_{test_case["id"]}', False):
        st.markdown("---")
        st.markdown("### ✏️ Edit Test Case")
        
        with st.form(f"edit_form_{test_case['id']}{key_suffix}"):
            # Editable fields
            new_title = st.text_input("Title", value=test_case.get('title', ''))
            new_description = st.text_area("Description", value=test_case.get('description', ''), height=100)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                new_priority = st.selectbox("Priority", 
                                           ["Critical", "High", "Medium", "Low"],
                                           index=["Critical", "High", "Medium", "Low"].index(test_case.get('priority', 'Medium')))
            with col2:
                new_category = st.selectbox("Category",
                                           ["Functional", "Security", "Performance", "Usability", "Integration", "Compliance"],
                                           index=["Functional", "Security", "Performance", "Usability", "Integration", "Compliance"].index(test_case.get('category', 'Functional')) if test_case.get('category', 'Functional') in ["Functional", "Security", "Performance", "Usability", "Integration", "Compliance"] else 0)
            with col3:
                new_duration = st.text_input("Duration", value=test_case.get('estimated_duration', '30 minutes'))
            
            new_preconditions = st.text_area("Preconditions", 
                                            value=test_case.get('preconditions', ''), 
                                            height=80)
            
            new_steps = st.text_area("Test Steps (one per line)", 
                                    value='\n'.join(test_case.get('test_steps', [])), 
                                    height=150)
            
            new_expected = st.text_area("Expected Results", 
                                       value=test_case.get('expected_results', ''), 
                                       height=100)
            
            # Compliance standards
            compliance_options = ["HIPAA", "GDPR", "FDA", "ISO-27001", "SOC2", "HITRUST"]
            current_compliance = test_case.get('compliance', [])
            new_compliance = st.multiselect("Compliance Standards", 
                                           compliance_options,
                                           default=[c for c in current_compliance if c in compliance_options])
            
            col1, col2 = st.columns(2)
            with col1:
                if st.form_submit_button("💾 Save Changes", type="primary"):
                    # Update the test case
                    test_case['title'] = new_title
                    test_case['description'] = new_description
                    test_case['priority'] = new_priority
                    test_case['category'] = new_category
                    test_case['estimated_duration'] = new_duration
                    test_case['preconditions'] = new_preconditions
                    test_case['test_steps'] = [step.strip() for step in new_steps.split('\n') if step.strip()]
                    test_case['expected_results'] = new_expected
                    test_case['compliance'] = new_compliance
                    test_case['version'] = test_case.get('version', 1) + 1
                    test_case['last_modified'] = datetime.now().isoformat()
                    test_case['manually_edited'] = True
                    
                    # Update in the session state
                    for i, tc in enumerate(st.session_state.generated_tests):
                        if tc['id'] == test_case['id']:
                            st.session_state.generated_tests[i] = test_case
                            break
                    
                    # Save to MongoDB if available
                    if st.session_state.db and st.session_state.get('user_id'):
                        save_test_to_mongodb(test_case)
                    
                    st.success("✅ Test case updated successfully!")
                    st.session_state[f'editing_{test_case["id"]}'] = False
                    time.sleep(1)
                    st.rerun()
            
            with col2:
                if st.form_submit_button("❌ Cancel"):
                    st.session_state[f'editing_{test_case["id"]}'] = False
                    st.rerun()
    
    # AI Refinement Form
    if st.session_state.get(f'refining_{test_case["id"]}', False):
        st.markdown("---")
        st.markdown("### 🤖 AI-Powered Refinement")
        
        refinement_type = st.radio("Refinement Type",
                                   ["General Improvement", "Specific Field"],
                                   key=f"refine_type_{test_case['id']}{key_suffix}")
        
        if refinement_type == "Specific Field":
            field_to_refine = st.selectbox("Field to Refine",
                                          ["description", "test_steps", "expected_results", 
                                           "preconditions", "edge_cases", "negative_tests"],
                                          key=f"field_{test_case['id']}{key_suffix}")
        else:
            field_to_refine = None
        
        feedback = st.text_area(
            "Provide your feedback or requirements for improvement:",
            placeholder="Example: Make the test steps more detailed, add security validation checks, include performance benchmarks...",
            height=100,
            key=f"feedback_{test_case['id']}{key_suffix}"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔧 Refine Test Case", type="primary", 
                        key=f"do_refine_{test_case['id']}{key_suffix}",
                        disabled=not feedback):
                if feedback:
                    with st.spinner("🤖 AI is refining your test case..."):
                        refined_test = refine_test_case_with_feedback(test_case, feedback, field_to_refine)
                        
                        # Update in session state
                        for i, tc in enumerate(st.session_state.generated_tests):
                            if tc['id'] == test_case['id']:
                                st.session_state.generated_tests[i] = refined_test
                                break
                        
                        # Save to MongoDB if available
                        if st.session_state.db and st.session_state.get('user_id'):
                            save_test_to_mongodb(refined_test)
                        
                        st.success(f"✅ Test case refined successfully! (Version {refined_test.get('version', 2)})")
                        st.session_state[f'refining_{test_case["id"]}'] = False
                        time.sleep(1.5)
                        st.rerun()
        
        with col2:
            if st.button("❌ Cancel", key=f"cancel_refine_{test_case['id']}{key_suffix}"):
                st.session_state[f'refining_{test_case["id"]}'] = False
                st.rerun()
    
    # Save options with download buttons
    st.markdown("#### 💾 Save & Download Options")
    col1, col2, col3, col4 = st.columns(4)
    
    # Track save status in session state
    save_json_key = f"saved_json_{test_case['id']}{key_suffix}"
    save_csv_key = f"saved_csv_{test_case['id']}{key_suffix}"
    
    with col1:
        if st.button(f"📁 Save JSON", key=f"save_json_{test_case['id']}{key_suffix}", help="Save to data folder"):
            try:
                filename = save_test_case(test_case, 'json')
                st.session_state[save_json_key] = filename
                st.success(f"✅ Saved: {filename}")
            except Exception as e:
                st.error(f"Failed to save: {str(e)}")
        elif save_json_key in st.session_state:
            st.info(f"✓ Saved: {st.session_state[save_json_key]}")
    
    with col2:
        # Download JSON
        json_data = json.dumps(test_case, indent=2, default=str)
        st.download_button(
            label="⬇️ Download JSON",
            data=json_data,
            file_name=f"test_{test_case['id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            key=f"download_json_{test_case['id']}{key_suffix}"
        )
    
    with col3:
        if st.button(f"📁 Save CSV", key=f"save_csv_{test_case['id']}{key_suffix}", help="Save to data folder"):
            try:
                filename = save_test_case(test_case, 'csv')
                st.session_state[save_csv_key] = filename
                st.success(f"✅ Saved: {filename}")
            except Exception as e:
                st.error(f"Failed to save: {str(e)}")
        elif save_csv_key in st.session_state:
            st.info(f"✓ Saved: {st.session_state[save_csv_key]}")
    
    with col4:
        # Download CSV
        df = pd.DataFrame([test_case])
        csv_data = df.to_csv(index=False)
        st.download_button(
            label="⬇️ Download CSV",
            data=csv_data,
            file_name=f"test_{test_case['id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            key=f"download_csv_{test_case['id']}{key_suffix}"
        )

# ===============================
# DevOps Export Functions (NEW IN V7)
# ===============================

def convert_to_jira_format(test_cases: List[Dict]) -> pd.DataFrame:
    """Convert test cases to Jira-compatible CSV format"""
    jira_data = []
    for tc in test_cases:
        jira_data.append({
            'Issue Type': 'Test',
            'Summary': tc.get('title', 'Untitled Test'),
            'Description': tc.get('description', ''),
            'Priority': tc.get('priority', 'Medium'),
            'Labels': ', '.join(tc.get('compliance', [])),
            'Test Type': 'Manual' if not tc.get('automation_feasible', False) else 'Automated',
            'Acceptance Criteria': tc.get('expected_results', ''),
            'Test Steps': '\n'.join([f"{i+1}. {step}" for i, step in enumerate(tc.get('test_steps', []))]),
            'Test Data': json.dumps(tc.get('test_data', {})),
            'Component': tc.get('category', 'Functional'),
            'Fix Version': tc.get('version', '1.0'),
            'Story Points': '3' if tc.get('priority') == 'High' else '2' if tc.get('priority') == 'Medium' else '1',
            'Epic Link': tc.get('traceability', ''),
            'Custom Field (Compliance)': ', '.join(tc.get('compliance', [])),
            'Custom Field (Duration)': tc.get('estimated_duration', ''),
        })
    return pd.DataFrame(jira_data)

def convert_to_azure_devops_format(test_cases: List[Dict]) -> pd.DataFrame:
    """Convert test cases to Azure DevOps Test Plans format"""
    azure_data = []
    for tc in test_cases:
        azure_data.append({
            'Work Item Type': 'Test Case',
            'Title': tc.get('title', 'Untitled Test'),
            'State': 'Design',
            'Priority': {'Critical': 1, 'High': 2, 'Medium': 3, 'Low': 4}.get(tc.get('priority', 'Medium'), 3),
            'Assigned To': '',
            'Area Path': f"Healthcare\\{tc.get('category', 'General')}",
            'Iteration Path': '',
            'Description': tc.get('description', ''),
            'Steps': json.dumps([{'action': step, 'expected': tc.get('expected_results', '')} 
                                for step in tc.get('test_steps', [])]),
            'Precondition': tc.get('preconditions', ''),
            'Postcondition': tc.get('expected_results', ''),
            'Tags': '; '.join(tc.get('compliance', [])),
            'Automation Status': 'Automated' if tc.get('automation_feasible', False) else 'Not Automated',
            'Test Suite': tc.get('category', 'General'),
            'Parameters': json.dumps(tc.get('test_data', {}))
        })
    return pd.DataFrame(azure_data)

def convert_to_postman_format(test_cases: List[Dict]) -> Dict:
    """Convert test cases to Postman Collection format"""
    postman_collection = {
        "info": {
            "name": "Healthcare Test Collection",
            "description": "Generated test cases for Healthcare/MedTech system",
            "schema": "https://schema.getpostman.com/json/collection/v2.1.0/collection.json"
        },
        "item": []
    }
    
    for tc in test_cases:
        # Only include API/Integration tests
        if tc.get('category') in ['Integration', 'API', 'Performance']:
            item = {
                "name": tc.get('title', 'Untitled Test'),
                "request": {
                    "method": "GET",  # Default, should be customized
                    "header": [],
                    "description": tc.get('description', ''),
                    "url": {
                        "raw": "{{base_url}}/api/endpoint",
                        "host": ["{{base_url}}"],
                        "path": ["api", "endpoint"]
                    }
                },
                "event": [
                    {
                        "listen": "test",
                        "script": {
                            "exec": [
                                f"// Test: {tc.get('title', '')}",
                                f"// Expected: {tc.get('expected_results', '')}",
                                "pm.test('Status code is 200', function () {",
                                "    pm.response.to.have.status(200);",
                                "});",
                                "",
                                "pm.test('Response time is less than 2000ms', function () {",
                                "    pm.expect(pm.response.responseTime).to.be.below(2000);",
                                "});"
                            ]
                        }
                    }
                ],
                "response": []
            }
            
            # Add test data as variables
            if tc.get('test_data'):
                item['event'].append({
                    "listen": "prerequest",
                    "script": {
                        "exec": [
                            "// Set test data",
                            *[f"pm.variables.set('{k}', '{v}');" 
                              for k, v in tc.get('test_data', {}).items()]
                        ]
                    }
                })
            
            postman_collection["item"].append(item)
    
    return postman_collection

def convert_to_junit_xml(test_cases: List[Dict]) -> str:
    """Convert test cases to JUnit XML format"""
    xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n'
    xml_content += f'<testsuites name="Healthcare Test Suite" tests="{len(test_cases)}">\n'
    xml_content += f'  <testsuite name="Automated Tests" tests="{len(test_cases)}">\n'
    
    for tc in test_cases:
        classname = tc.get('category', 'General').replace(' ', '')
        testname = tc.get('title', 'Untitled').replace(' ', '_')
        
        xml_content += f'    <testcase classname="{classname}" name="{testname}" time="{tc.get("estimated_duration", "0")}">\n'
        
        # Add test steps as system-out
        xml_content += '      <system-out><![CDATA[\n'
        xml_content += f'        Description: {tc.get("description", "")}\n'
        xml_content += f'        Priority: {tc.get("priority", "Medium")}\n'
        xml_content += '        Steps:\n'
        for i, step in enumerate(tc.get('test_steps', []), 1):
            xml_content += f'          {i}. {step}\n'
        xml_content += f'        Expected: {tc.get("expected_results", "")}\n'
        xml_content += '      ]]></system-out>\n'
        
        # Add properties for metadata
        xml_content += '      <properties>\n'
        xml_content += f'        <property name="priority" value="{tc.get("priority", "Medium")}"/>\n'
        xml_content += f'        <property name="automated" value="{tc.get("automation_feasible", False)}"/>\n'
        for compliance in tc.get('compliance', []):
            xml_content += f'        <property name="compliance" value="{compliance}"/>\n'
        xml_content += '      </properties>\n'
        
        xml_content += '    </testcase>\n'
    
    xml_content += '  </testsuite>\n'
    xml_content += '</testsuites>'
    
    return xml_content

def convert_to_testrail_format(test_cases: List[Dict]) -> List[Dict]:
    """Convert test cases to TestRail JSON format"""
    testrail_cases = []
    
    for tc in test_cases:
        testrail_case = {
            "title": tc.get('title', 'Untitled Test'),
            "type_id": 1,  # Manual test type
            "priority_id": {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1}.get(tc.get('priority', 'Medium'), 2),
            "estimate": tc.get('estimated_duration', ''),
            "refs": tc.get('traceability', ''),
            "custom_automation_type": 0 if not tc.get('automation_feasible', False) else 1,
            "custom_preconds": tc.get('preconditions', ''),
            "custom_steps_separated": [
                {
                    "content": step,
                    "expected": tc.get('expected_results', '') if i == len(tc.get('test_steps', [])) - 1 else ''
                }
                for i, step in enumerate(tc.get('test_steps', []))
            ],
            "custom_expected": tc.get('expected_results', ''),
            "custom_test_data": json.dumps(tc.get('test_data', {})),
            "custom_tags": ', '.join(tc.get('compliance', [])),
            "section_id": 1  # Default section
        }
        testrail_cases.append(testrail_case)
    
    return testrail_cases

def customize_export_with_ai(test_cases: List[Dict], target_format: str, custom_requirements: str) -> str:
    """Use AI to customize the export format based on user requirements"""
    
    # Sample of test data for AI
    sample_test = test_cases[0] if test_cases else {}
    
    prompt = f"""You are an expert in test management tools and DevOps integration.
    
I need to export test cases to {target_format} format with these custom requirements:
{custom_requirements}

Sample test case structure:
{json.dumps(sample_test, indent=2)[:1000]}

Current standard fields I'm using:
- Jira: Issue Type, Summary, Description, Priority, Labels, Test Steps, etc.
- Azure DevOps: Work Item Type, Title, State, Priority, Steps, Tags, etc.
- Postman: Collection with requests, tests, variables
- JUnit: XML with testsuites, testcases, properties

Based on the custom requirements, provide:
1. Modified field mappings
2. Additional fields to include
3. Field transformations needed
4. Any format-specific considerations

Return as JSON with structure:
{{
    "field_mappings": {{
        "our_field": "target_field",
        ...
    }},
    "additional_fields": [
        {{"name": "field_name", "value": "how to derive it"}}
    ],
    "transformations": [
        {{"field": "field_name", "transformation": "description"}}
    ],
    "format_notes": "specific considerations for this format"
}}
"""

    try:
        model = genai.GenerativeModel(
            GEMINI_MODEL_NAME,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.3
            )
        )
        
        response = model.generate_content(prompt)
        customization = json.loads(response.text)
        return customization
        
    except Exception as e:
        st.error(f"AI customization error: {e}")
        return {
            "field_mappings": {},
            "additional_fields": [],
            "transformations": [],
            "format_notes": "Using default format"
        }

def display_export_interface():
    """Display the DevOps export interface with preview and customization"""
    st.markdown("""
    Export test cases to industry-standard formats for seamless integration with existing 
    test management and DevOps platforms. Preview output formats and customize field mappings as needed.
    """)
    
    # Ensure generated_tests exists
    if 'generated_tests' not in st.session_state:
        st.session_state.generated_tests = []
    
    if not st.session_state.generated_tests:
        st.warning("⚠️ No test cases available to export. Generate or import tests first.")
        return
    
    # Export format selection
    col1, col2 = st.columns([2, 1])
    
    with col1:
        export_format = st.selectbox(
            "Select Export Format",
            ["Jira CSV", "Azure DevOps", "Postman Collection", "JUnit XML", "TestRail JSON", "Custom Format"],
            help="Choose the target platform for export"
        )
    
    with col2:
        tests_to_export = st.selectbox(
            "Tests to Export",
            ["All Tests", "Filtered Tests", "Selected Tests"],
            help="Choose which tests to export"
        )
    
    # Filter options
    if tests_to_export == "Filtered Tests":
        col1, col2, col3 = st.columns(3)
        with col1:
            filter_category = st.multiselect(
                "Categories",
                list(set(tc.get('category', 'Unknown') for tc in st.session_state.generated_tests))
            )
        with col2:
            filter_priority = st.multiselect(
                "Priorities",
                list(set(tc.get('priority', 'Unknown') for tc in st.session_state.generated_tests))
            )
        with col3:
            filter_automated = st.checkbox("Only Automated Tests")
            auto_status = "✅ Filtering" if filter_automated else "⭕ All Tests"
            st.caption(f"Status: {auto_status}")
        
        # Apply filters
        export_tests = st.session_state.generated_tests
        if filter_category:
            export_tests = [tc for tc in export_tests if tc.get('category') in filter_category]
        if filter_priority:
            export_tests = [tc for tc in export_tests if tc.get('priority') in filter_priority]
        if filter_automated:
            export_tests = [tc for tc in export_tests if tc.get('automation_feasible', False)]
    else:
        export_tests = st.session_state.generated_tests
    
    st.info(f"📊 {len(export_tests)} test cases will be exported")
    
    # Customization section
    with st.expander("⚙️ Customize Export Format", expanded=False):
        st.markdown("### Field Mapping Customization")
        
        custom_requirements = st.text_area(
            "Describe your customization needs",
            placeholder="Example: Add custom fields for sprint number, test environment, risk level. Map our 'compliance' field to Jira's 'Regulatory' field.",
            height=100
        )
        
        if st.button("🤖 Get AI Suggestions", key="ai_customize"):
            if custom_requirements:
                with st.spinner("AI is analyzing your requirements..."):
                    customization = customize_export_with_ai(export_tests, export_format, custom_requirements)
                    st.session_state['export_customization'] = customization
                    
                    st.success("AI suggestions ready!")
                    st.json(customization)
    
    # Preview section
    with st.expander("👁️ Preview Export Format", expanded=True):
        st.markdown(f"### Preview: {export_format}")
        
        # Generate preview based on format
        preview_data = None
        
        if export_format == "Jira CSV":
            preview_data = convert_to_jira_format(export_tests[:3])  # Preview first 3
            st.dataframe(preview_data, use_container_width=True)
            
        elif export_format == "Azure DevOps":
            preview_data = convert_to_azure_devops_format(export_tests[:3])
            st.dataframe(preview_data, use_container_width=True)
            
        elif export_format == "Postman Collection":
            preview_data = convert_to_postman_format(export_tests[:3])
            st.json(preview_data)
            
        elif export_format == "JUnit XML":
            preview_data = convert_to_junit_xml(export_tests[:3])
            st.code(preview_data[:1500], language='xml')  # Show first 1500 chars
            
        elif export_format == "TestRail JSON":
            preview_data = convert_to_testrail_format(export_tests[:3])
            st.json(preview_data[0] if preview_data else {})  # Show first test
        
        if len(export_tests) > 3:
            st.info(f"... and {len(export_tests) - 3} more tests")
    
    # Export buttons
    st.markdown("### 📥 Export Options")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Download Export", type="primary", key="download_export"):
            with st.spinner(f"Generating {export_format} export..."):
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                
                if export_format == "Jira CSV":
                    df = convert_to_jira_format(export_tests)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "Download Jira CSV",
                        csv,
                        f"jira_tests_{timestamp}.csv",
                        "text/csv"
                    )
                    
                elif export_format == "Azure DevOps":
                    df = convert_to_azure_devops_format(export_tests)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "Download Azure DevOps CSV",
                        csv,
                        f"azure_tests_{timestamp}.csv",
                        "text/csv"
                    )
                    
                elif export_format == "Postman Collection":
                    collection = convert_to_postman_format(export_tests)
                    json_str = json.dumps(collection, indent=2)
                    st.download_button(
                        "Download Postman Collection",
                        json_str,
                        f"postman_collection_{timestamp}.json",
                        "application/json"
                    )
                    
                elif export_format == "JUnit XML":
                    xml_content = convert_to_junit_xml(export_tests)
                    st.download_button(
                        "Download JUnit XML",
                        xml_content,
                        f"junit_tests_{timestamp}.xml",
                        "text/xml"
                    )
                    
                elif export_format == "TestRail JSON":
                    testrail_data = convert_to_testrail_format(export_tests)
                    json_str = json.dumps(testrail_data, indent=2)
                    st.download_button(
                        "Download TestRail JSON",
                        json_str,
                        f"testrail_tests_{timestamp}.json",
                        "application/json"
                    )
    
    with col2:
        if st.button("📋 Copy to Clipboard", key="copy_export"):
            st.info("Export copied! (Feature requires JavaScript integration)")
    
    with col3:
        if st.button("🔗 Send to Integration", key="send_integration"):
            st.info(f"Ready to integrate with {export_format.split()[0]} API")
    
    # Integration guide
    with st.expander("📚 Integration Guide", expanded=False):
        st.markdown(f"""
        ### How to import into {export_format.split()[0]}
        
        **Jira:**
        1. Go to Issues > Import Issues from CSV
        2. Upload the downloaded CSV file
        3. Map fields as needed
        4. Complete the import
        
        **Azure DevOps:**
        1. Open Test Plans
        2. Click Import Test Cases
        3. Upload the CSV file
        4. Map to your test plan
        
        **Postman:**
        1. Open Postman
        2. Click Import
        3. Select the JSON file
        4. Tests will appear in Collections
        
        **JUnit:**
        1. Place XML in your test results directory
        2. Configure CI/CD to read JUnit format
        3. Results will appear in test reports
        
        **TestRail:**
        1. Use TestRail API or Import tool
        2. Upload JSON via API endpoint
        3. Map to test suite
        """)

def save_test_case(test_case: Dict, format: str = 'json'):
    """Save test case to file with proper directory creation"""
    # Ensure directory exists
    os.makedirs(TEST_CASES_DIR, exist_ok=True)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if format == 'json':
        filename = f"{TEST_CASES_DIR}/manual_save_{test_case['id']}_{timestamp}.json"
        with open(filename, 'w') as f:
            json.dump(test_case, f, indent=2, default=str)
        return os.path.basename(filename)  # Return just filename for cleaner display
    elif format == 'csv':
        filename = f"{TEST_CASES_DIR}/manual_save_{test_case['id']}_{timestamp}.csv"
        df = pd.DataFrame([test_case])
        df.to_csv(filename, index=False)
        return os.path.basename(filename)  # Return just filename for cleaner display
    return None

# LOGIN/SIGNUP PAGES
def display_api_test_executor():
    """Display the API test execution interface"""
    
    # Initialize test executor in session state
    if 'api_executor' not in st.session_state:
        st.session_state.api_executor = None
    
    # Get API tests from generated tests - look for API-related keywords
    api_tests = []
    for tc in st.session_state.generated_tests:
        # Check category
        if tc.get('category') in ['API', 'Integration', 'Performance']:
            api_tests.append(tc)
        # Also check if test_data has API-specific fields
        elif tc.get('test_data', {}).get('endpoint') or tc.get('test_data', {}).get('method'):
            api_tests.append(tc)
        # Check title/description for API keywords
        elif any(keyword in str(tc.get('title', '')).lower() + str(tc.get('description', '')).lower() 
                for keyword in ['api', 'endpoint', 'rest', 'http', 'request', 'response', 'webhook', 'graphql']):
            api_tests.append(tc)
    
    if not api_tests:
        st.warning("No API tests available. Generate API/Integration tests first!")
        st.info("💡 Try generating tests with prompts like:")
        st.code("""
• "Generate API tests for appointment booking endpoint"
• "Create REST API tests for patient data retrieval"
• "Test API authentication with OAuth 2.0"
• "Generate API tests for the /appointments endpoint from the API spec"
        """)
        return
    
    # Configuration Section
    st.subheader("⚙️ API Test Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # API Endpoint Configuration
        st.markdown("**API Endpoint Settings**")
        
        # Replace checkbox with text status display
        use_mock = st.checkbox(
            "🧪 Use Mock API Server",
            value=True,
            help="Use built-in mock server for demonstration"
        )
        
        # Display text status for better visibility
        mock_status = "✅ **ON** - Mock Server Active" if use_mock else "❌ **OFF** - Using Custom API"
        st.markdown(f"Status: {mock_status}")
        
        if use_mock:
            base_url = "http://localhost:8000/mock"
            st.info("Using mock Healthcare API for demonstration")
        else:
            base_url = st.text_input(
                "Base URL",
                value="https://api.healthconnect.com/v2",
                help="Enter your API base URL"
            )
        
        # Slider with text display
        st.markdown("**Request Timeout**")
        col_slider, col_value = st.columns([3, 1])
        with col_slider:
            timeout = st.slider(
                "",  # Empty label
                min_value=5,
                max_value=60,
                value=30,
                key="timeout_slider"
            )
        with col_value:
            st.markdown(f"### `{timeout}s`")
    
    with col2:
        # Authentication Configuration
        st.markdown("**Authentication**")
        
        auth_type = st.selectbox(
            "Authentication Type",
            ["None", "Bearer Token", "API Key", "Basic Auth", "OAuth 2.0"]
        )
        
        credentials = {}
        if auth_type == "Bearer Token":
            token = st.text_input("Bearer Token", type="password")
            credentials = {'token': token}
        elif auth_type == "API Key":
            api_key = st.text_input("API Key", type="password")
            header_name = st.text_input("Header Name", value="X-API-Key")
            credentials = {'key': api_key, 'header_name': header_name}
        elif auth_type == "Basic Auth":
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            credentials = {'username': username, 'password': password}
        elif auth_type == "OAuth 2.0":
            access_token = st.text_input("Access Token", type="password")
            credentials = {'access_token': access_token}
    
    st.markdown("---")
    
    # Test Selection
    st.subheader("📋 Select Tests to Execute")
    
    # Filter options
    col1, col2, col3 = st.columns(3)
    
    with col1:
        filter_category = st.selectbox(
            "Category Filter",
            ["All"] + list(set(tc.get('category', 'API') for tc in api_tests))
        )
    
    with col2:
        filter_priority = st.selectbox(
            "Priority Filter",
            ["All"] + list(set(tc.get('priority', 'Medium') for tc in api_tests))
        )
    
    with col3:
        max_tests = st.number_input(
            "Max Tests to Run",
            min_value=1,
            max_value=len(api_tests),
            value=min(5, len(api_tests))
        )
    
    # Filter tests
    filtered_tests = api_tests
    if filter_category != "All":
        filtered_tests = [tc for tc in filtered_tests if tc.get('category') == filter_category]
    if filter_priority != "All":
        filtered_tests = [tc for tc in filtered_tests if tc.get('priority') == filter_priority]
    
    filtered_tests = filtered_tests[:max_tests]
    
    # Display test selection with API details
    st.info(f"Found {len(filtered_tests)} tests matching criteria")
    
    selected_tests = []
    for tc in filtered_tests:
        test_data = tc.get('test_data', {})
        method = test_data.get('method', 'GET')
        endpoint = test_data.get('endpoint', '/api/endpoint')
        
        # Create more informative checkbox label with text status
        label = f"**{tc['id']}** [{method}] {endpoint} - {tc.get('title', 'Untitled')[:60]}"
        
        col_check, col_status = st.columns([4, 1])
        with col_check:
            is_selected = st.checkbox(
                label,
                value=True,
                key=f"exec_{tc['id']}",
                help=f"Category: {tc.get('category', 'API')} | Priority: {tc.get('priority', 'Medium')}"
            )
        with col_status:
            status_text = "✅" if is_selected else "⭕"
            st.markdown(status_text)
        
        if is_selected:
            selected_tests.append(tc)
    
    st.markdown("---")
    
    # Execution Section
    st.subheader("🎯 Test Execution")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        execute_btn = st.button(
            "▶️ Execute Tests",
            type="primary",
            disabled=len(selected_tests) == 0
        )
    
    with col2:
        if st.button("🔄 Clear Results"):
            if 'test_execution_results' in st.session_state:
                del st.session_state['test_execution_results']
            st.rerun()
    
    with col3:
        st.metric("Tests Selected", len(selected_tests))
    
    # Execute tests
    if execute_btn and selected_tests:
        # Initialize executor
        if use_mock:
            # For mock mode, we'll simulate API responses
            executor = APITestExecutor(base_url="http://mock.healthcare.api")
        else:
            executor = APITestExecutor(base_url=base_url, timeout=timeout)
        
        # Set authentication if provided
        if auth_type != "None" and any(credentials.values()):
            executor.set_authentication(auth_type.lower().replace(' ', '_'), credentials)
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Execute tests
        results = []
        for i, test in enumerate(selected_tests):
            status_text.text(f"Executing: {test.get('title', 'Test')}...")
            progress_bar.progress((i + 1) / len(selected_tests))
            
            # For mock mode, simulate the response
            if use_mock:
                # Simulate API execution
                time.sleep(0.5)  # Simulate network delay
                
                # Create mock result
                result = {
                    'test_id': test['id'],
                    'test_title': test.get('title', 'Unknown'),
                    'status': 'passed' if np.random.random() > 0.2 else 'failed',
                    'request': {
                        'method': test.get('test_data', {}).get('method', 'GET'),
                        'url': f"{base_url}{test.get('test_data', {}).get('endpoint', '/mock/health')}",
                        'headers': {'Authorization': 'Bearer mock...oken'}
                    },
                    'response': {
                        'status_code': 200 if np.random.random() > 0.3 else 404,
                        'response_time_ms': np.random.randint(50, 500),
                        'body': {'status': 'success', 'data': 'mock_response'}
                    },
                    'assertions': [
                        {
                            'type': 'status_code',
                            'expected': 200,
                            'actual': 200 if np.random.random() > 0.3 else 404,
                            'passed': np.random.random() > 0.3,
                            'message': 'Status code validation'
                        },
                        {
                            'type': 'response_time',
                            'expected': '< 2 seconds',
                            'actual': f"{np.random.random() * 2:.2f} seconds",
                            'passed': np.random.random() > 0.2,
                            'message': 'Response time check'
                        }
                    ],
                    'compliance_checks': [
                        {
                            'standard': 'HIPAA',
                            'check': 'HTTPS encryption',
                            'passed': True,
                            'message': 'Data encrypted in transit'
                        }
                    ] if 'HIPAA' in test.get('compliance', []) else [],
                    'execution_time': np.random.random() * 2
                }
            else:
                # Execute real API test
                result = executor.execute_test(test)
            
            results.append(result)
        
        progress_bar.progress(1.0)
        status_text.text("Execution complete!")
        
        # Store results
        st.session_state['test_execution_results'] = {
            'timestamp': datetime.now(),
            'results': results,
            'summary': {
                'total': len(results),
                'passed': sum(1 for r in results if r.get('status') == 'passed'),
                'failed': sum(1 for r in results if r.get('status') == 'failed'),
                'errors': sum(1 for r in results if r.get('status') == 'error')
            }
        }
        
        time.sleep(1)
        st.rerun()
    
    # Display Results
    if 'test_execution_results' in st.session_state:
        results_data = st.session_state['test_execution_results']
        
        st.markdown("---")
        st.subheader("📊 Execution Results")
        
        # Summary metrics
        summary = results_data['summary']
        col1, col2, col3, col4, col5 = st.columns(5)
        
        with col1:
            st.metric("Total Tests", summary['total'])
        with col2:
            st.metric("✅ Passed", summary['passed'], 
                     delta=f"{(summary['passed']/summary['total']*100):.0f}%")
        with col3:
            st.metric("❌ Failed", summary['failed'])
        with col4:
            st.metric("⚠️ Errors", summary['errors'])
        with col5:
            pass_rate = (summary['passed'] / summary['total'] * 100) if summary['total'] > 0 else 0
            st.metric("Pass Rate", f"{pass_rate:.1f}%")
        
        # Detailed results
        st.markdown("### Detailed Test Results")
        
        for result in results_data['results']:
            status_icon = "✅" if result['status'] == 'passed' else "❌" if result['status'] == 'failed' else "⚠️"
            
            with st.expander(f"{status_icon} **{result['test_id']}** - {result['test_title']}"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Request Details:**")
                    st.json({
                        'method': result['request'].get('method'),
                        'url': result['request'].get('url'),
                        'headers': result['request'].get('headers', {})
                    })
                
                with col2:
                    st.markdown("**Response Details:**")
                    st.json({
                        'status_code': result['response'].get('status_code'),
                        'response_time': f"{result['response'].get('response_time_ms')} ms",
                        'body_preview': str(result['response'].get('body', ''))[:200] + '...'
                    })
                
                # Assertions
                st.markdown("**Assertions:**")
                for assertion in result.get('assertions', []):
                    icon = "✅" if assertion['passed'] else "❌"
                    st.write(f"{icon} {assertion['message']}")
                    st.caption(f"Expected: {assertion['expected']} | Actual: {assertion['actual']}")
                
                # Compliance checks
                if result.get('compliance_checks'):
                    st.markdown("**Compliance Checks:**")
                    for check in result['compliance_checks']:
                        icon = "✅" if check['passed'] else "❌"
                        st.write(f"{icon} **{check['standard']}**: {check['check']}")
                        st.caption(check['message'])
        
        # Export results
        st.markdown("---")
        st.subheader("📥 Export Results")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # JSON export
            json_str = json.dumps(results_data, default=str, indent=2)
            st.download_button(
                "📄 Download JSON Report",
                json_str,
                f"api_test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                "application/json"
            )
        
        with col2:
            # HTML report
            html_report = generate_api_test_html_report(results_data)
            st.download_button(
                "📊 Download HTML Report",
                html_report,
                f"api_test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                "text/html"
            )
        
        with col3:
            # JUnit XML for CI/CD
            junit_xml = convert_api_results_to_junit(results_data)
            st.download_button(
                "🔧 Download JUnit XML",
                junit_xml,
                f"api_test_junit_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xml",
                "application/xml"
            )

def generate_api_test_html_report(results_data: Dict) -> str:
    """Generate HTML report for API test results"""
    summary = results_data['summary']
    html = f"""
    <html>
    <head>
        <title>API Test Report - {results_data['timestamp']}</title>
        <style>
            body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 20px; }}
            .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; }}
            .summary {{ display: flex; gap: 20px; margin: 20px 0; }}
            .metric {{ background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); flex: 1; }}
            .metric h3 {{ margin: 0; color: #666; font-size: 14px; }}
            .metric .value {{ font-size: 32px; font-weight: bold; margin: 10px 0; }}
            .passed {{ color: #10b981; }}
            .failed {{ color: #ef4444; }}
            .error {{ color: #f59e0b; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
            th {{ background: #f3f4f6; padding: 12px; text-align: left; }}
            td {{ padding: 12px; border-bottom: 1px solid #e5e7eb; }}
            .status-passed {{ background: #d1fae5; }}
            .status-failed {{ background: #fee2e2; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>API Test Execution Report</h1>
            <p>Generated: {results_data['timestamp']}</p>
        </div>
        
        <div class="summary">
            <div class="metric">
                <h3>Total Tests</h3>
                <div class="value">{summary['total']}</div>
            </div>
            <div class="metric">
                <h3>Passed</h3>
                <div class="value passed">{summary['passed']}</div>
            </div>
            <div class="metric">
                <h3>Failed</h3>
                <div class="value failed">{summary['failed']}</div>
            </div>
            <div class="metric">
                <h3>Pass Rate</h3>
                <div class="value">{(summary['passed']/summary['total']*100):.1f}%</div>
            </div>
        </div>
        
        <h2>Test Results</h2>
        <table>
            <thead>
                <tr>
                    <th>Test ID</th>
                    <th>Title</th>
                    <th>Status</th>
                    <th>Response Time</th>
                    <th>Assertions</th>
                </tr>
            </thead>
            <tbody>
    """
    
    for result in results_data['results']:
        assertions_passed = sum(1 for a in result.get('assertions', []) if a['passed'])
        assertions_total = len(result.get('assertions', []))
        
        html += f"""
            <tr class="status-{result['status']}">
                <td>{result['test_id']}</td>
                <td>{result['test_title']}</td>
                <td>{result['status'].upper()}</td>
                <td>{result['response'].get('response_time_ms', 'N/A')} ms</td>
                <td>{assertions_passed}/{assertions_total}</td>
            </tr>
        """
    
    html += """
            </tbody>
        </table>
    </body>
    </html>
    """
    return html

def convert_api_results_to_junit(results_data: Dict) -> str:
    """Convert API test results to JUnit XML format"""
    summary = results_data['summary']
    timestamp = results_data['timestamp'].isoformat()
    
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<testsuites name="API Test Suite" tests="{summary['total']}" failures="{summary['failed']}" errors="{summary['errors']}" time="0">
    <testsuite name="Healthcare API Tests" tests="{summary['total']}" failures="{summary['failed']}" errors="{summary['errors']}" timestamp="{timestamp}">
"""
    
    for result in results_data['results']:
        execution_time = result.get('execution_time', 0)
        xml += f"""        <testcase classname="APITests" name="{result['test_title']}" time="{execution_time:.3f}">
"""
        
        if result['status'] == 'failed':
            failed_assertions = [a for a in result.get('assertions', []) if not a['passed']]
            failure_msg = "; ".join([a['message'] for a in failed_assertions])
            xml += f"""            <failure message="{failure_msg}" type="AssertionError"/>
"""
        elif result['status'] == 'error':
            xml += f"""            <error message="{result.get('error', 'Unknown error')}" type="Error"/>
"""
        
        xml += """        </testcase>
"""
    
    xml += """    </testsuite>
</testsuites>"""
    
    return xml

def show_login_page():
    """Display the login page"""
    st.markdown("""
    <div class="main-header" style='text-align: center; padding: 2.5rem 0;'>
        <h1 class="main-title">Test Case Generation Platform v14</h1>
        <p class="main-subtitle">Healthcare & MedTech Compliance Testing | Now with Full Editing & AI Refinement</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        tab1, tab2 = st.tabs(["Login", "Sign Up"])
        
        with tab1:
            st.markdown("### Welcome Back")
            with st.form("login_form"):
                email = st.text_input("Email", placeholder="Enter your email")
                password = st.text_input("Password", type="password", placeholder="Enter your password")
                submit = st.form_submit_button("Login", type="primary", use_container_width=True)
                
                if submit:
                    logger.info(f"[LOGIN_ATTEMPT] Email: {email}")
                    
                    if not email or not password:
                        logger.warning(f"[LOGIN_VALIDATION_FAILED] Missing email or password")
                        st.error("Please fill in all fields")
                    elif not validate_email(email):
                        logger.warning(f"[LOGIN_VALIDATION_FAILED] Invalid email format: {email}")
                        st.error("Please enter a valid email address")
                    else:
                        # Authenticate user
                        if st.session_state.db:
                            logger.info(f"[AUTH_START] Authenticating user: {email}")
                            hashed_password = hash_password(password)
                            user = st.session_state.db.authenticate_user(email, hashed_password)
                            
                            if user:
                                logger.info(f"[LOGIN_SUCCESS] User authenticated: {user['_id']}")
                                # Set session state
                                st.session_state.authenticated = True
                                st.session_state.user_id = user['_id']
                                st.session_state.user_email = user['email']
                                st.session_state.user_name = user.get('full_name', 'User')
                                st.session_state.user_role = user.get('role', 'user')
                                
                                logger.info(f"[SESSION_CREATED] User: {user['_id']}, Email: {email}, Name: {st.session_state.user_name}")
                                
                                st.success("Login successful!")
                                time.sleep(1)
                                st.rerun()
                            else:
                                # Try with plain password verification for existing users
                                logger.info(f"[AUTH_FALLBACK] Trying alternative authentication for: {email}")
                                user_data = st.session_state.db.get_user_by_email(email)
                                if user_data and verify_password(password, user_data.get('password', '')):
                                    logger.info(f"[LOGIN_SUCCESS_FALLBACK] User authenticated: {user_data['_id']}")
                                    # Update last login
                                    st.session_state.authenticated = True
                                    st.session_state.user_id = user_data['_id']
                                    st.session_state.user_email = user_data['email']
                                    st.session_state.user_name = user_data.get('full_name', 'User')
                                    st.session_state.user_role = user_data.get('role', 'user')
                                    
                                    logger.info(f"[SESSION_CREATED] User: {user_data['_id']}, Email: {email}, Name: {st.session_state.user_name}")
                                    
                                    st.success("Login successful!")
                                    time.sleep(1)
                                    st.rerun()
                                else:
                                    logger.warning(f"[LOGIN_FAILED] Invalid credentials for email: {email}")
                                    st.error("Invalid email or password")
                        else:
                            logger.error("[LOGIN_ERROR] Database connection not available")
                            st.error("Database connection not available")
        
        with tab2:
            st.markdown("### Create Your Account")
            with st.form("signup_form", clear_on_submit=False):
                full_name = st.text_input("Full Name", placeholder="Enter your full name", key="signup_name")
                email = st.text_input("Email", placeholder="Enter your email", key="signup_email")
                password = st.text_input("Password", type="password", placeholder="Create a password", key="signup_pass")
                confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password", key="signup_confirm")
                
                st.info("""
                **Password Requirements:**
                - At least 8 characters long
                - Contains uppercase and lowercase letters
                - Contains at least one number
                """)
                
                submit = st.form_submit_button("Sign Up", type="primary", use_container_width=True)
                
                if submit:
                    if not all([full_name, email, password, confirm_password]):
                        st.error("Please fill in all fields")
                    elif not validate_email(email):
                        st.error("Please enter a valid email address")
                    elif password != confirm_password:
                        st.error("Passwords do not match")
                    else:
                        is_valid, error_msg = validate_password(password)
                        if not is_valid:
                            st.error(error_msg)
                        else:
                            # Create user account
                            if st.session_state.db:
                                logger.info(f"[SIGNUP_START] Creating account for: {email}")
                                hashed_password = hash_password(password)
                                
                                # Encrypt sensitive fields
                                encrypted_name = encrypt_sensitive_data(full_name)
                                logger.info(f"[ENCRYPTION] User data encrypted for storage")
                                
                                success, result = st.session_state.db.create_user(
                                    email=email,
                                    password_hash=hashed_password,
                                    full_name=encrypted_name  # Store encrypted name
                                )
                                
                                if success:
                                    st.success("✅ Account created successfully! Please login.")
                                    st.balloons()
                                    time.sleep(2)
                                    st.rerun()
                                else:
                                    st.error(f"Failed to create account: {result}")
                            else:
                                st.error("Database connection not available")

# MAIN UI
def main():
    logger.info("[MAIN] Application v14.0 main() called - Full Editing & AI Refinement Enabled")
    
    # Apply modern CSS
    inject_modern_css()
    
    # Initialize MongoDB connection
    if 'db' not in st.session_state:
        logger.info("[MONGODB] Initializing database connection...")
        with st.spinner("Establishing database connection..."):
            st.session_state.db = init_mongodb()
            if st.session_state.db:
                logger.info("[MONGODB] Database connection established")
            else:
                logger.warning("[MONGODB] Database connection failed - running in local mode")
    
    # Check authentication
    if not check_authentication():
        logger.info("[AUTH_CHECK] User not authenticated, showing login page")
        show_login_page()
        return
    
    # User is authenticated, show main app
    logger.info(f"[AUTH_CHECK] User authenticated: {st.session_state.user_email}")
    
    # Get or create session for the authenticated user FIRST
    if 'session_id' not in st.session_state:
        logger.info("[SESSION] Creating new session for authenticated user")
        get_or_create_session()
    
    # Load user's test cases from MongoDB after authentication and session creation
    if 'tests_loaded_for_user' not in st.session_state or st.session_state.get('tests_loaded_user_id') != st.session_state.get('user_id'):
        if st.session_state.db and st.session_state.get('user_id'):
            with st.spinner("Loading your test cases..."):
                user_tests = load_tests_from_mongodb(limit=200)
                if user_tests:
                    st.session_state.generated_tests = user_tests
                    logger.info(f"[TESTS_LOADED] Loaded {len(user_tests)} test cases for user {st.session_state.get('user_id')}")
                    # Only show success message on first load, not on every rerun
                    if 'shown_load_success' not in st.session_state:
                        st.success(f"✅ Loaded {len(user_tests)} test case(s) from your workspace", icon="📂")
                        st.session_state.shown_load_success = True
                else:
                    # Initialize empty list if no tests found
                    st.session_state.generated_tests = []
                    logger.info(f"[TESTS_LOADED] No existing test cases found for user {st.session_state.get('user_id')}")
                
                st.session_state.tests_loaded_for_user = True
                st.session_state.tests_loaded_user_id = st.session_state.get('user_id')
        else:
            # Initialize empty list if no database connection
            if 'generated_tests' not in st.session_state:
                st.session_state.generated_tests = []
    
    # User header with logout and profile info
    col1, col2, col3 = st.columns([6, 3, 1.5])
    
    with col1:
        st.markdown(f"""
        <div style='padding: 1rem 0;'>
            <h3 style='margin: 0; color: var(--text-primary);'>Welcome, {st.session_state.user_name}!</h3>
            <p style='margin: 0.25rem 0; color: var(--text-secondary); font-size: 0.9rem;'>
                {st.session_state.user_email}
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        # Account settings button
        if st.button("⚙️ Account Settings", use_container_width=True):
            st.session_state.show_settings = True
    
    with col3:
        if st.button("🚪 Logout", type="secondary", use_container_width=True):
            logout()
    
    st.markdown("---")
    
    # Clean, professional header
    st.markdown("""
    <div class="main-header" style='text-align: center; padding: 1.5rem 0; animation: fadeIn 0.8s ease;'>
        <h1 class="main-title">
            AI-Powered Test Case Generation Platform
        </h1>
        <p class="main-subtitle">
            Healthcare & MedTech Compliance Testing Solution | Enterprise Edition v14.0
        </p>
        <div style='margin-top: 1rem; display: flex; justify-content: center; gap: 0.75rem; flex-wrap: wrap;'>
            <span class="feature-badge">
                ✏️ Manual Test Editing
            </span>
            <span class="feature-badge">
                🤖 AI-Powered Refinement
            </span>
            <span class="feature-badge">
                📊 Version Control
            </span>
            <span class="feature-badge">
                🏥 Healthcare Compliant
            </span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Show account settings if requested
    if 'show_settings' in st.session_state and st.session_state.show_settings:
        with st.expander("⚙️ Account Settings", expanded=True):
            st.subheader("Profile Information")
            
            with st.form("profile_form"):
                new_name = st.text_input("Full Name", value=st.session_state.user_name)
                new_email = st.text_input("Email", value=st.session_state.user_email, disabled=True, 
                                        help="Email cannot be changed for security reasons")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.form_submit_button("Update Profile", type="primary"):
                        if st.session_state.db:
                            # Log profile update attempt
                            logger.info(f"[PROFILE_UPDATE] User {st.session_state.user_id} updating profile")
                            
                            success = st.session_state.db.update_user_profile(
                                st.session_state.user_id,
                                {'full_name': new_name}
                            )
                            if success:
                                st.session_state.user_name = new_name
                                logger.info(f"[PROFILE_UPDATE_SUCCESS] User {st.session_state.user_id} profile updated")
                                st.success("✅ Profile updated successfully")
                            else:
                                logger.error(f"[PROFILE_UPDATE_FAILED] User {st.session_state.user_id} profile update failed")
                                st.error("Failed to update profile")
                
                with col2:
                    if st.form_submit_button("Change Password", type="secondary"):
                        st.info("Password change feature coming soon")
            
            st.markdown("---")
            st.caption(f"User ID: {st.session_state.user_id[:8]}...")
            st.caption(f"Member Since: {datetime.now().strftime('%B %Y')}")
            
            if st.button("Close Settings", type="secondary", use_container_width=True):
                st.session_state.show_settings = False
                st.rerun()
        
        st.markdown("---")
    
    # Sidebar removed - uncomment below to restore MongoDB status display
    # with st.sidebar:
    #     st.markdown("### 🗄️ Database Status")
    #     if st.session_state.db and st.session_state.db.ping():
    #         st.success("MongoDB: Online ✅")
    #         
    #         # Show statistics
    #         if st.button("📊 View Stats"):
    #             stats = st.session_state.db.get_statistics(get_or_create_session())
    #             st.metric("Total Tests", stats.get('total_test_cases', 0))
    #             st.metric("Documents", stats.get('total_documents', 0))
    #             st.metric("Compliance Reports", stats.get('total_compliance_reports', 0))
    #     else:
    #         st.warning("MongoDB: Offline ⚠️")
    #         st.info("Using local file storage")
    #     
    #     st.markdown("---")
    #     if 'session_id' in st.session_state:
    #         st.caption(f"Session: {st.session_state.session_id[:8]}...")
    
    # NOTE: Test loading is now handled immediately after authentication (see line ~5211)
    # This ensures user-specific test cases are loaded correctly with data isolation
    # The old logic below has been commented out to prevent interference
    
    # # Load saved tests on first run - Try MongoDB first, then files
    # if 'tests_loaded' not in st.session_state:
    #     if st.session_state.db:
    #         # Try to load from MongoDB
    #         mongo_tests = load_tests_from_mongodb()
    #         if mongo_tests:
    #             st.session_state.generated_tests = mongo_tests
    #             st.info(f"📥 Loaded {len(mongo_tests)} tests from MongoDB")
    #         else:
    #             # Fallback to file system
    #             all_saved_tests = initialize_saved_tests()
    #             if all_saved_tests:
    #                 st.session_state.generated_tests = all_saved_tests
    #                 # Optionally migrate to MongoDB
    #                 if st.session_state.db:
    #                     with st.spinner("Migrating tests to MongoDB..."):
    #                         user_id = st.session_state.get('user_id')
    #                         success, ids = st.session_state.db.save_test_cases_batch(
    #                             all_saved_tests, get_or_create_session(), user_id
    #                         )
    #                         if success:
    #                             st.success(f"✅ Migrated {len(ids)} tests to MongoDB for your account")
    #     else:
    #         # No MongoDB, use file system
    #         all_saved_tests = initialize_saved_tests()
    #         if all_saved_tests:
    #             st.session_state.generated_tests = all_saved_tests
    #     st.session_state.tests_loaded = True
    
    # Load RAG system (needed for functionality but not displayed)
    embedding_model, index, doc_chunks, doc_metadata = load_rag_system()

    # Enhanced Navigation Bar with Custom Styling
    st.markdown("""
    <style>
        /* Modern Navigation Bar Styling */
        .stTabs {
            background: transparent !important;
            padding: 0 !important;
            margin-top: 2rem !important;
        }
        
        .stTabs [data-baseweb="tab-list"] {
            background: var(--gradient-surface) !important;
            border: 2px solid var(--border-primary) !important;
            border-radius: 24px !important;
            padding: 12px !important;
            gap: 8px !important;
            box-shadow: var(--shadow-xl) !important;
            backdrop-filter: blur(20px) !important;
            min-height: 80px !important;
            display: flex !important;
            align-items: center !important;
            justify-content: space-evenly !important;
            position: relative !important;
            overflow: hidden !important;
        }
        
        /* Animated background gradient for tab list */
        .stTabs [data-baseweb="tab-list"]::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 200%;
            height: 100%;
            background: linear-gradient(
                90deg, 
                transparent, 
                var(--primary-subtle), 
                transparent
            );
            animation: navShimmer 8s infinite linear;
            opacity: 0.3;
            pointer-events: none;
        }
        
        @keyframes navShimmer {
            0% { transform: translateX(-100%); }
            100% { transform: translateX(100%); }
        }
        
        /* Individual Tab Styling */
        .stTabs [data-baseweb="tab"] {
            background: var(--bg-elevated) !important;
            border: 2px solid transparent !important;
            border-radius: 18px !important;
            padding: 16px 28px !important;
            font-size: 1rem !important;
            font-weight: 600 !important;
            letter-spacing: 0.3px !important;
            color: var(--text-secondary) !important;
            transition: all 0.4s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            position: relative !important;
            overflow: hidden !important;
            min-width: 140px !important;
            text-align: center !important;
            white-space: nowrap !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            gap: 8px !important;
            cursor: pointer !important;
            flex: 1 !important;
            max-width: 200px !important;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05) !important;
            transform: translateY(0) scale(1) !important;
            z-index: 1 !important;
        }
        
        /* Light mode tab background */
        @media (prefers-color-scheme: light) {
            .stTabs [data-baseweb="tab"] {
                background: linear-gradient(135deg, 
                    rgba(255, 255, 255, 0.95) 0%, 
                    rgba(249, 250, 251, 0.95) 100%) !important;
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04),
                           inset 0 1px 0 rgba(255, 255, 255, 1) !important;
            }
        }
        
        /* Dark mode tab background */
        @media (prefers-color-scheme: dark) {
            .stTabs [data-baseweb="tab"] {
                background: linear-gradient(135deg, 
                    rgba(31, 41, 55, 0.95) 0%, 
                    rgba(20, 24, 36, 0.95) 100%) !important;
                box-shadow: 0 2px 12px rgba(0, 0, 0, 0.3),
                           inset 0 1px 0 rgba(255, 255, 255, 0.03) !important;
            }
        }
        
        /* Hover effect with ripple animation */
        .stTabs [data-baseweb="tab"]::before {
            content: '';
            position: absolute;
            top: 50%;
            left: 50%;
            width: 0;
            height: 0;
            border-radius: 50%;
            background: var(--primary-subtle);
            transform: translate(-50%, -50%);
            transition: width 0.6s, height 0.6s;
            z-index: -1;
        }
        
        .stTabs [data-baseweb="tab"]:hover {
            background: var(--gradient-primary) !important;
            color: var(--text-on-primary) !important;
            border-color: var(--primary) !important;
            transform: translateY(-4px) scale(1.08) !important;
            box-shadow: 0 12px 28px var(--shadow-glow) !important;
            z-index: 10 !important;
        }
        
        .stTabs [data-baseweb="tab"]:hover::before {
            width: 300px;
            height: 300px;
            opacity: 0;
        }
        
        /* Active/Selected Tab with animation */
        .stTabs [aria-selected="true"] {
            background: var(--gradient-primary) !important;
            color: var(--text-on-primary) !important;
            border: 2px solid var(--primary) !important;
            box-shadow: 0 8px 32px rgba(124, 58, 237, 0.25),
                       0 0 60px rgba(124, 58, 237, 0.15) !important;
            transform: translateY(-6px) scale(1.1) !important;
            z-index: 20 !important;
            animation: tabPulse 2s ease-in-out infinite !important;
        }
        
        @keyframes tabPulse {
            0%, 100% {
                box-shadow: 0 8px 32px rgba(124, 58, 237, 0.25),
                           0 0 60px rgba(124, 58, 237, 0.15);
            }
            50% {
                box-shadow: 0 12px 40px rgba(124, 58, 237, 0.35),
                           0 0 80px rgba(124, 58, 237, 0.2);
            }
        }
        
        /* Icon animations in tabs */
        .stTabs [data-baseweb="tab"] span {
            display: inline-block !important;
            transition: transform 0.3s ease !important;
        }
        
        .stTabs [data-baseweb="tab"]:hover span:first-child {
            animation: iconBounce 0.5s ease !important;
        }
        
        @keyframes iconBounce {
            0%, 100% { transform: scale(1) rotate(0deg); }
            25% { transform: scale(1.2) rotate(-5deg); }
            75% { transform: scale(1.2) rotate(5deg); }
        }
        
        /* Tab panel fade-in animation */
        .stTabs [data-baseweb="tab-panel"] {
            animation: panelSlideIn 0.6s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            padding: 2rem 0 !important;
        }
        
        @keyframes panelSlideIn {
            from {
                opacity: 0;
                transform: translateY(20px) scale(0.98);
            }
            to {
                opacity: 1;
                transform: translateY(0) scale(1);
            }
        }
        
        /* Mobile responsive adjustments */
        @media (max-width: 768px) {
            .stTabs [data-baseweb="tab-list"] {
                flex-wrap: wrap !important;
                min-height: auto !important;
                padding: 8px !important;
            }
            
            .stTabs [data-baseweb="tab"] {
                min-width: calc(50% - 8px) !important;
                padding: 12px 16px !important;
                font-size: 0.9rem !important;
            }
        }
        
        /* Tab notification badge (for future use) */
        .tab-badge {
            position: absolute;
            top: -4px;
            right: -4px;
            background: var(--danger);
            color: white;
            border-radius: 50%;
            width: 20px;
            height: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 0.7rem;
            font-weight: bold;
            animation: badgePulse 2s infinite;
        }
        
        @keyframes badgePulse {
            0%, 100% {
                transform: scale(1);
                opacity: 1;
            }
            50% {
                transform: scale(1.1);
                opacity: 0.8;
            }
        }
    </style>
    """, unsafe_allow_html=True)

    # Main content area with enhanced tabs - ORDERED BY WORKFLOW
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "Document Upload",
        "Import Test Cases",
        "Generate Test Cases",
        "Test Suite Management", 
        "API Test Execution",
        "Export & Integration",
        "Document Repository"
    ])

    # Tab 1: UPLOAD DOCUMENTS (First Step)
    with tab1:
        st.header("Document Upload Center")
        st.info("""
        **Required First Step:** Upload project documentation including requirements specifications, user stories, 
        API documentation, and technical specifications. These documents provide essential context for 
        intelligent test case generation.
        """)
        display_upload_interface()

    # Tab 2: IMPORT EXISTING TESTS (Optional Step)
    with tab2:
        st.header("Import Existing Test Cases")
        st.info("""
        **Optional:** Import your existing test cases from CSV, Excel, or JSON formats. 
        The system will standardize them to NASSCOM-compliant format and integrate them into your test suite.
        """)
        display_import_interface()

    # Tab 3: GENERATE NEW TESTS (Main Feature)
    with tab3:
        st.header("AI-Powered Test Case Generation")
        st.info("""
        **Primary Function:** Describe your testing requirements in natural language. 
        The AI engine leverages uploaded documentation to generate comprehensive, context-aware test cases 
        that comply with healthcare industry standards.
        """)

        # Predefined examples
        examples = [
            "Test user registration with HIPAA-compliant data handling",
            "Verify patient data encryption meets FDA and ISO 27001 standards",
            "Test API authentication with OAuth 2.0 for medical records access",
            "Validate HL7 message processing for lab results",
            "Test DICOM image upload with PHI anonymization",
            "Verify audit trail generation for GDPR compliance"
        ]
        
        selected_example = st.selectbox("Select a template or define custom requirements:", 
                                       ["Custom requirement..."] + examples)
        
        if selected_example == "Custom requirement...":
            requirement = st.text_area(
                "Test Requirement Specification:",
                height=100,
                placeholder="Describe the functionality or scenario requiring test coverage..."
            )
        else:
            requirement = selected_example
            st.info(f"Selected template: {requirement}")
        
        # Generation button
        if st.button("Generate Test Case", type="primary"):
            if not requirement or requirement == "Custom requirement...":
                st.warning("Please provide test requirements to proceed.")
            elif not index:
                st.error("Knowledge base initialization required. Please upload relevant documentation in the Document Upload section.")
            else:
                logger.info(f"[TEST_GENERATION_START] User: {st.session_state.user_id}, Requirement: {requirement[:100]}...")
                
                # Configure API key
                genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
                logger.info("[API_CONFIG] Using default Gemini API key")
                
                # Show AI thinking animation - minimal style for first phase
                show_ai_thinking(duration=2, style="minimal")
                
                with st.spinner("Analyzing requirement specification..."):
                    logger.info("[CONTEXT_RETRIEVAL] Starting RAG context retrieval")
                    # Retrieve context
                    context_docs = retrieve_context(
                        requirement, 
                        embedding_model, 
                        index, 
                        doc_chunks, 
                        doc_metadata,
                        k=5
                    )
                    logger.info(f"[CONTEXT_RETRIEVED] Found {len(context_docs)} relevant documents")
                    
                # More AI thinking for generation - detailed style for main processing
                show_ai_thinking(duration=3, style="detailed")
                
                with st.spinner("Generating compliance-validated test case..."):
                    logger.info("[AI_GENERATION] Starting test case generation with Gemini")
                    # Generate test case
                    test_case = generate_test_case_with_gemini(requirement, context_docs)
                    logger.info(f"[AI_GENERATION_SUCCESS] Test case generated: {test_case.get('id', 'unknown')}")
                    
                    # Add context to test case object
                    test_case['retrieved_context'] = context_docs
                    
                    # Clean numpy types from test case
                    test_case = convert_numpy_to_python(test_case)
                    
                    # Store in session
                    st.session_state.generated_tests.insert(0, test_case)
                    
                    # Save to MongoDB if available, otherwise to file
                    if st.session_state.db:
                        if save_test_to_mongodb(test_case):
                            st.success(f"✅ Test case saved to MongoDB: {test_case['id']}")
                        else:
                            # Fallback to file save
                            auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                    else:
                        # AUTO-SAVE all test cases to unified location
                        auto_save_test_cases(st.session_state.generated_tests, "all_tests")
                        # Clean up old saves to prevent disk bloat
                        cleanup_old_saves("all_tests", keep_count=10)
                    
                # Display results
                if test_case.get('nasscom_compliant', False):
                    logger.info(f"[TEST_COMPLIANCE] Test case {test_case['id']} is NASSCOM compliant")
                    st.success(f"Test case successfully generated and saved: {test_case['id']}")
                else:
                    logger.warning(f"[TEST_COMPLIANCE] Test case {test_case['id']} has partial compliance")
                    st.warning(f"Test case generated with partial compliance: {test_case['id']}")
                
                # Store in session state for persistence
                st.session_state['last_generated_test'] = test_case
                st.session_state['last_generated_context'] = context_docs
                st.session_state['last_requirement'] = requirement
                logger.info(f"[SESSION_STATE] Stored test case {test_case['id']} in session")
                
                # Show success message with professional animation
                success_placeholder = st.empty()
                success_placeholder.markdown("""
                <div class='ai-thinking-container' style='border-color: var(--success);'>
                    <div class='ai-thinking-icon' style='animation: none; color: var(--success); font-size: 2rem;'>✓</div>
                    <p class='loading-message' style='color: var(--success); font-weight: 600;'>Test Case Generated Successfully</p>
                </div>
                """, unsafe_allow_html=True)
                time.sleep(1.2)
                success_placeholder.empty()
        
        # Display the last generated test case (persists across reruns)
        if 'last_generated_test' in st.session_state:
            st.markdown("---")
            col1, col2, col3 = st.columns([5, 1, 1])
            with col1:
                st.markdown("### Generated Test Case")
            with col2:
                if 'last_requirement' in st.session_state:
                    if st.button("Regenerate", key="regenerate_test", help="Generate a new test for the same requirement"):
                        # Clear and trigger regeneration
                        if 'last_generated_test' in st.session_state:
                            del st.session_state['last_generated_test']
                        if 'last_generated_context' in st.session_state:
                            del st.session_state['last_generated_context']
                        # Clear save status
                        keys_to_clear = [k for k in st.session_state.keys() if k.startswith('saved_')]
                        for key in keys_to_clear:
                            del st.session_state[key]
                        st.rerun()
            with col3:
                if st.button("Clear", key="clear_test", help="Clear the generated test case"):
                    # Clear all related session state
                    if 'last_generated_test' in st.session_state:
                        del st.session_state['last_generated_test']
                    if 'last_generated_context' in st.session_state:
                        del st.session_state['last_generated_context']
                    if 'last_requirement' in st.session_state:
                        del st.session_state['last_requirement']
                    # Clear save status
                    keys_to_clear = [k for k in st.session_state.keys() if k.startswith('saved_')]
                    for key in keys_to_clear:
                        del st.session_state[key]
                    st.rerun()
            
            # Display the persisted test case
            display_test_case(
                st.session_state['last_generated_test'], 
                st.session_state.get('last_generated_context', []), 
                "_gen"
            )

    # Tab 4: VIEW TEST SUITE (Review Generated Tests)
    with tab4:
        st.header("Test Suite Management")
        st.info("""
        **Quality Review:** Examine and manage your complete test suite. Apply filters by category, priority, 
        or compliance standards to organize and validate test coverage.
        """)
        
        # Ensure generated_tests exists
        if 'generated_tests' not in st.session_state:
            st.session_state.generated_tests = []
        
        if not st.session_state.generated_tests:
            st.warning("No test cases available. Please upload documents and generate test cases to proceed.")
        else:
            st.markdown(f"### Current Test Suite: **{len(st.session_state.generated_tests)}** Test Case(s)")
            
            # Filter options
            col1, col2, col3 = st.columns(3)
            with col1:
                filter_category = st.selectbox("Filter by Category", 
                                              ["All"] + list(set(tc.get('category', 'Unknown') for tc in st.session_state.generated_tests)))
            with col2:
                filter_priority = st.selectbox("Filter by Priority",
                                              ["All"] + list(set(tc.get('priority', 'Unknown') for tc in st.session_state.generated_tests)))
            with col3:
                filter_compliance = st.selectbox("Filter by Compliance",
                                                ["All", "NASSCOM Compliant", "Non-Compliant"])
            
            # Apply filters
            filtered_tests = st.session_state.generated_tests
            if filter_category != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('category') == filter_category]
            if filter_priority != "All":
                filtered_tests = [tc for tc in filtered_tests if tc.get('priority') == filter_priority]
            if filter_compliance == "NASSCOM Compliant":
                filtered_tests = [tc for tc in filtered_tests if tc.get('nasscom_compliant', False)]
            elif filter_compliance == "Non-Compliant":
                filtered_tests = [tc for tc in filtered_tests if not tc.get('nasscom_compliant', False)]
            
            # Bulk Actions
            st.markdown("---")
            bulk_col1, bulk_col2, bulk_col3, bulk_col4 = st.columns(4)
            
            with bulk_col1:
                if st.button("📥 Export All Filtered", help="Export filtered tests to CSV"):
                    if filtered_tests:
                        df = pd.DataFrame(filtered_tests)
                        csv = df.to_csv(index=False)
                        st.download_button(
                            "⬇️ Download CSV",
                            csv,
                            f"filtered_tests_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                            "text/csv",
                            key="bulk_download_csv"
                        )
            
            with bulk_col2:
                if st.button("🗑️ Clear Non-Compliant", help="Remove all non-compliant tests"):
                    if st.session_state.generated_tests:
                        non_compliant_tests = [tc for tc in st.session_state.generated_tests 
                                              if not tc.get('nasscom_compliant', False)]
                        compliant_tests = [tc for tc in st.session_state.generated_tests 
                                         if tc.get('nasscom_compliant', False)]
                        
                        # Delete non-compliant from MongoDB
                        if st.session_state.db and non_compliant_tests:
                            session_id = get_or_create_session()
                            for tc in non_compliant_tests:
                                st.session_state.db.delete_test_case(tc['id'], session_id)
                            logger.info(f"[CLEANUP] Deleted {len(non_compliant_tests)} non-compliant tests from MongoDB")
                        
                        removed = len(st.session_state.generated_tests) - len(compliant_tests)
                        st.session_state.generated_tests = compliant_tests
                        st.success(f"Removed {removed} non-compliant test(s)")
                        time.sleep(1)
                        st.rerun()
            
            with bulk_col3:
                if st.button("📊 Suite Statistics", help="View test suite analytics"):
                    st.session_state.show_suite_stats = not st.session_state.get('show_suite_stats', False)
                    st.rerun()
            
            with bulk_col4:
                if st.button("🔄 Refresh Suite", help="Reload test cases"):
                    if st.session_state.db and st.session_state.get('user_id'):
                        tests = load_tests_from_mongodb(limit=100)
                        if tests:
                            st.session_state.generated_tests = tests
                            st.success(f"Loaded {len(tests)} test(s) from database")
                    st.rerun()
            
            # Suite Statistics
            if st.session_state.get('show_suite_stats', False):
                st.markdown("---")
                st.markdown("### 📊 Test Suite Analytics")
                
                col1, col2, col3, col4, col5 = st.columns(5)
                
                with col1:
                    total = len(st.session_state.generated_tests)
                    st.metric("Total Tests", total)
                
                with col2:
                    compliant = sum(1 for tc in st.session_state.generated_tests 
                                  if tc.get('nasscom_compliant', False))
                    st.metric("Compliant", compliant, f"{compliant/total*100:.0f}%")
                
                with col3:
                    categories = {}
                    for tc in st.session_state.generated_tests:
                        cat = tc.get('category', 'Unknown')
                        categories[cat] = categories.get(cat, 0) + 1
                    st.metric("Categories", len(categories))
                
                with col4:
                    high_priority = sum(1 for tc in st.session_state.generated_tests 
                                      if tc.get('priority') in ['Critical', 'High'])
                    st.metric("High Priority", high_priority)
                
                with col5:
                    edited = sum(1 for tc in st.session_state.generated_tests 
                               if tc.get('manually_edited', False) or tc.get('version', 1) > 1)
                    st.metric("Modified", edited)
                
                # Category Distribution
                if categories:
                    st.markdown("**Test Distribution by Category:**")
                    for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
                        st.progress(count / total, text=f"{cat}: {count} tests")
            
            st.markdown("---")
            st.markdown(f"### Showing **{len(filtered_tests)}** test case(s)")
            
            # Test Case Display with enhanced actions
            for i, tc in enumerate(filtered_tests):
                # Add version indicator and edit status to expander title
                version_badge = f"v{tc.get('version', 1)}"
                edit_badge = " ✏️" if tc.get('manually_edited', False) else ""
                refinement_badge = " 🤖" if tc.get('refinement_feedback') else ""
                
                with st.expander(
                    f"**{tc['id']}**: {tc.get('title', 'Untitled')} ({version_badge}{edit_badge}{refinement_badge})", 
                    expanded=(i==0)
                ):
                    # Quick Actions Bar
                    quick_col1, quick_col2, quick_col3, quick_col4 = st.columns(4)
                    
                    with quick_col1:
                        if st.button("🗑️", key=f"delete_{tc['id']}_{i}", help="Delete this test case"):
                            # Delete from MongoDB first
                            if st.session_state.db:
                                session_id = get_or_create_session()
                                if st.session_state.db.delete_test_case(tc['id'], session_id):
                                    logger.info(f"[DELETE] Test case {tc['id']} deleted from MongoDB")
                            
                            # Remove from session state
                            st.session_state.generated_tests.remove(tc)
                            st.success(f"Deleted: {tc['id']}")
                            time.sleep(0.5)
                            st.rerun()
                    
                    with quick_col2:
                        # Show version history if available
                        if tc.get('version', 1) > 1:
                            if st.button("📜", key=f"history_{tc['id']}_{i}", help="View version history"):
                                st.session_state[f'show_history_{tc["id"]}'] = True
                                st.rerun()
                    
                    with quick_col3:
                        # Mark as favorite/important
                        is_favorite = tc.get('is_favorite', False)
                        if st.button("⭐" if is_favorite else "☆", 
                                   key=f"fav_{tc['id']}_{i}", 
                                   help="Mark as favorite"):
                            tc['is_favorite'] = not is_favorite
                            # Save to MongoDB
                            if st.session_state.db and st.session_state.get('user_id'):
                                save_test_to_mongodb(tc)
                            st.rerun()
                    
                    with quick_col4:
                        # Quick compliance toggle
                        is_compliant = tc.get('nasscom_compliant', False)
                        if st.button("✅" if is_compliant else "⚠️", 
                                   key=f"comp_{tc['id']}_{i}",
                                   help="Toggle compliance status"):
                            tc['nasscom_compliant'] = not is_compliant
                            # Save to MongoDB
                            if st.session_state.db and st.session_state.get('user_id'):
                                save_test_to_mongodb(tc)
                            st.rerun()
                    
                    # Show version history if requested
                    if st.session_state.get(f'show_history_{tc["id"]}', False):
                        st.markdown("---")
                        st.markdown("#### 📜 Version History")
                        
                        history_data = []
                        if tc.get('generation_timestamp'):
                            history_data.append(f"• **v1** - Created: {tc['generation_timestamp'][:19]}")
                        if tc.get('last_modified'):
                            history_data.append(f"• **v{tc.get('version', 2)}** - Modified: {tc['last_modified'][:19]} {'(Manual)' if tc.get('manually_edited') else '(AI)'}")
                        if tc.get('refinement_timestamp'):
                            history_data.append(f"• **Refinement** - {tc['refinement_timestamp'][:19]}: {tc.get('refinement_feedback', 'N/A')[:50]}...")
                        
                        for item in history_data:
                            st.write(item)
                        
                        if st.button("Close History", key=f"close_history_{tc['id']}_{i}"):
                            st.session_state[f'show_history_{tc["id"]}'] = False
                            st.rerun()
                        st.markdown("---")
                    
                    # Display the test case with all new features
                    display_test_case(tc, tc.get('retrieved_context', []), f"_suite_{i}")

    # Tab 5: API TEST EXECUTION (NEW!)
    with tab5:
        st.header("🚀 API Test Execution")
        st.info("""
        **Execute API Tests:** Run your generated API test cases against real or mock endpoints.
        Monitor response times, validate status codes, and verify healthcare compliance in real-time.
        """)
        display_api_test_executor()
    
    # Tab 6: EXPORT TESTS (Final Step)
    with tab6:
        st.header("Export and Integration")
        st.info("""
        **Export Options:** Generate output files compatible with your existing test management infrastructure.
        Supported formats include CSV, JSON, Jira, Azure DevOps, and TestRail for seamless integration.
        """)
        display_export_interface()

    # Tab 7: DOCUMENT LIBRARY (Reference)
    with tab7:
        st.header("Document Repository")
        st.markdown("View and manage your personal document library and access shared knowledge base documents.")
        
        # Create tabs for different document views
        doc_tab1, doc_tab2 = st.tabs(["My Documents", "Shared Knowledge Base"])
        
        with doc_tab1:
            st.subheader("📤 My Uploaded Documents")
            
            # Load user's documents from MongoDB
            if st.session_state.db:
                user_id = st.session_state.get('user_id')
                user_documents = st.session_state.db.get_all_documents(user_id=user_id, limit=100)
                
                if user_documents:
                    st.info(f"Found {len(user_documents)} documents in your library")
                    
                    # Create a selectbox to choose document
                    doc_names = [doc.get('filename', 'Unnamed') for doc in user_documents]
                    selected_doc_idx = st.selectbox(
                        "Select a document to view:",
                        range(len(doc_names)),
                        format_func=lambda x: doc_names[x],
                        key="my_doc_selector"
                    )
                    
                    if selected_doc_idx is not None:
                        selected_doc = user_documents[selected_doc_idx]
                        
                        # Display document metadata
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("File Type", selected_doc.get('doc_type', 'Unknown'))
                        with col2:
                            st.metric("Size", f"{selected_doc.get('content_length', 0):,} chars")
                        with col3:
                            upload_date = selected_doc.get('uploaded_at', 'Unknown')
                            if upload_date != 'Unknown' and hasattr(upload_date, 'strftime'):
                                upload_date = upload_date.strftime('%Y-%m-%d %H:%M')
                            st.metric("Uploaded", upload_date)
                        
                        # Show document content directly
                        st.markdown("---")
                        st.markdown("### 📄 Document Content")
                        
                        content = selected_doc.get('content', 'No content available')
                        # Decrypt if encrypted
                        try:
                            if content and content.startswith('gAAAAA'):  # Fernet encrypted data signature
                                content = decrypt_sensitive_data(content)
                        except:
                            pass  # Not encrypted or decryption failed, show as is
                        
                        filename = selected_doc.get('filename', '')
                        
                        # Display content based on file type
                        if filename.endswith('.md'):
                            st.markdown(content)
                        elif filename.endswith(('.yaml', '.yml')):
                            st.code(content, language='yaml')
                        elif filename.endswith('.json'):
                            try:
                                st.json(json.loads(content))
                            except:
                                st.code(content, language='json')
                        else:
                            st.text_area("", content, height=400, key=f"my_doc_content_{selected_doc_idx}")
                        
                        # Download button
                        st.download_button(
                            label="⬇️ Download Document",
                            data=content,
                            file_name=selected_doc.get('filename', 'document.txt'),
                            mime="text/plain",
                            key=f"dl_my_{selected_doc_idx}"
                        )
                else:
                    st.info("You haven't uploaded any documents yet. Go to 'Document Upload' tab to add documents to your library.")
            else:
                st.warning("Database connection required to view documents.")
        
        with doc_tab2:
            st.subheader("📚 Shared Knowledge Base")
            st.info("These are pre-loaded healthcare documents available to all users for quick demo and RAG-based test generation.")
            
            # Load shared documents from MongoDB
            if st.session_state.db:
                shared_docs = st.session_state.db.get_all_shared_documents(limit=50)
                if shared_docs:
                    # Create selectbox for document selection
                    doc_options = []
                    doc_map = {}
                    
                    for doc in shared_docs:
                        doc_type = doc.get('doc_type', '').replace('_', ' ').title()
                        filename = doc.get('filename', 'Unknown')
                        doc_id = doc.get('_id')
                        metadata = doc.get('metadata', {})
                        
                        label = f"[{doc_type}] {filename}"
                        doc_options.append(label)
                        doc_map[label] = (doc_id, filename, metadata)
                    
                    selected_label = st.selectbox(
                        "Select a document to view:",
                        doc_options,
                        key="shared_doc_selector"
                    )
                    
                    if selected_label:
                        doc_id, filename, metadata = doc_map[selected_label]
                        
                        # Display metadata
                        col1, col2 = st.columns(2)
                        with col1:
                            if metadata.get('description'):
                                st.info(metadata['description'])
                        with col2:
                            if metadata.get('tags'):
                                tag_str = " ".join([f"`{tag}`" for tag in metadata['tags'][:5]])
                                st.caption(f"Tags: {tag_str}")
                        
                        # Load and display the document content
                        st.markdown("---")
                        st.markdown("### 📖 Document Content")
                        
                        try:
                            full_doc = st.session_state.db.get_shared_document(doc_id)
                            if full_doc:
                                content = full_doc.get('content', 'No content available')
                                
                                # Display content based on file type
                                if filename.endswith('.md'):
                                    st.markdown(content)
                                elif filename.endswith(('.yaml', '.yml')):
                                    st.code(content, language='yaml')
                                elif filename.endswith('.json'):
                                    try:
                                        st.json(json.loads(content))
                                    except:
                                        st.code(content, language='json')
                                else:
                                    st.text_area("", content, height=500, key=f"shared_content_{doc_id}")
                                
                                # Download button
                                st.download_button(
                                    label="⬇️ Download Document",
                                    data=content,
                                    file_name=filename,
                                    mime="text/plain",
                                    key=f"dl_shared_{doc_id}"
                                )
                            else:
                                st.error("Failed to load document content")
                        except Exception as e:
                            st.error(f"Error loading document: {e}")
                else:
                    st.info("No shared documents loaded. Run the populate script to load demo documents.")
            else:
                st.warning("MongoDB not connected. Shared documents require database connection.")
        
        # Statistics - User specific
        st.markdown("---")
        st.subheader("📊 Your Document Statistics")
        
        if st.session_state.db:
            user_id = st.session_state.get('user_id')
            user_docs = st.session_state.db.get_all_documents(user_id=user_id, limit=100)
            shared_docs_mongo = st.session_state.db.get_all_shared_documents(limit=100)
        
        col1, col2, col3 = st.columns(3)
        with col1:
                st.metric("Your Documents", len(user_docs) if user_docs else 0)
        with col2:
                st.metric("Shared Documents", len(shared_docs_mongo) if shared_docs_mongo else 0)
        with col3:
                total_size = sum(doc.get('content_length', 0) for doc in (user_docs or []))
                st.metric("Total Size", f"{total_size:,} chars")


if __name__ == "__main__":
    main()
